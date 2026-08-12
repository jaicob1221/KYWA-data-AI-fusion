import streamlit as st
import streamlit.components.v1 as components
import os
import ssl
import json
import hashlib
import time
import re
from datetime import datetime
from pathlib import Path
from io import BytesIO
from streamlit_autorefresh import st_autorefresh

ADMIN_USER = "yjchoi"

os.environ["CURL_SSL_NO_REVOKE"] = "1"
ssl._create_default_https_context = ssl._create_unverified_context

st.set_page_config(page_title="KYWA 데이터 융복합 서비스", page_icon="🔗", layout="wide")

ASSETS = Path("assets")
USER_DB = Path("users.json")
USAGE_DB = Path("usage_stats.json")
FEEDBACK_DB = Path("qa_feedback.json")
REGS_DIR = Path("regs")
ERP_DIR = Path("regs_erp")
VECTOR_DIR = Path("vector_db")
CACHE_DIR = Path("cache")
CERT_CACHE_FILE = CACHE_DIR / "cert_programs.json.gz"
CERT_CACHE_TTL_HOURS = 24  # 파일 캐시 유효 시간(시간)
YOUTH_CACHE_FILE = CACHE_DIR / "youth_programs.json.gz"
YOUTH_CACHE_TTL_DAYS = 30  # 청소년활동 캐시 유효(일)
GGOMGIL_CSV = CACHE_DIR / "ggomgil_programs.csv"  # 꿈길 진로체험프로그램
CN_MAJOR_CACHE = CACHE_DIR / "careernet_majors.json"
CN_SCHOOL_CACHE = CACHE_DIR / "careernet_schools.json"
CN_COSE_CACHE = CACHE_DIR / "careernet_cose.json"
CN_CACHE_META = CACHE_DIR / "careernet_cache_meta.json"
CN_CACHE_TTL_DAYS = 30
JINSOL_CACHE = CACHE_DIR / "jinsol_solutions.json"
JINSOL_META = CACHE_DIR / "jinsol_meta.json"
SESSION_TIMEOUT_MIN = 60  # 세션 1시간

REGS_DIR.mkdir(exist_ok=True)
ERP_DIR.mkdir(exist_ok=True)
VECTOR_DIR.mkdir(exist_ok=True)
CACHE_DIR.mkdir(exist_ok=True)

st.markdown("""
<style>
button[kind="header"], button[data-testid="baseButton-header"],
[data-testid="stSidebarCollapsedControl"], [data-testid="collapsedControl"] {
    display: flex !important; visibility: visible !important; opacity: 1 !important;
    color: #e5eefb !important; background: rgba(30, 41, 59, 0.95) !important;
    border: 1px solid rgba(148, 163, 184, 0.35) !important; border-radius: 10px !important;
    z-index: 999999 !important;
}
header[data-testid="stHeader"] { background: transparent !important; border-bottom: none !important; box-shadow: none !important; }
html, body, [class*="css"] { font-family: "Pretendard", "Noto Sans KR", "Malgun Gothic", sans-serif; }
.stApp { background: radial-gradient(circle at top, #1e293b 0%, #0b1220 45%, #020617 100%); color: #e5eefb; }
.block-container { padding-top: 1.2rem; padding-bottom: 2rem; max-width: 1120px; }
div[data-testid="stDecoration"] { display: none !important; }
section[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #020617 0%, #0f172a 55%, #111827 100%) !important;
    border-right: 1px solid rgba(96, 165, 250, 0.25) !important;
}
section[data-testid="stSidebar"] .stRadio label {
    display: block !important; padding: 0.62rem 0.75rem !important; margin: 0.22rem 0 !important;
    border-radius: 12px !important; border: 1px solid rgba(148, 163, 184, 0.18) !important;
    background: rgba(15, 23, 42, 0.9) !important; color: #f8fafc !important; font-weight: 600 !important;
}
section[data-testid="stSidebar"] .stRadio label:hover {
    background: rgba(37, 99, 235, 0.22) !important; border-color: rgba(96, 165, 250, 0.55) !important;
}
section[data-testid="stSidebar"] .stRadio div[role="radiogroup"] label:has(input:checked) {
    background: linear-gradient(90deg, rgba(37,99,235,0.45), rgba(29,78,216,0.35)) !important;
    border-color: #60a5fa !important;
}
section[data-testid="stSidebar"] .stRadio > label {
    display: none !important; height: 0 !important; margin: 0 !important; padding: 0 !important; overflow: hidden !important;
}
section[data-testid="stSidebar"] .stRadio div[role="radiogroup"] label {
    display: block !important; height: auto !important; overflow: visible !important;
}
section[data-testid="stSidebar"] .stButton > button { min-height: 38px !important; width: 100%; }
.stTextInput input, .stTextArea textarea, input[type="password"] {
    background: #ffffff !important; color: #0f172a !important;
    border: 1px solid #94a3b8 !important; border-radius: 10px !important;
}
.stTextInput input::placeholder, .stTextArea textarea::placeholder {
    color: #64748b !important; opacity: 1 !important;
}
/* 진로맵 등 채팅 입력창: 흰 배경 + 진한 글자 */
[data-testid="stChatInput"],
[data-testid="stChatInput"] > div,
[data-testid="stChatInput"] textarea,
[data-testid="stChatInputTextArea"],
[data-testid="stChatInputTextArea"] textarea,
div[data-testid="stBottomBlockContainer"] textarea,
div[data-testid="stChatInputContainer"] textarea,
.stChatInput textarea,
.stChatFloatingInputContainer textarea {
    background: #ffffff !important;
    color: #0f172a !important;
    -webkit-text-fill-color: #0f172a !important;
    caret-color: #0f172a !important;
    border: 1px solid #94a3b8 !important;
    border-radius: 12px !important;
}
[data-testid="stChatInput"] textarea::placeholder,
[data-testid="stChatInputTextArea"] textarea::placeholder,
div[data-testid="stBottomBlockContainer"] textarea::placeholder,
.stChatInput textarea::placeholder {
    color: #64748b !important;
    -webkit-text-fill-color: #64748b !important;
    opacity: 1 !important;
}
/* 채팅 입력 하단 바 전체 */
div[data-testid="stBottomBlockContainer"],
div[data-testid="stChatInputContainer"],
[data-testid="stChatInput"] {
    background: rgba(15, 23, 42, 0.96) !important;
}
div[data-testid="stBottomBlockContainer"] * {
    color: inherit;
}
div[data-baseweb="select"] > div {
    background: #ffffff !important; color: #0f172a !important; border: 1px solid #94a3b8 !important;
}
div[data-baseweb="select"] *, ul[role="listbox"] *, li[role="option"], li[role="option"] *,
div[data-baseweb="popover"] li, div[data-baseweb="popover"] li * { color: #0f172a !important; }
ul[role="listbox"], div[data-baseweb="popover"] { background: #ffffff !important; }
[data-testid="stFileUploader"] section {
    background: rgba(15, 23, 42, 0.6) !important;
    border: 1px dashed rgba(148, 163, 184, 0.35) !important;
    border-radius: 12px !important; padding: 0.75rem !important;
}
[data-testid="stFileUploader"] section,
[data-testid="stFileUploader"] section > div,
[data-testid="stFileUploader"] button {
    background: #0f172a !important; color: #e5eefb !important;
    border-color: rgba(148, 163, 184, 0.35) !important;
}
[data-testid="stFileUploader"] section * { color: #e5eefb !important; }
.main-title {
    font-size: 3.4rem; font-weight: 800; letter-spacing: -0.03em; text-align: center; line-height: 1.2;
    background: linear-gradient(90deg, #93c5fd, #e2e8f0 55%, #bfdbfe);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent; margin-bottom: 0.25rem;
}
.sub-title { color: #94a3b8 !important; font-size: 1.35rem; font-weight: 600; text-align: center; margin-bottom: 0.7rem; }
.section-line {
    height: 1px; margin: 0.8rem 0 1rem 0;
    background: linear-gradient(90deg, rgba(96,165,250,0), rgba(96,165,250,0.45), rgba(96,165,250,0));
}
.feature-card {
    background: linear-gradient(180deg, rgba(30,41,59,0.95), rgba(15,23,42,0.92));
    border: 1px solid rgba(148, 163, 184, 0.18); border-radius: 16px; padding: 1.1rem 1rem; min-height: 120px;
}
.feature-card h4 { margin: 0 0 0.45rem 0; }
.feature-card p { margin: 0; color: #cbd5e1 !important; font-size: 0.9rem; }
.badge-ok {
    display: inline-block; padding: 0.25rem 0.6rem; border-radius: 999px;
    background: rgba(16, 185, 129, 0.15); border: 1px solid rgba(16, 185, 129, 0.35);
    color: #6ee7b7 !important; font-size: 0.82rem; font-weight: 600;
}
.page-header {
    background: linear-gradient(180deg, rgba(30,41,59,0.75), rgba(15,23,42,0.55));
    border: 1px solid rgba(148, 163, 184, 0.18);
    border-radius: 16px; padding: 1rem 1.15rem; margin-bottom: 1rem;
}
.page-header h2 { margin: 0 0 0.25rem 0; font-size: 1.55rem; color: #f8fafc !important; }
.page-header p { margin: 0; color: #94a3b8 !important; font-size: 0.92rem; }
.result-card {
    background: linear-gradient(180deg, rgba(30,41,59,0.95), rgba(15,23,42,0.92));
    border: 1px solid rgba(96, 165, 250, 0.28);
    border-radius: 16px; padding: 1rem 1.15rem; margin-top: 0.8rem; margin-bottom: 0.8rem;
}
.esult-card h3, .result-card h3 { margin: 0 0 0.6rem 0; font-size: 1.1rem; color: #93c5fd !important; }
label, label p, label span, div[data-testid="stWidgetLabel"] p, div[data-testid="stWidgetLabel"] span,
.stMarkdown p, .stCaption, [data-testid="stCaption"] { color: #e5eefb !important; }
button[data-baseweb="tab"], button[data-baseweb="tab"] * { color: #f8fafc !important; font-weight: 700 !important; }
button[data-baseweb="tab"][aria-selected="true"], button[data-baseweb="tab"][aria-selected="true"] * { color: #93c5fd !important; }
.stButton > button {
    min-height: 42px !important; border-radius: 12px !important; font-weight: 700 !important;
    background: #1e293b; color: #e5eefb; border: 1px solid rgba(148, 163, 184, 0.2);
}
.stButton > button[kind="primary"] {
    background: linear-gradient(90deg, #3b82f6, #2563eb) !important; color: #fff !important; border: none !important;
}
div[data-testid="stAlert"] {
    border-radius: 14px !important; background: rgba(15, 23, 42, 0.9);
    border: 1px solid rgba(148, 163, 184, 0.18);
}
.src-chip {
    display: inline-block; background: rgba(37, 99, 235, 0.25);
    border: 1px solid rgba(96, 165, 250, 0.45); color: #93c5fd !important;
    border-radius: 8px; padding: 2px 8px; margin: 2px 4px 2px 0; font-size: 0.8rem;
}
.example-box {
    background: rgba(15, 23, 42, 0.85); border-left: 3px solid #3b82f6;
    padding: 0.55rem 0.85rem; margin: 0.35rem 0; border-radius: 8px;
    color: #cbd5e1 !important; font-size: 0.9rem;
}

/* 진로맵 결과 계층·카드 */
.cm-h1 {
  font-size: 1.55rem !important; font-weight: 800 !important; color: #7dd3fc !important;
  margin: 1.45rem 0 0.7rem 0 !important; padding: 0.4rem 0 0.4rem 0.85rem !important;
  border-left: 5px solid #38bdf8 !important; letter-spacing: -0.02em;
}
.cm-h2 {
  font-size: 1.05rem !important; font-weight: 600 !important; color: #a5b4fc !important;
  margin: 0.85rem 0 0.4rem 0.7rem !important;
}
.cm-h3 {
  font-size: 0.95rem !important; font-weight: 600 !important; color: #e2e8f0 !important;
  margin: 0.3rem 0 0.2rem 0 !important;
}
.cm-body {
  font-size: 0.9rem !important; color: #cbd5e1 !important; line-height: 1.55 !important;
  margin: 0.15rem 0 0.35rem 0 !important;
}
.cm-meta { font-size: 0.78rem !important; color: #94a3b8 !important; margin: 0.15rem 0 !important; }
.cm-card {
  background: linear-gradient(180deg, #1e293b 0%, #0f172a 100%) !important;
  border: 1px solid #334155 !important; border-radius: 14px !important;
  padding: 0.95rem 1.15rem !important; margin: 0.45rem 0 0.7rem 0.5rem !important;
  box-shadow: 0 4px 14px rgba(0,0,0,0.25);
}
.cm-card-title { font-size: 0.98rem !important; font-weight: 650 !important; color: #f8fafc !important; margin: 0 0 0.35rem 0 !important; }
.cm-indent { margin-left: 0.75rem !important; }
.cm-dot { color: #64748b; margin-right: 0.35rem; }
/* 진로 분석 카드 안 마크다운: 메인 타이틀보다 작게 */
.cm-card h1, .cm-card h2, .cm-card h3,
div[data-testid="stMarkdownContainer"] .cm-card h1,
div[data-testid="stMarkdownContainer"] .cm-card h2 {
  font-size: 1.0rem !important; font-weight: 600 !important; color: #e2e8f0 !important;
  margin: 0.65rem 0 0.35rem 0 !important; border: none !important; padding: 0 !important;
}
.cm-card h1 { font-size: 1.08rem !important; color: #bae6fd !important; }
.cm-card p, .cm-card li {
  font-size: 0.9rem !important; color: #cbd5e1 !important; line-height: 1.55 !important;
}
.cm-analysis-box {
  background: linear-gradient(180deg, #1e293b 0%, #0f172a 100%) !important;
  border: 1px solid #334155 !important; border-radius: 14px !important;
  padding: 1rem 1.2rem !important; margin: 0.45rem 0 0.7rem 0.5rem !important;
}

/* 로그인/회원가입 폼·버튼 가시성 */
div[data-testid="stForm"] {
  background: #1e293b !important;
  border: 1px solid #475569 !important;
  border-radius: 14px !important;
  padding: 1rem 1.1rem 1.15rem !important;
}
div[data-testid="stForm"] label p {
  color: #e2e8f0 !important;
  font-weight: 600 !important;
}
/* 폼 제출 버튼 기본 */
div[data-testid="stForm"] button[kind="secondaryFormSubmit"],
div[data-testid="stForm"] button[kind="primaryFormSubmit"],
div[data-testid="stForm"] button {
  background: #2563eb !important;
  color: #ffffff !important;
  border: 2px solid #93c5fd !important;
  font-weight: 700 !important;
  font-size: 1.05rem !important;
  min-height: 2.6rem !important;
  border-radius: 10px !important;
}
div[data-testid="stForm"] button:hover {
  background: #1d4ed8 !important;
  border-color: #bfdbfe !important;
  color: #fff !important;
}
/* 회원가입 모드 강조용 클래스 대체: 두 번째 폼 버튼 초록 */
form[action*="signup"] button,
div[data-testid="stForm"]:has(button[kind*="FormSubmit"]) button {
  background: #2563eb !important;
}

</style>
""", unsafe_allow_html=True)


def hash_pw(pw: str) -> str:
    return hashlib.sha256(pw.encode("utf-8")).hexdigest()

def _seed_admin_user(users: dict) -> dict:
    """재배포 후에도 secrets의 관리자 비밀번호로 yjchoi 복구"""
    try:
        admin_pw = (st.secrets.get("ADMIN_PASSWORD", "") or "").strip()
    except Exception:
        admin_pw = ""
    if not admin_pw:
        return users
    u = users.get(ADMIN_USER) or {}
    users[ADMIN_USER] = {
        "password": hash_pw(admin_pw),
        "name": ADMIN_USER,
        "approved": True,
        "role": "admin",
        **{k: v for k, v in u.items() if k not in ("password", "approved", "role")},
    }
    return users


def normalize_users(users: dict) -> dict:
    """구버전 users.json에 approved 등이 없어도 표준 필드로 맞춤"""
    out = {}
    for uid, info in (users or {}).items():
        if not uid:
            continue
        if not isinstance(info, dict):
            # 아주 옛 형식: {"id": "해시비번"} 만 있는 경우
            info = {"password": str(info)}
        is_admin = (uid == ADMIN_USER) or (info.get("role") == "admin")
        approved = info.get("approved", None)
        if approved is None:
            # 필드가 없으면: 관리자만 승인, 나머지는 미승인
            approved = True if is_admin else False
        else:
            approved = bool(approved)
        out[uid] = {
            "password": info.get("password") or "",
            "name": info.get("name") or uid,
            "approved": True if is_admin else approved,
            "role": "admin" if is_admin else (info.get("role") or "user"),
        }
    return out


def load_users():
    users = {}
    if USER_DB.exists():
        try:
            with open(USER_DB, "r", encoding="utf-8") as f:
                users = json.load(f) or {}
        except Exception:
            users = {}
    # secrets에 USERS_JSON 백업이 있으면 병합 (재배포 복구용)
    try:
        raw = st.secrets.get("USERS_JSON", "")
        if raw:
            if isinstance(raw, str) and raw.strip():
                backup = json.loads(raw)
            elif isinstance(raw, dict):
                backup = raw
            else:
                backup = {}
            for uid, info in (backup or {}).items():
                if uid not in users:
                    users[uid] = info
    except Exception:
        pass
    users = normalize_users(users)
    users = _seed_admin_user(users)
    return users


def save_users(users):
    users = normalize_users(dict(users or {}))
    users = _seed_admin_user(users)
    USER_DB.parent.mkdir(parents=True, exist_ok=True)
    with open(USER_DB, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)

def load_usage():
    if not USAGE_DB.exists():
        return {"total_logins": 0, "menu_counts": {}}
    try:
        with open(USAGE_DB, "r", encoding="utf-8") as f:
            data = json.load(f)
        data.setdefault("total_logins", 0)
        data.setdefault("menu_counts", {})
        return data
    except Exception:
        return {"total_logins": 0, "menu_counts": {}}

def save_usage(data):
    with open(USAGE_DB, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def inc_login_count():
    data = load_usage()
    data["total_logins"] = int(data.get("total_logins", 0)) + 1
    save_usage(data)

def inc_menu_count(menu_name: str):
    data = load_usage()
    counts = data.setdefault("menu_counts", {})
    counts[menu_name] = int(counts.get(menu_name, 0)) + 1
    save_usage(data)

def touch_session():
    st.session_state.last_active = time.time()


# ---------- 진로맵: 꿈길 CSV + 여가부 청소년활동 검색 ----------
def _ggomgil_col(df, *candidates, partial: bool = True):
    """컬럼명 공백/인코딩 차이 대응 — 정확일치 우선, 필요 시 부분일치"""
    cols = list(df.columns)
    stripped = {str(c).strip(): c for c in cols}
    # 1) 정확 일치
    for cand in candidates:
        if cand in df.columns:
            return cand
        if cand in stripped:
            return stripped[cand]
    # 2) 부분일치 (짧은 후보가 NO 등을 잡지 않도록 길이 제한)
    if partial:
        for cand in candidates:
            if len(str(cand)) < 3:
                continue
            for c in cols:
                cs = str(c).strip()
                if cs == cand or cand in cs:
                    return c
    return None


def load_ggomgil_df():
    """꿈길 진로체험프로그램 CSV 로드"""
    import pandas as pd
    if not GGOMGIL_CSV.exists():
        return None
    try:
        df = None
        for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
            try:
                df = pd.read_csv(GGOMGIL_CSV, encoding=enc)
                break
            except Exception:
                continue
        if df is None or df.empty:
            return None
        # 컬럼명 정리
        df.columns = [str(c).strip().replace("\ufeff", "") for c in df.columns]
        lat_c = _ggomgil_col(df, "위도")
        lng_c = _ggomgil_col(df, "경도")
        if lat_c:
            df[lat_c] = pd.to_numeric(df[lat_c], errors="coerce")
            if lat_c != "위도":
                df["위도"] = df[lat_c]
        if lng_c:
            df[lng_c] = pd.to_numeric(df[lng_c], errors="coerce")
            if lng_c != "경도":
                df["경도"] = df[lng_c]
        return df
    except Exception:
        return None


def search_ggomgil(school_level: str = "", region_kw: str = "", keywords=None, top_n: int = 12):
    """
    꿈길 프로그램 검색.
    school_level: 초/중/고
    region_kw: 체험지역명 부분일치 (예: 서울, 경기, 용인)
    keywords: 직업·프로그램 키워드 리스트
    """
    import pandas as pd
    df = load_ggomgil_df()
    if df is None or df.empty:
        return []
    m = df.copy()
    level = (school_level or "").strip()

    col_elem = _ggomgil_col(m, "초등학교대상여부", "초등대상여부")
    col_mid = _ggomgil_col(m, "중학교대상여부", "중학대상여부")
    col_high = _ggomgil_col(m, "고등학교대상여부", "고등대상여부")
    col_region = _ggomgil_col(m, "체험지역명")
    col_name = _ggomgil_col(m, "체험프로그램명", "프로그램명")
    col_job = _ggomgil_col(m, "체험프로그램 직업유형", "직업유형")
    col_type = _ggomgil_col(m, "체험유형")
    col_place = _ggomgil_col(m, "체험처명")
    # 필수 컬럼 없으면 검색 불가
    if not col_name:
        return []

    try:
        if level.startswith("초") and col_elem:
            m = m[m[col_elem].astype(str).str.upper().str.contains("Y", na=False)]
        elif level.startswith("중") and col_mid:
            m = m[m[col_mid].astype(str).str.upper().str.contains("Y", na=False)]
        elif level.startswith("고") and col_high:
            m = m[m[col_high].astype(str).str.upper().str.contains("Y", na=False)]
    except Exception:
        pass  # 교급 필터 실패 시 전체에서 검색

    if region_kw and col_region:
        r = str(region_kw).strip()
        if r:
            try:
                m = m[m[col_region].astype(str).str.contains(r, na=False, regex=False)]
            except Exception:
                pass

    kws = keywords or []
    if isinstance(kws, str):
        kws = [kws]
    kws = expand_interest_keywords(kws)
    kws = [str(k).strip() for k in kws if str(k).strip()]
    negs = interest_negative_terms(keywords or kws)
    # 키워드가 있으면 반드시 제목·직업유형 등에 매칭 (없으면 빈 결과)
    if kws:
        mask = None
        for k in kws:
            part = None
            for col in (col_name, col_job, col_type, col_place):
                if not col or col not in m.columns:
                    continue
                try:
                    p = m[col].astype(str).str.contains(k, case=False, na=False, regex=False)
                    part = p if part is None else (part | p)
                except Exception:
                    continue
            if part is not None:
                mask = part if mask is None else (mask | part)
        if mask is not None:
            try:
                m = m[mask]
            except Exception:
                pass
        else:
            return []
    else:
        return []

    if m is None or m.empty:
        return []
    try:
        import pandas as pd
        # 점수 순 정렬 (부정어 감점)
        scored_rows = []
        for idx, row in m.iterrows():
            blob = " ".join(
                str(row.get(c, "") or "")
                for c in (col_name, col_job, col_type, col_place)
                if c
            )
            sc = sum(2 for k in kws if k and k in blob)
            sc -= sum(3 for n in negs if n and n in blob)
            if sc > 0:
                scored_rows.append((sc, idx))
        scored_rows.sort(key=lambda x: -x[0])
        if not scored_rows:
            return []
        m = m.loc[[i for _, i in scored_rows[: max(top_n * 3, top_n)]]]
        out = []
        for _, row in m.head(top_n).iterrows():
            def _v(col):
                if not col or col not in m.columns:
                    return ""
                val = row.get(col, "")
                try:
                    if val is None or (isinstance(val, float) and pd.isna(val)):
                        return ""
                except Exception:
                    pass
                return str(val).strip()

            lat_c = _ggomgil_col(m, "위도")
            lng_c = _ggomgil_col(m, "경도")
            lat = lng = None
            try:
                if lat_c and pd.notna(row.get(lat_c)):
                    lat = float(row[lat_c])
            except Exception:
                pass
            try:
                if lng_c and pd.notna(row.get(lng_c)):
                    lng = float(row[lng_c])
            except Exception:
                pass
            title = _v(col_name)
            # NO·숫자만 있는 행은 잘못된 컬럼 매핑으로 보고 제외
            if not title or title.upper() == "NO" or title.isdigit():
                continue
            out.append({
                "체험프로그램명": title,
                "체험처명": _v(col_place),
                "체험처유형": _v(_ggomgil_col(m, "체험처유형")),
                "체험프로그램 직업유형": _v(col_job),
                "체험유형": _v(col_type),
                "체험지역명": _v(col_region),
                "체험진행장소": _v(_ggomgil_col(m, "체험진행장소")),
                "유무료구분": _v(_ggomgil_col(m, "유무료구분")),
                "유료금액": _v(_ggomgil_col(m, "유료금액")),
                "체험시간": _v(_ggomgil_col(m, "체험시간")),
                "체험인원": _v(_ggomgil_col(m, "체험인원")),
                "대면비대면구분": _v(_ggomgil_col(m, "대면비대면구분")),
                "안전도": _v(_ggomgil_col(m, "안전도")),
                "위도": lat,
                "경도": lng,
            })
        return out
    except Exception:
        return []


def classify_ggomgil_bucket(row: dict) -> str:
    """진로체험 vs 진로교육 구분"""
    t = f"{row.get('체험유형') or ''} {row.get('체험프로그램명') or ''} {row.get('체험진행장소') or ''}"
    edu_keys = ("교육", "강의", "수업", "학습", "진로교육", "특강", "워크숍", "워크샵")
    if any(k in t for k in edu_keys) and "체험" not in (row.get("체험유형") or ""):
        return "교육"
    if any(k in t for k in edu_keys) and any(k in t for k in ("체험", "실습", "탐방")):
        # 혼합이면 체험 우선
        return "체험"
    if any(k in t for k in edu_keys):
        return "교육"
    return "체험"


def region_matches_user(item_region: str, user_region: str) -> bool:
    """사용자 시도/도를 벗어나면 False (느슨한 포함 매칭)"""
    u = (user_region or "").strip()
    if not u:
        return True  # 지역 미상 시 필터 안 함
    r = (item_region or "").strip()
    if not r:
        return False  # 지역 정보 없는 항목은 배제
    # 공통 시도 키
    keys = [
        "서울", "부산", "대구", "인천", "광주", "대전", "울산", "세종",
        "경기", "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주",
        "고양", "성남", "용인", "수원", "부천", "안양", "화성", "남양주",
    ]
    u_hits = [k for k in keys if k in u]
    if not u_hits:
        return u in r or r in u
    return any(k in r for k in u_hits)


def filter_items_by_keywords(items, keywords, text_keys=None, min_score: int = 1, limit: int = 12):
    """목록형 결과에서 키워드 관련도 낮은 항목 제거 (무관 폴백 없음)"""
    if not items:
        return []
    kws = expand_interest_keywords(keywords)
    kws = [str(k).strip() for k in kws if str(k).strip()]
    negs = interest_negative_terms(keywords)
    if not kws:
        return []
    text_keys = text_keys or []
    scored = []
    for it in items:
        if not isinstance(it, dict):
            continue
        if text_keys:
            blob = " ".join(str(it.get(k) or "") for k in text_keys)
        else:
            blob = " ".join(str(v) for v in it.values() if v is not None)
        score = sum(1 for k in kws if k and k.lower() in blob.lower())
        score -= sum(2 for n in negs if n and n in blob)
        if score >= min_score:
            scored.append((score, it))
    scored.sort(key=lambda x: -x[0])
    return [x for _, x in scored[:limit]]


def fetch_youth_singo_programs(
    service_key: str,
    sido: str = "",
    sigungu: str = "",
    pgm: str = "",
    org: str = "",
    page_no: int = 1,
    num_of_rows: int = 20,
    sdate: str = "",
    edate: str = "",
    endpoint_mode: str = "auto",
):
    """
    청소년활동 프로그램 검색 — 여러 공공 API 후보를 순차 시도
    1) YouthProgramSearchService (사용자 지정)
    2) 자원봉사 프로그램 getVolProgrmList
    3) 기타 yhis / YouthActivInfo 계열
    """
    import requests
    from urllib.parse import unquote
    import xml.etree.ElementTree as ET
    from datetime import datetime, timedelta

    key = (service_key or "").strip()
    if "%" in key:
        try:
            key = unquote(key)
        except Exception:
            pass
    if not key:
        return [], {"error": "DATA_GO_KR_KEY 없음"}

    def _parse(text: str):
        items, meta = [], {}
        if not text or not str(text).strip():
            return items, {"empty": True}
        t = str(text).strip()
        try:
            if t.startswith("{") or t.startswith("["):
                data = json.loads(t)
                body = data.get("response", data.get("body", data))
                if isinstance(body, dict):
                    header = body.get("header") or data.get("header") or {}
                    meta["resultCode"] = header.get("resultCode")
                    meta["resultMsg"] = header.get("resultMsg")
                    b = body.get("body") or body
                    it = (b.get("items") or {}) if isinstance(b, dict) else {}
                    raw = it.get("item") if isinstance(it, dict) else it
                    if isinstance(raw, list):
                        items = raw
                    elif isinstance(raw, dict):
                        items = [raw]
                    if isinstance(b, dict) and b.get("totalCount") is not None:
                        meta["totalCount"] = b.get("totalCount")
            else:
                root = ET.fromstring(t)
                for tag in ("resultCode", "resultMsg", "totalCount"):
                    el = root.find(f".//{tag}")
                    if el is not None and el.text:
                        meta[tag] = el.text
                for item in root.findall(".//item"):
                    items.append({c.tag: (c.text or "") for c in list(item)})
        except Exception as e:
            meta["parse_error"] = str(e)
            meta["preview"] = t[:400]
        return items, meta

    # edate 필수인 구 API용
    edate = (datetime.now() + timedelta(days=365)).strftime("%Y%m%d")
    sdate = (datetime.now() - timedelta(days=30)).strftime("%Y%m%d")

    # End Point (bulk는 search_only 권장 — Vol API는 건수 적어 조기종료됨)
    primary = (
        "https://apis.data.go.kr/1383000/yhis/YouthProgramSearchService/getYouthProgramSearchList"
    )
    singo = "https://apis.data.go.kr/1383000/YouthActivInfoSingoSrvc/getSingoProgrmList"
    vol = "https://apis.data.go.kr/1383000/YouthActivInfoVolSrvc/getVolProgrmList"
    mode = (endpoint_mode or "auto").strip().lower()
    if mode in ("search", "search_only", "primary"):
        candidates = [(primary, {})]
    elif mode in ("singo", "singo_only"):
        candidates = [(singo, {})]
    elif mode in ("vol", "vol_only"):
        candidates = [(vol, {})]
    else:
        candidates = [(primary, {}), (singo, {}), (vol, {})]

    last_meta = {}
    for url, extra in candidates:
        for type_fmt in ("json", "xml"):
            params = {
                "serviceKey": key,
                "pageNo": str(page_no),
                "numOfRows": str(num_of_rows),
                "type": type_fmt,
            }
            params.update(extra)
            if sido:
                params["sido"] = sido
            if sigungu:
                params["sigungu"] = sigungu
            if pgm:
                params["pgm"] = pgm
                params["keyword"] = pgm
                params["progrmNm"] = pgm
                params["atName"] = pgm
            if org:
                params["org"] = org
                params["organNm"] = org
                params["orgName"] = org
            if sdate:
                params["sdate"] = sdate
                params["startDate"] = sdate
            if edate:
                params["edate"] = edate
                params["endDate"] = edate
            try:
                res = requests.get(url, params=params, timeout=45, verify=False)
                text = res.text or ""
                items, meta = _parse(text)
                meta["http_status"] = res.status_code
                meta["url"] = url
                meta["type"] = type_fmt
                code = str(meta.get("resultCode") or "")
                if items:
                    meta["hit"] = len(items)
                    return items, meta
                if code in ("00", "0", "000") and meta.get("totalCount") in ("0", 0):
                    last_meta = meta
                    continue
                last_meta = meta
                last_meta["preview"] = text[:250]
            except Exception as e:
                last_meta = {"url": url, "error": str(e), "type": type_fmt}
                continue

    return [], last_meta or {"error": "모든 청소년활동 API 후보 실패"}



def get_youth_api_key() -> str:
    dummy = {"0000", "0", "test", "YOUR_KEY", "changeme", ""}
    try:
        y_raw = (st.secrets.get("YOUTH_ACTIV_KEY", "") or "").strip()
        d_raw = (st.secrets.get("DATA_GO_KR_KEY", "") or "").strip()
        if y_raw and y_raw not in dummy:
            return y_raw
        if d_raw and d_raw not in dummy:
            return d_raw
    except Exception:
        pass
    return ""


def load_youth_file_cache(ttl_days: int = None):
    """청소년활동 전체 캐시 로드. 유효하면 items, meta / 아니면 None, meta"""
    import gzip
    from datetime import datetime, timezone
    ttl_days = YOUTH_CACHE_TTL_DAYS if ttl_days is None else ttl_days
    meta = {"path": str(YOUTH_CACHE_FILE), "ttl_days": ttl_days}
    if not YOUTH_CACHE_FILE.exists():
        meta["status"] = "missing"
        return None, meta
    try:
        with gzip.open(YOUTH_CACHE_FILE, "rt", encoding="utf-8") as f:
            payload = json.load(f)
        items = payload.get("items") or []
        saved_at = payload.get("saved_at") or ""
        meta["saved_at"] = saved_at
        meta["count"] = len(items)
        meta["last_page"] = payload.get("last_page") or 0
        meta["complete"] = bool(payload.get("complete", False))
        meta["num_of_rows"] = payload.get("num_of_rows")
        meta["totalCount"] = payload.get("totalCount")
        try:
            if isinstance(saved_at, (int, float)):
                saved_ts = float(saved_at)
            else:
                saved_ts = datetime.fromisoformat(str(saved_at).replace("Z", "+00:00")).timestamp()
            age_d = (time.time() - saved_ts) / 86400.0
            meta["age_days"] = round(age_d, 2)
            if age_d > ttl_days:
                meta["status"] = "expired"
                # 만료여도 일단 데이터는 반환 (매칭용) — 상태만 expired
                return items, meta
        except Exception as e:
            meta["date_error"] = str(e)
        if not items:
            meta["status"] = "empty"
            return None, meta
        meta["status"] = meta.get("status") or "ok"
        return items, meta
    except Exception as e:
        meta["status"] = f"load_error:{e}"
        return None, meta


def _youth_item_key(it: dict) -> str:
    return (
        str(it.get("prgrmNm") or it.get("pgmNm") or it.get("atName") or it.get("programNm") or "")
        + "|"
        + str(it.get("fcltyNm") or it.get("organNm") or it.get("orgName") or it.get("operInstNm") or "")
        + "|"
        + str(it.get("actvBgngYmd") or it.get("sdate") or it.get("certNo") or "")
    )


def save_youth_file_cache(items: list, extra_meta: dict = None):
    import gzip
    from datetime import datetime, timezone
    CACHE_DIR.mkdir(exist_ok=True)
    payload = {
        "saved_at": datetime.now(timezone.utc).isoformat(),
        "count": len(items or []),
        "items": items or [],
    }
    if extra_meta:
        payload.update(extra_meta)
    tmp = YOUTH_CACHE_FILE.with_suffix(".tmp.gz")
    with gzip.open(tmp, "wt", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False)
    tmp.replace(YOUTH_CACHE_FILE)
    return {
        "path": str(YOUTH_CACHE_FILE),
        "count": payload["count"],
        "saved_at": payload["saved_at"],
        "last_page": payload.get("last_page"),
        "complete": payload.get("complete"),
    }


def fetch_youth_programs_all(
    service_key: str,
    num_of_rows: int = 100,
    max_pages: int = 2000,
    progress_cb=None,
    start_page: int = 1,
    existing_items: list = None,
    checkpoint_every: int = 10,
    sdate: str = "",
    edate: str = "",
    endpoint_mode: str = "search_only",
    by_sido: bool = True,
):
    """
    청소년활동 전체 수집.
    - 전국 단일 조회는 API가 ~수천 건에서 막히거나 pageNo를 무시하는 경우가 많음
    - by_sido=True 이면 시도별로 나눠 조회해 한도를 우회
    """
    SIDO_LIST = [
        "서울특별시", "부산광역시", "대구광역시", "인천광역시", "광주광역시",
        "대전광역시", "울산광역시", "세종특별자치시", "경기도", "강원특별자치도",
        "강원도", "충청북도", "충청남도", "전북특별자치도", "전라북도",
        "전라남도", "경상북도", "경상남도", "제주특별자치도", "제주도",
    ]

    all_items = list(existing_items or [])
    seen = set()
    for it in all_items:
        if isinstance(it, dict):
            seen.add(_youth_item_key(it))
    last_meta = {}
    num_of_rows = max(10, min(int(num_of_rows or 100), 1000))
    max_pages = max(1, min(int(max_pages or 1), 5000))
    complete = False
    global_page = 0

    def _one_scope(sido_name: str, pages_limit: int):
        nonlocal global_page, last_meta, complete
        empty_streak = 0
        dup_streak = 0
        total = None
        scope_new = 0
        first_page_fp = None
        for page in range(1, pages_limit + 1):
            global_page += 1
            part, meta = fetch_youth_singo_programs(
                service_key,
                sido=sido_name or "",
                pgm="",
                page_no=page,
                num_of_rows=num_of_rows,
                sdate=sdate or "",
                edate=edate or "",
                endpoint_mode=endpoint_mode or "search_only",
            )
            last_meta = dict(meta or {})
            last_meta["sido"] = sido_name or "전국"
            if total is None and meta.get("totalCount") not in (None, ""):
                try:
                    total = int(str(meta.get("totalCount")).replace(",", ""))
                except Exception:
                    total = None

            if not part:
                empty_streak += 1
                if empty_streak >= 2:
                    break
                time.sleep(0.25)
                continue
            empty_streak = 0

            # pageNo 무시 여부 감지: 1페이지 지문과 동일하면 중단
            fp = "|".join(_youth_item_key(x) for x in part[:5] if isinstance(x, dict))
            if page == 1:
                first_page_fp = fp
            elif first_page_fp and fp == first_page_fp:
                last_meta["page_ignored"] = True
                break

            new_in_page = 0
            for it in part:
                if not isinstance(it, dict):
                    continue
                key = _youth_item_key(it)
                if key in seen:
                    continue
                seen.add(key)
                all_items.append(it)
                new_in_page += 1
                scope_new += 1

            if progress_cb:
                progress_cb(
                    global_page,
                    len(all_items),
                    total,
                    extra={
                        "sido": sido_name or "전국",
                        "page": page,
                        "new": new_in_page,
                        "url": (meta or {}).get("url"),
                    },
                )

            if new_in_page == 0:
                dup_streak += 1
                if dup_streak >= 2:
                    break
            else:
                dup_streak = 0

            if checkpoint_every and global_page % int(checkpoint_every) == 0:
                try:
                    save_youth_file_cache(
                        all_items,
                        extra_meta={
                            "last_page": global_page,
                            "num_of_rows": num_of_rows,
                            "totalCount": total,
                            "complete": False,
                            "sdate": sdate or "",
                            "edate": edate or "",
                            "by_sido": by_sido,
                            "last_sido": sido_name or "",
                        },
                    )
                except Exception:
                    pass

            if total is not None and scope_new >= total and sido_name:
                break
            if len(part) < num_of_rows and new_in_page == 0:
                break
            time.sleep(0.12)
        return scope_new

    if by_sido:
        # 시도별 수집 (전국 단일 한도 우회)
        pages_per_sido = max(20, min(max_pages // max(len(SIDO_LIST), 1), 200))
        for sido_name in SIDO_LIST:
            _one_scope(sido_name, pages_per_sido)
        # 시도 누락분 보완: 전국 1회 (선택)
        _one_scope("", min(50, max_pages))
        complete = True
    else:
        # 전국 단일 페이징
        start = max(1, int(start_page or 1))
        empty_streak = 0
        dup_streak = 0
        total = None
        first_page_fp = None
        for page in range(start, max_pages + 1):
            global_page = page
            part, meta = fetch_youth_singo_programs(
                service_key,
                sido="",
                pgm="",
                page_no=page,
                num_of_rows=num_of_rows,
                sdate=sdate or "",
                edate=edate or "",
                endpoint_mode=endpoint_mode or "search_only",
            )
            last_meta = dict(meta or {})
            if total is None and meta.get("totalCount") not in (None, ""):
                try:
                    total = int(str(meta.get("totalCount")).replace(",", ""))
                except Exception:
                    total = None
            if not part:
                empty_streak += 1
                if empty_streak >= 3:
                    complete = True
                    break
                continue
            empty_streak = 0
            fp = "|".join(_youth_item_key(x) for x in part[:5] if isinstance(x, dict))
            if page == start:
                first_page_fp = fp
            elif first_page_fp and fp == first_page_fp:
                last_meta["page_ignored"] = True
                complete = True
                break
            new_in_page = 0
            for it in part:
                if not isinstance(it, dict):
                    continue
                key = _youth_item_key(it)
                if key in seen:
                    continue
                seen.add(key)
                all_items.append(it)
                new_in_page += 1
            if progress_cb:
                progress_cb(page, len(all_items), total, extra={"new": new_in_page, "sido": "전국"})
            if new_in_page == 0:
                dup_streak += 1
                if dup_streak >= 3:
                    complete = True
                    break
            else:
                dup_streak = 0
            if checkpoint_every and page % int(checkpoint_every) == 0:
                try:
                    save_youth_file_cache(
                        all_items,
                        extra_meta={
                            "last_page": page,
                            "complete": False,
                            "sdate": sdate or "",
                            "edate": edate or "",
                        },
                    )
                except Exception:
                    pass
            if total is not None and len(all_items) >= total:
                complete = True
                break
            time.sleep(0.12)
        if page >= max_pages:
            complete = True

    last_meta["collected"] = len(all_items)
    last_meta["pages"] = global_page
    last_meta["complete"] = complete
    last_meta["last_page"] = global_page
    last_meta["endpoint_mode"] = endpoint_mode
    last_meta["by_sido"] = by_sido
    try:
        save_youth_file_cache(
            all_items,
            extra_meta={
                "last_page": global_page,
                "complete": complete,
                "sdate": sdate or "",
                "edate": edate or "",
                "by_sido": by_sido,
            },
        )
    except Exception:
        pass
    return all_items, last_meta


def _school_to_target_tokens(school: str) -> list:
    """교급 → 대상 연령/학년 매칭용 토큰"""
    s = str(school or "")
    tokens = []
    if any(x in s for x in ("초", "초등", "4학년", "5학년", "6학년", "3학년", "2학년", "1학년")) and "중" not in s and "고" not in s:
        tokens += ["초등", "초등학생", "어린이", "아동", "초1", "초2", "초3", "초4", "초5", "초6", "저학년", "고학년"]
    if any(x in s for x in ("중", "중학", "중학생")):
        tokens += ["중학", "중학생", "청소년", "중1", "중2", "중3"]
    if any(x in s for x in ("고", "고등", "고등학생")):
        tokens += ["고등", "고등학생", "청소년", "고1", "고2", "고3"]
    if not tokens:
        tokens += ["청소년", "초등", "중학", "학생"]
    return tokens


def _interest_type_tokens(keywords: list) -> list:
    """관심 키워드 → 유사 활동유형/영역 토큰 (프로그램명 일치가 약할 때 보완)"""
    blob = " ".join(str(k) for k in (keywords or []))
    types = []
    rules = [
        (("요리", "조리", "베이킹", "제빵", "셰프", "쿠킹", "음식", "푸드"), ["요리", "조리", "제빵", "베이킹", "음식", "급식", "바리스타", "커피", "체험"]),
        (("새", "조류", "동물", "생물", "자연", "생태", "환경", "식물"), ["자연", "생태", "환경", "동물", "탐방", "숲", "체험", "과학"]),
        (("로봇", "코딩", "AI", "인공지능", "소프트웨어", "컴퓨터", "IT"), ["로봇", "코딩", "소프트웨어", "AI", "정보", "과학", "메이커", "SW"]),
        (("미술", "그림", "디자인", "공예", "예술", "음악", "댄스", "연극"), ["미술", "예술", "공예", "문화", "음악", "공연", "체험"]),
        (("운동", "체육", "축구", "농구", "수영", "스포츠"), ["체육", "스포츠", "운동", "캠프", "체험"]),
        (("역사", "문화", "전통", "박물관"), ["역사", "문화", "전통", "탐방", "체험", "인문"]),
        (("진로", "직업", "꿈"), ["진로", "직업", "체험", "캠프"]),
    ]
    for keys, vals in rules:
        if any(k in blob for k in keys):
            types.extend(vals)
    if not types and keywords:
        types = [str(k) for k in keywords if k][:8]
    # 공통 완화 토큰
    types += ["체험", "활동", "캠프", "프로그램"]
    # 중복 제거
    out, seen = [], set()
    for t in types:
        if t and t not in seen:
            seen.add(t)
            out.append(t)
    return out


def filter_youth_cache_by_keywords(
    items: list,
    keywords: list,
    region: str = "",
    school: str = "",
    limit: int = 60,
):
    """
    캐시 1차 후보 선별 (완화 버전)
    - 관심 키워드 (프로그램명·유형·개요)
    - 동일 지역
    - 대상 연령대(교급)
    - 유사 유형
    하나라도 맞으면 후보에 포함, 점수로 정렬
    """
    kws = [str(k).strip() for k in (keywords or []) if str(k).strip()]
    try:
        kws = expand_interest_keywords(kws) if kws else []
    except Exception:
        pass
    type_tokens = _interest_type_tokens(kws)
    age_tokens = _school_to_target_tokens(school)

    region_tokens = []
    if region:
        region_tokens.append(region)
        if region.endswith("시") or region.endswith("도") or region.endswith("군") or region.endswith("구"):
            region_tokens.append(region[:-1])
        for a in ("특별시", "광역시", "특별자치시", "특별자치도", "자치시", "자치도"):
            if a in region:
                region_tokens.append(region.replace(a, "").strip())

    scored = []
    for it in items or []:
        if not isinstance(it, dict):
            continue
        info = normalize_youth_program(it)
        name = str(info.get("name") or "")
        target = str(info.get("target") or "")
        facility = str(info.get("facility") or "")
        sido = str(info.get("sido") or "")
        sgg = str(info.get("sgg") or "")
        addr = str(info.get("addr") or "")
        domain = str(it.get("actvRelmNm") or it.get("actvTypeNm") or "")
        outline = str(it.get("prgrmOtlnCn") or it.get("eduGoalCn") or "")[:200]
        blob = " ".join([name, target, facility, sido, sgg, addr, domain, outline])

        sc = 0
        reasons = []

        # 1) 관심 키워드
        kw_hit = 0
        for k in kws:
            if k and k in blob:
                kw_hit += 3 if k in name else 1
        if kw_hit:
            sc += kw_hit
            reasons.append("관심")

        # 2) 유사 유형
        type_hit = sum(1 for t in type_tokens if t and t in blob)
        if type_hit:
            sc += min(type_hit, 5)
            reasons.append("유사유형")

        # 3) 동일 지역
        loc = f"{sido} {sgg} {addr} {facility}"
        region_hit = any(t and t in loc for t in region_tokens) if region_tokens else False
        if region_hit:
            sc += 4
            reasons.append("지역")

        # 4) 대상 연령대
        age_blob = f"{target} {name} {outline}"
        age_hit = any(t and t in age_blob for t in age_tokens)
        if age_hit:
            sc += 3
            reasons.append("연령")

        # 포함 조건: 관심/유형 중 하나 + (지역 또는 연령) 이거나, 관심만 강하거나, 지역+연령
        include = False
        if kw_hit or type_hit:
            include = True
        elif region_hit and age_hit:
            include = True
        elif region_hit and sc >= 4:
            include = True

        if include and sc > 0:
            scored.append((sc, it, reasons))

    scored.sort(key=lambda x: -x[0])
    # 다양성: 상위 limit
    return [x[1] for x in scored[:limit]]


def ai_pick_youth_from_cache(
    provider, api_key, keywords: list, region: str, candidates: list, limit: int = 12, school: str = ""
):
    """AI 선별: 관심·지역·연령·유사유형을 넓게 인정"""
    if not candidates:
        return []
    if not api_key:
        return candidates[:limit]
    slim = []
    for i, it in enumerate(candidates[:60]):
        info = normalize_youth_program(it)
        slim.append({
            "i": i,
            "name": info.get("name") or "",
            "facility": info.get("facility") or "",
            "target": info.get("target") or "",
            "sido": info.get("sido") or "",
            "type": str(it.get("actvRelmNm") or it.get("actvTypeNm") or "")[:30],
            "addr": (info.get("addr") or "")[:40],
        })
    try:
        prompt = (
            f"관심키워드: {keywords}\n교급: {school or '미상'}\n지역: {region or '전국'}\n"
            f"후보활동: {json.dumps(slim, ensure_ascii=False)}\n\n"
            "선정 기준(완화):\n"
            "1) 관심과 직접 관련되거나 유사 유형(체험·캠프·문화 등)\n"
            "2) 동일·인근 지역\n"
            "3) 교급에 맞는 대상(초/중/고/청소년)\n"
            "프로그램명에 관심 단어가 없어도 유형·대상·지역이 맞으면 포함.\n"
            f'JSON만: {{"idx":[...], "reason":"한줄"}} 최대 {limit}개.'
        )
        raw = get_ai_response(
            provider,
            api_key,
            "청소년활동 추천 선별기. 관련성 기준을 완화해 실용적 추천. JSON만 출력.",
            prompt,
        )
        m = re.search(r"\{[\s\S]*\}", raw or "")
        if m:
            idxs = json.loads(m.group(0)).get("idx") or []
            out = []
            for i in idxs:
                try:
                    ii = int(i)
                    if 0 <= ii < len(candidates):
                        out.append(candidates[ii])
                except Exception:
                    pass
            if out:
                return out[:limit]
    except Exception:
        pass
    return candidates[:limit]


def naver_map_link(lat, lng, name=""):
    try:
        la, ln = float(lat), float(lng)
        q = name or f"{la},{ln}"
        from urllib.parse import quote
        return f"https://map.naver.com/v5/search/{quote(str(q))}"
    except Exception:
        return ""


def get_career_net_key():
    try:
        return (st.secrets.get("CAREER_NET_API_KEY", "") or "").strip()
    except Exception:
        return ""


def _job_text(j: dict) -> str:
    parts = [
        j.get("job_nm"), j.get("JOB_NM"), j.get("jobNm"),
        j.get("work"), j.get("JOB_SUMMARY"), j.get("job_summary"),
        j.get("top_nm"), j.get("aptit_name"), j.get("rel_job_nm"),
    ]
    return " ".join(str(p) for p in parts if p)


def is_bird_interest(keywords: list) -> bool:
    """'조류'가 새(鳥類) 의미인지 — 대화에 새/탐조 등이 있으면 새, 없으면 기본 새로 간주"""
    blob = " ".join(str(k) for k in (keywords or []))
    if any(x in blob for x in ("해양조류", "조력발전", "조력", "해류", "파력")):
        return False
    if any(x in blob for x in ("새", "탐조", "야생조류", "조류학", "새집", "철새")):
        return True
    if "조류" in blob:
        return True  # 사용자 답 '조류'는 진로맵에서 새로 해석 (해류는 해양 키워드와 같이 올 때)
    return False


def expand_interest_keywords(keywords: list) -> list:
    """관심 키워드 → 검색용 확장어. 동음이의(조류=새 vs 해류) 분리."""
    base = []
    for k in keywords or []:
        k = str(k).strip()
        if k and k not in base:
            base.append(k)
    bird = is_bird_interest(base)
    # 주제별 확장 (조류=새일 때 해양·환경공학 쪽으로 확장하지 않음)
    syn = {
        "새": ["새", "야생조류", "탐조", "철새", "조류학", "동물", "야생동물", "생물", "수의사", "동물원", "생태학"],
        "조류": (
            ["야생조류", "탐조", "철새", "조류학", "새", "동물", "생물", "수의사", "동물원"]
            if bird else
            ["조류", "조력발전", "해양", "해류", "파력"]
        ),
        "동물": ["동물", "수의사", "동물원", "야생동물", "반려", "축산", "생물", "수의학"],
        "생물": ["생물", "생명", "생명과학", "동물", "식물", "생태학"],
        "수의": ["수의사", "수의학", "동물"],
        "요리": ["요리", "조리", "셰프", "식품", "제과", "제빵", "영양"],
        "자동차": ["자동차", "기계", "정비", "모빌리티", "자동차공학"],
        "코딩": ["코딩", "프로그래밍", "소프트웨어", "개발", "컴퓨터", "AI", "데이터"],
        "그림": ["그림", "미술", "디자인", "일러스트", "애니메이션"],
        "음악": ["음악", "연주", "작곡", "음향", "가수"],
        "운동": ["운동", "체육", "스포츠", "트레이너", "선수"],
        "우주": ["우주", "항공", "천문", "물리", "과학자"],
        "로봇": ["로봇", "기계", "전자", "공학", "AI"],
        "환경": ["환경", "기후", "에너지", "재활용", "환경공학"],
    }
    out = list(base)
    for k in list(base):
        for key, vals in syn.items():
            if k == key or key in k or k in key:
                for v in vals:
                    if v not in out:
                        out.append(v)
    # 새(鳥) 관심이면 검색어에서 단독 '조류'보다 야생조류·탐조 우선
    if bird:
        out = [x for x in out if x not in ("해양", "조력발전", "해류", "파력", "모빌리티")]
        for prefer in ("야생조류", "탐조", "철새", "조류학", "수의사", "동물원", "생물"):
            if prefer not in out:
                out.insert(0, prefer)
        if "새" not in out:
            out.insert(0, "새")
    return out[:16]


def interest_negative_terms(keywords: list) -> list:
    """관심과 반대·혼동되기 쉬운 단어 (필터 감점용)"""
    if is_bird_interest(keywords):
        return [
            "해양에너지", "조력발전", "파력", "해류", "친환경건축", "건축컨설턴트",
            "창업 컨설턴트", "창업컨설턴트", "조향", "상품기획", "크라우드펀딩",
            "모빌리티", "의료코디", "환경공학", "건설환경", "ICT환경",
            "국회의원", "쇼핑", "가사도우미",
            "반도체", "로봇", "인공지능", "AI&", "AI ", "SW개발", "소프트웨어",
            "프로그래머", "데이터사이언스", "빅데이터", "정보보안",
        ]
    return ["쇼핑", "가사도우미", "경비", "청소", "호스트", "국회의원"]


def filter_jobs_by_keywords(jobs, keywords, min_score: int = 2, limit: int = 8):
    """검색어와 직업명·설명 관련도 필터. 무관·혼동 결과는 버림."""
    if not jobs:
        return []
    kws = expand_interest_keywords(keywords)
    kws = [str(k).strip() for k in kws if str(k).strip()]
    negs = interest_negative_terms(keywords)
    if not kws:
        return []
    # 핵심어(짧은 확장 전 원본) 가중
    core = [str(k).strip() for k in (keywords or []) if str(k).strip()]
    scored = []
    for j in jobs:
        t = _job_text(j)
        nm = str(j.get("job_nm") or j.get("JOB_NM") or j.get("jobNm") or "")
        score = 0
        for k in kws:
            kl = k.lower()
            if not kl:
                continue
            if kl in nm.lower():
                score += 5 if k in core or len(k) >= 2 else 4
            elif kl in t.lower():
                score += 2
        for bad in negs:
            if bad in t or bad in nm:
                score -= 5
        if score >= min_score:
            scored.append((score, j))
    scored.sort(key=lambda x: -x[0])
    return [j for _, j in scored[:limit]]


def filter_majors_by_keywords(majors, keywords, min_score: int = 2, limit: int = 10):
    """학과 캐시 결과 중 관심 무관·혼동(환경공학 등) 제거"""
    if not majors:
        return []
    kws = expand_interest_keywords(keywords)
    negs = interest_negative_terms(keywords)
    scored = []
    for m in majors:
        if not isinstance(m, dict):
            continue
        blob = " ".join(str(m.get(k) or "") for k in (
            "mClass", "facilName", "lClass", "major", "majorNm", "mjrNm", "title", "name"
        ))
        score = 0
        for k in kws:
            if k and k in blob:
                score += 3
        for bad in negs:
            if bad in blob:
                score -= 4
        # 새/동물 관심인데 환경·해양·건설만 있으면 탈락
        if is_bird_interest(keywords):
            animalish = any(x in blob for x in (
                "생물", "생명", "동물", "수의", "생태학", "자원생물", "야생", "산림", "농"
            ))
            env_only = any(x in blob for x in ("환경공학", "건설환경", "해양환경", "기후변화융합"))
            if env_only and not animalish:
                score -= 6
            if animalish:
                score += 4
        if score >= min_score:
            scored.append((score, m))
    scored.sort(key=lambda x: -x[0])
    return [m for _, m in scored[:limit]]


def fetch_careernet_jobs(api_key: str, keyword: str = "", page_index: int = 1):
    """직업백과 목록"""
    import requests
    key = (api_key or "").strip()
    if not key:
        return [], {"error": "CAREER_NET_API_KEY 없음"}
    url = "https://www.career.go.kr/cnet/front/openapi/jobs.json"
    params = {"apiKey": key, "pageIndex": page_index}
    if keyword:
        params["searchJobNm"] = keyword
    try:
        res = requests.get(url, params=params, timeout=30, verify=False)
        data = res.json() if res.text else {}
        jobs = data.get("jobs") or []
        if isinstance(jobs, dict):
            jobs = [jobs]
        return jobs, {"count": data.get("count"), "http_status": res.status_code}
    except Exception as e:
        return [], {"error": str(e)}


def fetch_careernet_junior_jobs(api_key: str, keyword: str = "", page_index: int = 1):
    """주니어직업정보 (초·중)"""
    import requests
    key = (api_key or "").strip()
    if not key:
        return [], {"error": "CAREER_NET_API_KEY 없음"}
    url = "https://www.career.go.kr/cnet/front/openapi/juniorjobsinfo.json"
    params = {"apiKey": key, "pageIndex": page_index}
    if keyword:
        params["searchJobNm"] = keyword
    try:
        res = requests.get(url, params=params, timeout=30, verify=False)
        data = res.json() if res.text else {}
        jobs = data.get("jobs") or []
        if isinstance(jobs, dict):
            jobs = [jobs]
        return jobs, {"count": data.get("count"), "http_status": res.status_code}
    except Exception as e:
        return [], {"error": str(e)}


def fetch_careernet_counsel(api_key: str, gubun: str = ""):
    """진로상담사례 목록 COUNSEL"""
    import requests
    import xml.etree.ElementTree as ET
    key = (api_key or "").strip()
    if not key:
        return [], {"error": "CAREER_NET_API_KEY 없음"}
    url = "https://www.career.go.kr/cnet/openapi/getOpenApi"
    params = {
        "apiKey": key,
        "svcType": "api",
        "svcCode": "COUNSEL",
        "contentType": "json",
    }
    if gubun:
        params["gubun"] = gubun
    try:
        res = requests.get(url, params=params, timeout=30, verify=False)
        text = res.text or ""
        items, meta = [], {"http_status": res.status_code}
        if text.strip().startswith("{") or text.strip().startswith("["):
            data = json.loads(text)
            if isinstance(data, dict):
                ds = data.get("dataSearch") or data
                content = ds.get("content") if isinstance(ds, dict) else None
                if isinstance(content, list):
                    items = content
                elif isinstance(content, dict):
                    items = [content]
                elif isinstance(data.get("content"), list):
                    items = data["content"]
            elif isinstance(data, list):
                items = data
        else:
            root = ET.fromstring(text)
            for node in root.findall(".//content"):
                row = {c.tag: (c.text or "") for c in list(node)}
                if row:
                    items.append(row)
        return items, meta
    except Exception as e:
        return [], {"error": str(e)}


def fetch_careernet_counsel_detail(api_key: str, con_cd: str):
    """상담사례 상세 COUNSEL_VIEW"""
    import requests
    import xml.etree.ElementTree as ET
    key = (api_key or "").strip()
    if not key or not con_cd:
        return {}, {"error": "키 또는 con_cd 없음"}
    url = "https://www.career.go.kr/cnet/openapi/getOpenApi"
    params = {
        "apiKey": key,
        "svcType": "api",
        "svcCode": "COUNSEL_VIEW",
        "contentType": "json",
        "con_cd": con_cd,
    }
    try:
        res = requests.get(url, params=params, timeout=30, verify=False)
        text = res.text or ""
        meta = {"http_status": res.status_code}
        if text.strip().startswith("{"):
            data = json.loads(text)
            ds = data.get("dataSearch") or data
            content = ds.get("content") if isinstance(ds, dict) else data.get("content")
            if isinstance(content, list) and content:
                return content[0], meta
            if isinstance(content, dict):
                return content, meta
            return data, meta
        root = ET.fromstring(text)
        node = root.find(".//content")
        if node is not None:
            return {c.tag: (c.text or "") for c in list(node)}, meta
        return {}, meta
    except Exception as e:
        return {}, {"error": str(e)}


def _parse_careernet_list(text: str):
    """OpenAPI 목록/단건 파싱 → (items:list, meta)"""
    import xml.etree.ElementTree as ET
    items, meta = [], {}
    t = (text or "").strip()
    if not t:
        return [], {"error": "empty"}
    if t.startswith("{") or t.startswith("["):
        data = json.loads(t)
        if isinstance(data, list):
            return data, meta
        ds = data.get("dataSearch") or data
        content = ds.get("content") if isinstance(ds, dict) else None
        if isinstance(content, list):
            items = content
        elif isinstance(content, dict):
            items = [content]
        elif isinstance(data.get("content"), list):
            items = data["content"]
        return items, meta
    root = ET.fromstring(t)
    for node in root.findall(".//content"):
        row = {c.tag: (c.text or "") for c in list(node)}
        if row:
            items.append(row)
    return items, meta



# 학과/학교 API 필수 gubun 후보
_CN_MAJOR_GUBUN = ["univ_list", "대학교", "전문대학", "high_list", "고등학교"]
_CN_SCHOOL_GUBUN = ["high_list", "univ_list", "midd_list", "elem_list", "고등학교", "대학교", "중학교", "초등학교"]
_CN_SUBJECT_CODES = [
    "", "100391", "100392", "100393", "100394", "100395", "100396", "100397",
]


def fetch_careernet_majors(
    api_key: str,
    keyword: str = "",
    gubun: str = "대학교",
    subject: str = "",
    page: int = 1,
    per_page: int = 50,
):
    """학과정보 목록 MAJOR — gubun 필수, searchTitle 검색"""
    import requests
    key = (api_key or "").strip()
    if not key:
        return [], {"error": "CAREER_NET_API_KEY 없음"}
    params = {
        "apiKey": key,
        "svcType": "api",
        "svcCode": "MAJOR",
        "contentType": "json",
        "gubun": gubun or "대학교",
        "thisPage": str(page),
        "perPage": str(per_page),
    }
    if keyword:
        params["searchTitle"] = keyword
    if subject:
        params["subject"] = subject
    try:
        res = requests.get(
            "https://www.career.go.kr/cnet/openapi/getOpenApi",
            params=params, timeout=30, verify=False,
        )
        items, meta = _parse_careernet_list(res.text or "")
        meta["http_status"] = res.status_code
        meta["gubun"] = gubun
        norm = []
        for it in items or []:
            if not isinstance(it, dict):
                continue
            d = dict(it)
            name = (
                d.get("facilName") or d.get("mClass") or d.get("major")
                or d.get("majorNm") or d.get("name") or ""
            )
            if name:
                d["major"] = str(name).strip()
            d["lclsfNm"] = d.get("lClass") or d.get("lclsfNm") or ""
            d["mclsfNm"] = d.get("mClass") or d.get("mclsfNm") or ""
            norm.append(d)
        return norm, meta
    except Exception as e:
        return [], {"error": str(e)}


def fetch_careernet_major_view(api_key: str, major_seq: str = "", major_nm: str = "", gubun: str = "대학교"):
    """학과 상세 MAJOR_VIEW"""
    import requests
    key = (api_key or "").strip()
    if not key:
        return {}, {"error": "CAREER_NET_API_KEY 없음"}
    params = {
        "apiKey": key,
        "svcType": "api",
        "svcCode": "MAJOR_VIEW",
        "contentType": "json",
        "gubun": gubun or "대학교",
    }
    if major_seq:
        params["majorSeq"] = major_seq
    if major_nm:
        params["searchTitle"] = major_nm
    try:
        res = requests.get(
            "https://www.career.go.kr/cnet/openapi/getOpenApi",
            params=params, timeout=30, verify=False,
        )
        items, meta = _parse_careernet_list(res.text or "")
        meta["http_status"] = res.status_code
        if items:
            d = dict(items[0])
            if not d.get("major"):
                d["major"] = d.get("facilName") or d.get("mClass") or d.get("name") or ""
            return d, meta
        return {}, meta
    except Exception as e:
        return {}, {"error": str(e)}


def fetch_careernet_schools(api_key: str, keyword: str = "", region: str = "", gubun: str = "고등학교"):
    """학교정보 목록 SCHOOL — gubun 필수"""
    import requests
    key = (api_key or "").strip()
    if not key:
        return [], {"error": "CAREER_NET_API_KEY 없음"}
    params = {
        "apiKey": key,
        "svcType": "api",
        "svcCode": "SCHOOL",
        "contentType": "json",
        "gubun": gubun or "고등학교",
        "thisPage": "1",
        "perPage": "50",
    }
    if keyword:
        params["searchSchulNm"] = keyword
    if region:
        region_codes = {
            "서울": "100260", "부산": "100267", "대구": "100272", "인천": "100269",
            "광주": "112691", "대전": "100271", "울산": "100273", "세종": "100704",
            "경기": "100276", "강원": "100278", "충북": "100280", "충남": "100281",
            "전북": "100282", "전남": "100285", "경북": "100285", "경남": "100291",
        }
        code = None
        for k, v in region_codes.items():
            if k in (region or ""):
                code = v
                break
        params["region"] = code or region
    try:
        res = requests.get(
            "https://www.career.go.kr/cnet/openapi/getOpenApi",
            params=params, timeout=30, verify=False,
        )
        items, meta = _parse_careernet_list(res.text or "")
        meta["http_status"] = res.status_code
        norm = []
        for it in items or []:
            if not isinstance(it, dict):
                continue
            d = dict(it)
            d["schoolNm"] = (
                d.get("schoolName") or d.get("schoolNm") or d.get("schulNm")
                or d.get("name") or ""
            )
            norm.append(d)
        return norm, meta
    except Exception as e:
        return [], {"error": str(e)}


def fetch_careernet_cose(api_key: str, keyword: str = "", page: int = 1):
    """진로교육자료 목록 COSE — searchTitleWord, dataTitle, attFile"""
    import requests
    key = (api_key or "").strip()
    if not key:
        return [], {"error": "CAREER_NET_API_KEY 없음"}
    params = {
        "apiKey": key,
        "svcType": "api",
        "svcCode": "COSE",
        "contentType": "json",
        "thisPage": str(page),
        "perPage": "50",
    }
    if keyword:
        params["searchTitleWord"] = keyword
    try:
        res = requests.get(
            "https://www.career.go.kr/cnet/openapi/getOpenApi",
            params=params, timeout=30, verify=False,
        )
        items, meta = _parse_careernet_list(res.text or "")
        meta["http_status"] = res.status_code
        norm = []
        for it in items or []:
            if not isinstance(it, dict):
                continue
            d = dict(it)
            title = (
                d.get("dataTitle") or d.get("title") or d.get("subject")
                or d.get("name") or d.get("boardTitle") or ""
            )
            d["title"] = str(title).strip() or str(d.get("dataTitle") or "").strip()
            seq = str(d.get("seq") or d.get("boardSeq") or "").strip()
            d["seq"] = seq
            att = str(d.get("attFile") or d.get("fileUrl") or d.get("url") or "").strip()
            d["attFile"] = att
            # 썸네일(no=0)보다 실제 첨부(no>=1) 우선. content seq로 URL 생성 금지
            files = parse_cose_att_files(att, seq)
            d["url"] = files[0] if files else ""
            d["files"] = files
            d["author"] = d.get("author") or d.get("selAuthor") or ""
            d["year"] = d.get("year") or ""
            norm.append(d)
        return norm, meta
    except Exception as e:
        return [], {"error": str(e)}


def fetch_careernet_cose_view(api_key: str, seq: str):
    """진로교육자료 상세 COSE_VIEW"""
    import requests
    key = (api_key or "").strip()
    if not key or not seq:
        return {}, {"error": "키 또는 seq 없음"}
    params = {
        "apiKey": key,
        "svcType": "api",
        "svcCode": "COSE_VIEW",
        "contentType": "json",
        "seq": str(seq),
    }
    try:
        res = requests.get(
            "https://www.career.go.kr/cnet/openapi/getOpenApi",
            params=params, timeout=30, verify=False,
        )
        items, meta = _parse_careernet_list(res.text or "")
        meta["http_status"] = res.status_code
        if items:
            d = dict(items[0])
            d["title"] = d.get("dataTitle") or d.get("title") or ""
            att = str(d.get("attFile") or "").strip()
            d["attFile"] = att
            files = parse_cose_att_files(att, str(seq))
            d["url"] = files[0] if files else ""
            d["files"] = files
            d["content"] = d.get("dataContent") or d.get("content") or ""
            return d, meta
        return {}, meta
    except Exception as e:
        return {}, {"error": str(e)}



def parse_cose_att_files(att_file: str, content_seq: str = ""):
    """
    attFile은 콤마 구분 다중 URL.
    - no=0 만 있으면 표지/썸네일일 수 있어 같은 file seq의 no=1,2 후보를 추가
    - 목록에 no>=1이 있으면 그것을 앞에 둠
    - 자료 content seq로 다운로드 URL을 만들지 않음 (파일 seq와 다름)
    """
    import re
    urls = []
    seen = set()
    for part in str(att_file or "").split(","):
        u = part.strip()
        if u.startswith("http") and u not in seen:
            seen.add(u)
            urls.append(u)

    # no=0만 있는 경우 → no=1, no=2 후보 추가 (본문 PDF일 가능성)
    expanded = list(urls)
    for u in urls:
        m = re.search(r"(fileDownload2\.do\?seq=\d+)&no=0\b", u)
        if not m:
            m = re.search(r"(https?://[^\s]+fileDownload2\.do\?seq=\d+)&no=0\b", u)
        if m or ("no=0" in u and "fileDownload2.do" in u and "seq=" in u):
            base = re.sub(r"&no=\d+", "", u)
            if "fileDownload2.do" in base and "seq=" in base:
                for n in (1, 2, 3):
                    cand = f"{base}&no={n}"
                    if cand not in seen:
                        seen.add(cand)
                        expanded.append(cand)

    preferred, thumbs, others = [], [], []
    for u in expanded:
        if re.search(r"no=0\b", u):
            thumbs.append(u)
        elif re.search(r"no=[1-9]\d*\b", u):
            preferred.append(u)
        else:
            others.append(u)
    # 본문 후보(no>=1) → 기타 → 썸네일(no=0)
    return preferred + others + thumbs


def cose_download_url(item: dict) -> str:
    """실제 자료 파일 URL (본문 후보 우선)"""
    if not isinstance(item, dict):
        return ""
    att = item.get("attFile") or item.get("url") or item.get("fileUrl") or ""
    files = parse_cose_att_files(att, str(item.get("seq") or ""))
    return files[0] if files else ""


def cose_detail_page_url(seq: str) -> str:
    """진로교육자료 웹 상세/목록 검색"""
    s = str(seq or "").strip()
    if not s:
        return ""
    # 자료 게시판 상세 (seq)
    return f"https://www.career.go.kr/cnet/front/web/retriveEducDataView.do?seq={s}"


def cose_search_page_url(title: str) -> str:
    """제목으로 자료실 검색"""
    import urllib.parse as _up
    q = _up.quote(str(title or "").strip())
    if not q:
        return "https://www.career.go.kr/cloud/w/edudata/list"
    return f"https://www.career.go.kr/cloud/w/edudata/list?searchWrd={q}"


def _cn_item_key(item: dict, name_keys):
    for k in name_keys:
        v = item.get(k)
        if v and str(v).strip():
            return str(v).strip()
    return json.dumps(item, ensure_ascii=False, sort_keys=True)[:120]


def _merge_cn_items(bucket: dict, items: list, name_keys):
    for it in items or []:
        if not isinstance(it, dict):
            continue
        k = _cn_item_key(it, name_keys)
        if k and k not in bucket:
            bucket[k] = it


# 캐시 시드 키워드 (빈 검색 + 분야 훑기)
_CN_SEED_KEYWORDS = [
    "", "가", "나", "다", "라", "마", "바", "사", "아", "자", "차", "카", "타", "파", "하",
    "요리", "조리", "제과", "제빵", "식품", "영양", "간호", "의료", "약학", "보건",
    "기계", "전기", "전자", "컴퓨터", "소프트웨어", "정보", "디자인", "미술", "음악",
    "체육", "교육", "유아", "사회복지", "경영", "경제", "관광", "호텔", "항공",
    "건축", "토목", "환경", "농업", "생명", "화학", "물리", "수학", "영어", "일본어",
    "미디어", "방송", "영상", "광고", "패션", "뷰티", "자동차", "로봇", "AI", "데이터",
]


def build_careernet_cache(api_key: str, force: bool = False):
    """
    MAJOR / SCHOOL / COSE 전체 구분 수집 → cache JSON.
    갱신 주기: CN_CACHE_TTL_DAYS(기본 30일)
    """
    from datetime import datetime, timedelta
    CACHE_DIR.mkdir(exist_ok=True)
    summary = {"majors": 0, "schools": 0, "cose": 0, "errors": [], "refreshed": False}

    # 월 단위 TTL: 만료 시 force
    if not force and CN_CACHE_META.exists():
        try:
            meta = json.loads(CN_CACHE_META.read_text(encoding="utf-8"))
            built = datetime.fromisoformat(meta.get("built_at", "2000-01-01"))
            if datetime.now() - built > timedelta(days=CN_CACHE_TTL_DAYS):
                force = True
                summary["errors"].append("cache_expired_force_refresh")
        except Exception:
            force = True
    elif not force:
        # 메타 없고 캐시 파일도 빈약하면 강제
        force = True

    def _load_existing(path):
        if path.exists() and not force:
            try:
                return json.loads(path.read_text(encoding="utf-8"))
            except Exception:
                return []
        return []

    majors_map, schools_map, cose_map = {}, {}, {}
    for it in _load_existing(CN_MAJOR_CACHE):
        if isinstance(it, dict):
            _merge_cn_items(majors_map, [it], ["major", "majorNm", "mjrNm", "name", "facilName"])
    for it in _load_existing(CN_SCHOOL_CACHE):
        if isinstance(it, dict):
            _merge_cn_items(schools_map, [it], ["schoolNm", "schulNm", "name", "schoolName"])
    for it in _load_existing(CN_COSE_CACHE):
        if isinstance(it, dict):
            _merge_cn_items(cose_map, [it], ["title", "subject", "name", "boardSeq", "seq"])

    need_major = force or len(majors_map) < 30
    need_school = force or len(schools_map) < 20
    need_cose = force or len(cose_map) < 20

    if need_major:
        for gubun in _CN_MAJOR_GUBUN:
            for subj in _CN_SUBJECT_CODES:
                for page in range(1, 6):
                    try:
                        items, meta = fetch_careernet_majors(
                            api_key, keyword="", gubun=gubun, subject=subj, page=page, per_page=50
                        )
                        _merge_cn_items(
                            majors_map, items,
                            ["major", "facilName", "mClass", "mclsfNm", "majorNm", "name"],
                        )
                        if not items:
                            break
                    except Exception as e:
                        summary["errors"].append(f"major:{gubun}:{subj}:{e}")
                if len(majors_map) >= 500:
                    break
            if len(majors_map) >= 500:
                break
        for kw in _CN_SEED_KEYWORDS:
            if not kw:
                continue
            try:
                items, _ = fetch_careernet_majors(api_key, keyword=kw, gubun="대학교")
                _merge_cn_items(
                    majors_map, items,
                    ["major", "facilName", "mClass", "majorNm", "name"],
                )
            except Exception as e:
                summary["errors"].append(f"major_kw:{kw}:{e}")
            if len(majors_map) >= 600:
                break

    if need_school:
        for gubun in _CN_SCHOOL_GUBUN:
            try:
                items, _ = fetch_careernet_schools(api_key, keyword="", region="", gubun=gubun)
                _merge_cn_items(
                    schools_map, items,
                    ["schoolNm", "schoolName", "schulNm", "name"],
                )
            except Exception as e:
                summary["errors"].append(f"school:{gubun}:{e}")
        for kw in ("자동차", "공고", "특성화", "마이스터", "요리", "조리"):
            try:
                items, _ = fetch_careernet_schools(api_key, keyword=kw, gubun="고등학교")
                _merge_cn_items(
                    schools_map, items,
                    ["schoolNm", "schoolName", "schulNm", "name"],
                )
            except Exception as e:
                summary["errors"].append(f"school_kw:{kw}:{e}")

    if need_cose:
        for page in range(1, 8):
            try:
                items, _ = fetch_careernet_cose(api_key, keyword="", page=page)
                _merge_cn_items(
                    cose_map, items,
                    ["title", "subject", "name", "boardSeq", "seq"],
                )
                if not items:
                    break
            except Exception as e:
                summary["errors"].append(f"cose:{page}:{e}")
        for kw in _CN_SEED_KEYWORDS:
            if not kw:
                continue
            try:
                items, _ = fetch_careernet_cose(api_key, keyword=kw)
                _merge_cn_items(
                    cose_map, items,
                    ["title", "subject", "name", "boardSeq", "seq"],
                )
            except Exception as e:
                summary["errors"].append(f"cose_kw:{kw}:{e}")
            if len(cose_map) >= 400:
                break

    majors = list(majors_map.values())
    schools = list(schools_map.values())
    cose = list(cose_map.values())
    try:
        CN_MAJOR_CACHE.write_text(json.dumps(majors, ensure_ascii=False), encoding="utf-8")
        CN_SCHOOL_CACHE.write_text(json.dumps(schools, ensure_ascii=False), encoding="utf-8")
        CN_COSE_CACHE.write_text(json.dumps(cose, ensure_ascii=False), encoding="utf-8")
        from datetime import datetime
        CN_CACHE_META.write_text(
            json.dumps({
                "built_at": datetime.now().isoformat(timespec="seconds"),
                "majors": len(majors),
                "schools": len(schools),
                "cose": len(cose),
            }, ensure_ascii=False),
            encoding="utf-8",
        )
        summary["refreshed"] = True
    except Exception as e:
        summary["errors"].append(f"save:{e}")
    summary["majors"] = len(majors)
    summary["schools"] = len(schools)
    summary["cose"] = len(cose)
    return summary


def load_cn_cache(kind: str):
    path = {
        "majors": CN_MAJOR_CACHE,
        "schools": CN_SCHOOL_CACHE,
        "cose": CN_COSE_CACHE,
    }.get(kind)
    if not path or not path.exists():
        return []
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except Exception:
        return []


def score_cn_item(item: dict, keywords: list) -> int:
    if not keywords:
        return 0
    blob = " ".join(str(v) for v in item.values() if v is not None).lower()
    score = 0
    for k in keywords:
        kk = str(k).strip().lower()
        if kk and kk in blob:
            score += 2 if len(kk) >= 2 else 1
    return score


def pick_from_cn_cache(kind: str, keywords: list, limit: int = 10):
    items = load_cn_cache(kind)
    if not items:
        return []
    scored = []
    for it in items:
        if not isinstance(it, dict):
            continue
        sc = score_cn_item(it, keywords)
        if sc > 0:
            scored.append((sc, it))
    scored.sort(key=lambda x: -x[0])
    if scored:
        return [x[1] for x in scored[:limit]]
    # 키워드 미매칭 시 상위 일부만 (완전 공백 방지)
    return items[: min(5, limit)]


def ai_pick_from_cache(provider, api_key, kind_label: str, keywords: list, candidates: list, limit: int = 8):
    """캐시 후보 중 관심과 연관된 항목을 AI가 고름"""
    if not api_key or not candidates:
        return candidates[:limit]
    slim = []
    for i, c in enumerate(candidates[:80]):
        if not isinstance(c, dict):
            continue
        name = (
            c.get("dataTitle") or c.get("title") or c.get("subject")
            or c.get("major") or c.get("majorNm") or c.get("mjrNm") or c.get("mClass") or c.get("facilName")
            or c.get("schoolNm") or c.get("schulNm") or c.get("name")
            or f"item{i}"
        )
        extra = c.get("author") or c.get("year") or c.get("activityType") or c.get("achieveType") or ""
        slim.append({"i": i, "name": str(name)[:50], "extra": str(extra)[:30]})
    try:
        edu_hint = ""
        if "교육" in kind_label or "cose" in kind_label.lower() or "자료" in kind_label:
            edu_hint = (
                "진로교육자료는 흥미·적성 이해, 직업탐색, 진로설계, 체험·수업 자료처럼 "
                "관심 분야 역량·적성을 키우는 자료를 우선 선택. 제목만 보고 무관한 것은 제외.\n"
            )
        raw = get_ai_response(
            provider,
            api_key,
            "관련 항목 선택기. JSON만 출력. 관심과 무관한 항목은 넣지 말 것.",
            (
                f"관심키워드: {keywords}\n"
                f"종류: {kind_label}\n"
                f"{edu_hint}"
                f"후보: {json.dumps(slim, ensure_ascii=False)}\n"
                f'관련 높은 순 index만. 예: {{"idx":[0,3,5]}} 최대 {limit}개. '
                "관련 없으면 빈 배열."
            ),
        )
        m = re.search(r"\{[\s\S]*\}", raw or "")
        if m:
            data = json.loads(m.group(0))
            idxs = data.get("idx") or data.get("indexes") or []
            picked = []
            for i in idxs:
                try:
                    ii = int(i)
                    if 0 <= ii < len(candidates):
                        picked.append(candidates[ii])
                except Exception:
                    pass
            if picked:
                return picked[:limit]
            # 교육자료는 무관 시 빈 목록 허용
            if edu_hint:
                return []
    except Exception:
        pass
    return candidates[:limit]



def load_jinsol_cache():
    if not JINSOL_CACHE.exists():
        return []
    try:
        data = json.loads(JINSOL_CACHE.read_text(encoding="utf-8"))
        return data if isinstance(data, list) else []
    except Exception:
        return []


def scrape_jinsol_solutions(max_pages: int = 50, extract_pdf: bool = True, progress_cb=None):
    """
    진로솔루션 목록·상세 수집 (www.career.go.kr cloud API)
    목록: POST /cloud/api/counsel/jinsolSearch
    상세: POST /cloud/api/counsel/jinsolView  body {"seq": "..."}
    """
    import time
    import requests
    from datetime import datetime

    CACHE_DIR.mkdir(exist_ok=True)
    session = requests.Session()
    session.headers.update({
        "User-Agent": "Mozilla/5.0 (compatible; KYWA-AI-Tools/1.0)",
        "Accept": "application/json",
        "Content-Type": "application/json",
    })
    base = "https://www.career.go.kr"
    list_url = base + "/cloud/api/counsel/jinsolSearch"
    view_url = base + "/cloud/api/counsel/jinsolView"

    items_meta = []
    page = 0
    total_pages = 1
    while page < total_pages and page < max_pages:
        res = session.post(
            list_url,
            params={"size": 12, "page": page, "sort": "B.REGIST_DT,desc"},
            json={"searchWord": ""},
            timeout=30,
            verify=False,
        )
        res.raise_for_status()
        data = res.json()
        total_pages = int(data.get("totalPages") or 1)
        content = data.get("content") or []
        for row in content:
            if isinstance(row, dict) and row.get("seq"):
                items_meta.append(row)
        if progress_cb:
            progress_cb(f"목록 {page+1}/{total_pages} · 누적 {len(items_meta)}건")
        page += 1
        time.sleep(0.15)

    results = []
    for i, row in enumerate(items_meta):
        seq = str(row.get("seq") or "")
        rec = {
            "seq": seq,
            "job_nm": row.get("job_nm") or "",
            "subject": row.get("subject") or "",
            "year": row.get("year") or "",
            "month": row.get("month") or "",
            "thumbnail": row.get("thumbnail") or "",
            "downloadfile": row.get("downloadfile") or "",
            "rcnt": row.get("rcnt"),
            "url": f"{base}/cloud/w/counsel/jinsolView?seq={seq}",
            "text": "",
        }
        try:
            vr = session.post(view_url, json={"seq": seq}, timeout=30, verify=False)
            if vr.ok:
                det = vr.json() if vr.text else {}
                if isinstance(det, dict):
                    rec["job_nm"] = det.get("job_nm") or rec["job_nm"]
                    rec["subject"] = det.get("subject") or rec["subject"]
                    rec["year"] = det.get("year") or rec["year"]
                    rec["month"] = det.get("month") or rec["month"]
                    rec["downloadfile"] = det.get("downloadfile") or rec["downloadfile"]
                    rec["thumbnail"] = det.get("thumbnail") or rec["thumbnail"]
        except Exception as e:
            rec["view_error"] = str(e)

        # PDF 텍스트 일부 추출 (상세 안내용)
        pdf_path = rec.get("downloadfile") or ""
        if extract_pdf and pdf_path:
            try:
                pdf_url = pdf_path if str(pdf_path).startswith("http") else (base + str(pdf_path))
                pr = session.get(pdf_url, timeout=40, verify=False)
                if pr.ok and pr.content[:4] == b"%PDF":
                    try:
                        import io
                        from pypdf import PdfReader
                        reader = PdfReader(io.BytesIO(pr.content))
                        parts = []
                        for pg in reader.pages[:8]:
                            t = pg.extract_text() or ""
                            if t.strip():
                                parts.append(t.strip())
                        rec["text"] = "\n".join(parts)[:6000]
                    except Exception as e:
                        rec["pdf_error"] = f"extract:{e}"
                else:
                    rec["pdf_error"] = f"http:{getattr(pr,'status_code',0)}"
            except Exception as e:
                rec["pdf_error"] = str(e)

        results.append(rec)
        if progress_cb and (i % 5 == 0 or i == len(items_meta) - 1):
            progress_cb(f"상세 {i+1}/{len(items_meta)} · {rec.get('job_nm') or seq}")
        time.sleep(0.12)

    JINSOL_CACHE.write_text(json.dumps(results, ensure_ascii=False), encoding="utf-8")
    from datetime import datetime as _dt
    JINSOL_META.write_text(
        json.dumps({
            "built_at": _dt.now().isoformat(timespec="seconds"),
            "count": len(results),
            "with_text": sum(1 for r in results if r.get("text")),
        }, ensure_ascii=False),
        encoding="utf-8",
    )
    return {"count": len(results), "with_text": sum(1 for r in results if r.get("text"))}



def normalize_major_rows(majors: list) -> list:
    """학과 레코드를 대/중/소 분류 행으로 정규화 (긴 콤마 나열 분리)"""
    rows = []
    seen = set()
    for m in majors or []:
        if not isinstance(m, dict):
            continue
        large = str(m.get("lClass") or m.get("lclsfNm") or m.get("large") or "").strip()
        mid = str(m.get("mClass") or m.get("mclsfNm") or m.get("middle") or "").strip()
        facil = str(m.get("facilName") or "").strip()
        main = str(
            m.get("major") or m.get("majorNm") or m.get("mjrNm") or m.get("name") or mid or ""
        ).strip()
        # facilName이 콤마/·로 여러 세부학과면 분리
        parts = []
        if facil and ("," in facil or "·" in facil or "、" in facil):
            for p in facil.replace("·", ",").replace("、", ",").split(","):
                p = p.strip()
                if p and len(p) >= 2:
                    parts.append(p)
        if not parts:
            parts = [main or facil or mid]
        for p in parts:
            key = (large, mid, p)
            if key in seen or not p:
                continue
            seen.add(key)
            rows.append({
                "대분류": large,
                "중분류": mid,
                "소분류(학과)": p,
                "원본계열": large,
            })
    return rows


def ai_pick_top_majors(provider, api_key, keywords: list, rows: list, limit: int = 5):
    """유망·관련 학과 5개 선별"""
    if not rows:
        return []
    if not api_key:
        return rows[:limit]
    slim = [
        {"i": i, "대": r.get("대분류"), "중": r.get("중분류"), "소": r.get("소분류(학과)")}
        for i, r in enumerate(rows[:80])
    ]
    try:
        prompt = (
            f"관심: {keywords}\n"
            "최근·미래에 유망하거나 관심(야생조류·동물·생물 등)과 직접 관련된 학과만 "
            f"최대 {limit}개 index 선택. 환경공학·해양·IT만 해당되면 제외.\n"
            f"후보: {json.dumps(slim, ensure_ascii=False)}\n"
            'JSON만: {"idx":[0,2,5]}'
        )
        raw = get_ai_response(
            provider,
            api_key,
            "진로 학과 선별기. JSON만. 관심과 무관한 학과는 제외.",
            prompt,
        )
        m = re.search(r"\{[\s\S]*\}", raw or "")
        if m:
            idxs = json.loads(m.group(0)).get("idx") or []
            out = []
            for i in idxs:
                try:
                    ii = int(i)
                    if 0 <= ii < len(rows):
                        out.append(rows[ii])
                except Exception:
                    pass
            if out:
                return out[:limit]
    except Exception:
        pass
    return rows[:limit]


def ai_filter_ggomgil(provider, api_key, keywords: list, hits: list, limit: int = 10):
    """꿈길 체험 중 관심·직업 관련도 높은 것만"""
    if not hits:
        return []
    if not api_key:
        return hits[:limit]
    slim = []
    for i, h in enumerate(hits[:40]):
        slim.append({
            "i": i,
            "name": h.get("체험프로그램명") or "",
            "job": h.get("체험프로그램 직업유형") or "",
            "type": h.get("체험유형") or "",
            "place": h.get("체험처명") or "",
        })
    try:
        prompt = (
            f"관심키워드: {keywords}\n"
            f"후보: {json.dumps(slim, ensure_ascii=False)}\n"
            f'관련 높은 순 {{"idx":[...]}} 최대 {limit}개. 없으면 빈 배열.'
        )
        raw = get_ai_response(
            provider,
            api_key,
            "체험 프로그램 선별기. JSON만. 관심과 무관하면 넣지 말 것.",
            prompt,
        )
        m = re.search(r"\{[\s\S]*\}", raw or "")
        if m:
            idxs = json.loads(m.group(0)).get("idx") or []
            out = []
            for i in idxs:
                try:
                    ii = int(i)
                    if 0 <= ii < len(hits):
                        out.append(hits[ii])
                except Exception:
                    pass
            return out[:limit]
    except Exception:
        pass
    return hits[:limit]


def pick_jinsol_for_keywords(keywords: list, limit: int = 5):
    """캐시된 진로솔루션 중 관심 키워드 매칭 — 무관하면 빈 목록"""
    items = load_jinsol_cache()
    if not items:
        return []
    core = [str(k).strip().lower() for k in (keywords or []) if str(k).strip()]
    kws = expand_interest_keywords(keywords)
    kws = [str(k).strip().lower() for k in kws if str(k).strip()]
    negs = [n.lower() for n in interest_negative_terms(keywords)]
    bird = is_bird_interest(keywords)
    must_any = []
    if bird:
        must_any = ["새", "조류", "탐조", "동물", "야생", "생물", "수의", "동물원", "생태학", "철새", "반려"]
    scored = []
    for it in items:
        job = str(it.get("job_nm") or "")
        sub = str(it.get("subject") or "")
        blob = " ".join([job, sub, str(it.get("text") or "")[:800]]).lower()
        sc = 0
        for k in kws:
            if not k:
                continue
            if k in job.lower():
                sc += 6
            elif k in sub.lower():
                sc += 3
            elif k in blob:
                sc += 1
        for bad in negs:
            if bad and (bad in job.lower() or bad in sub.lower() or bad in blob[:200]):
                sc -= 8
        # 새/동물 관심이면 직업명·제목에 동물·생물 계열이 하나도 없으면 제외
        if must_any and not any(m in job.lower() or m in sub.lower() for m in must_any):
            continue
        if sc >= 5:
            scored.append((sc, it))
    scored.sort(key=lambda x: -x[0])
    return [x[1] for x in scored[:limit]]



# ---------- 지자체 복지서비스 API + 청소년 통합 지원 ----------
def get_welfare_api_key() -> str:
    dummy = {"0000", "0", "test", "YOUR_KEY", "changeme", ""}
    try:
        for k in ("WELFARE_API_KEY", "DATA_GO_KR_KEY", "YOUTH_ACTIV_KEY"):
            v = (st.secrets.get(k, "") or "").strip()
            if v and v not in dummy:
                return v
    except Exception:
        pass
    return ""


def fetch_welfare_list(
    service_key: str,
    ctpv_nm: str = "",
    sgg_nm: str = "",
    life_array: str = "",
    age: str = "",
    search_wrd: str = "",
    intrs_thema: str = "",
    page_no: int = 1,
    num_of_rows: int = 30,
):
    """
    사회보장정보원 지자체복지서비스 목록조회
    End Point: https://apis.data.go.kr/B554287/LocalGovernmentWelfareInformations/LcgvWelfarelist
    데이터포맷: XML (JSON 응답 시에도 파싱)
    """
    import requests
    from urllib.parse import unquote, quote
    import xml.etree.ElementTree as ET

    raw_key = (service_key or "").strip()
    if not raw_key:
        return [], {"error": "복지/공공데이터 API 키 없음"}

    keys_try = [raw_key]
    try:
        dec = unquote(raw_key)
        if dec and dec != raw_key:
            keys_try.append(dec)
    except Exception:
        pass
    if "%" not in raw_key:
        try:
            keys_try.append(quote(raw_key, safe=""))
        except Exception:
            pass
    seen_k, keys = set(), []
    for k in keys_try:
        if k and k not in seen_k:
            seen_k.add(k)
            keys.append(k)

    url = "https://apis.data.go.kr/B554287/LocalGovernmentWelfareInformations/LcgvWelfarelist"
    base_params = {
        "pageNo": str(page_no),
        "numOfRows": str(num_of_rows),
    }
    if ctpv_nm:
        base_params["ctpvNm"] = ctpv_nm
    if sgg_nm:
        base_params["sggNm"] = sgg_nm
    if life_array:
        base_params["lifeArray"] = life_array
    if age:
        base_params["age"] = str(age)
    if search_wrd:
        base_params["searchWrd"] = search_wrd
    if intrs_thema:
        base_params["intrsThemaArray"] = intrs_thema

    def _parse_xml(text_body: str):
        items, meta = [], {}
        root = ET.fromstring(text_body)
        for tag in ("resultCode", "resultMessage", "totalCount", "numOfRows", "pageNo"):
            el = root.find(f".//{tag}")
            if el is not None and el.text:
                meta[tag] = el.text
        # servList 노드들
        nodes = root.findall(".//servList")
        if not nodes:
            nodes = root.findall(".//item")
        for node in nodes:
            row = {c.tag: (c.text or "") for c in list(node)}
            if row.get("servNm") or row.get("servId"):
                items.append(row)
        return items, meta

    def _parse_json(text_body: str):
        items, meta = [], {}
        data = json.loads(text_body)
        body = data.get("response") or data
        header = body.get("header") if isinstance(body, dict) else {}
        if isinstance(header, dict):
            meta["resultCode"] = header.get("resultCode")
            meta["resultMessage"] = header.get("resultMsg") or header.get("resultMessage")
        if isinstance(body, dict) and "body" in body:
            body = body["body"]
        if isinstance(body, dict):
            meta["totalCount"] = body.get("totalCount")
            serv = body.get("servList") or body.get("item") or body.get("items")
        else:
            serv = data.get("servList")
        if isinstance(serv, dict):
            if serv.get("servNm") or serv.get("servId"):
                items = [serv]
            else:
                inner = serv.get("item") or serv.get("servList")
                if isinstance(inner, list):
                    items = inner
                elif isinstance(inner, dict):
                    items = [inner]
        elif isinstance(serv, list):
            items = serv
        return items, meta

    last_meta = {"url": url}
    for key in keys:
        params = dict(base_params)
        params["serviceKey"] = key
        try:
            # requests가 키를 재인코딩하지 않도록 일부 환경은 수동 쿼리 필요
            res = requests.get(url, params=params, timeout=40, verify=False)
            text_body = res.text or ""
            last_meta["http_status"] = res.status_code
            items = []
            meta = dict(last_meta)
            # XML 우선
            try:
                items, pmeta = _parse_xml(text_body)
                meta.update(pmeta)
            except Exception as e:
                meta["xml_err"] = str(e)
                try:
                    items, pmeta = _parse_json(text_body)
                    meta.update(pmeta)
                except Exception as e2:
                    meta["json_err"] = str(e2)
                    meta["preview"] = text_body[:300]
            if items:
                meta["hit"] = len(items)
                meta["key_mode"] = "encoded" if "%" in key else "plain"
                return items, meta
            last_meta = meta
            last_meta["preview"] = text_body[:250]
        except Exception as e:
            last_meta = {"url": url, "error": str(e)}
            continue
    return [], last_meta or {"error": "복지 목록 조회 실패"}


def fetch_welfare_detail(service_key: str, serv_id: str):
    """
    지자체복지서비스 상세조회
    /LcgvWelfaredetailed — 데이터포맷 XML
    """
    import requests
    from urllib.parse import unquote, quote
    import xml.etree.ElementTree as ET

    raw_key = (service_key or "").strip()
    if not raw_key or not serv_id:
        return {}, {"error": "key 또는 servId 없음"}
    keys_try = [raw_key]
    try:
        dec = unquote(raw_key)
        if dec != raw_key:
            keys_try.append(dec)
    except Exception:
        pass
    if "%" not in raw_key:
        try:
            keys_try.append(quote(raw_key, safe=""))
        except Exception:
            pass
    seen_k, keys = set(), []
    for k in keys_try:
        if k and k not in seen_k:
            seen_k.add(k)
            keys.append(k)

    url = "https://apis.data.go.kr/B554287/LocalGovernmentWelfareInformations/LcgvWelfaredetailed"
    last_meta = {"url": url}
    for key in keys:
        params = {"serviceKey": key, "servId": serv_id}
        try:
            res = requests.get(url, params=params, timeout=40, verify=False)
            text_body = res.text or ""
            last_meta["http_status"] = res.status_code
            # XML
            try:
                root = ET.fromstring(text_body)
                row = {}
                for el in root.iter():
                    if len(list(el)) == 0 and el.tag:
                        if el.tag not in row or not row[el.tag]:
                            row[el.tag] = el.text or ""
                if row.get("servNm") or row.get("servId") or row.get("sprtTrgtCn"):
                    return row, last_meta
            except Exception as e:
                last_meta["xml_err"] = str(e)
            try:
                data = json.loads(text_body)
                body = data.get("response") or data
                if isinstance(body, dict) and "body" in body:
                    body = body["body"]
                if isinstance(body, dict) and (body.get("servNm") or body.get("servId")):
                    return body, last_meta
            except Exception as e:
                last_meta["json_err"] = str(e)
            last_meta["preview"] = text_body[:250]
        except Exception as e:
            last_meta = {"url": url, "error": str(e)}
    return {}, last_meta



def map_school_to_life_age(school: str, age: str = ""):
    """교급 → 복지 lifeArray 코드(코드표) · 나이
    002 아동 / 003 청소년
    """
    s = str(school or "")
    age_n = (age or "").strip()
    if any(x in s for x in ("초", "초등")):
        life = "002"  # 아동
        if not age_n:
            age_n = "11"
    elif any(x in s for x in ("중", "중학", "고", "고등")):
        life = "003"  # 청소년
        if not age_n:
            age_n = "14" if "중" in s else "17"
    else:
        life = "003"
        if not age_n:
            age_n = "15"
    return life, age_n


def map_interest_to_thema_codes(interest: str) -> list:
    """관심 키워드 → intrsThemaArray 코드 (복수 가능)"""
    blob = str(interest or "")
    codes = []
    rules = [
        (("건강", "운동", "체육", "의료"), "010"),
        (("심리", "상담", "정신"), "020"),
        (("생활", "지원", "바우처"), "030"),
        (("주거", "주택"), "040"),
        (("일자리", "취업", "직업", "알바"), "050"),
        (("문화", "여가", "예술", "음악", "공연", "요리", "체험"), "060"),
        (("안전", "위기"), "070"),
        (("보육", "돌봄"), "090"),
        (("교육", "학습", "진로", "학교", "코딩", "로봇"), "100"),
        (("보호", "돌봄"), "120"),
    ]
    for keys, code in rules:
        if any(k in blob for k in keys):
            codes.append(code)
    if not codes:
        codes = ["100", "060"]  # 교육·문화여가 기본
    # 중복 제거
    out = []
    for c in codes:
        if c not in out:
            out.append(c)
    return out[:3]


def split_region_ctpv_sgg(region: str):
    """'경기도 고양시' → (경기도, 고양시)"""
    r = str(region or "").strip()
    if not r:
        return "", ""
    parts = r.replace(",", " ").split()
    ctpv, sgg = "", ""
    for p in parts:
        if any(x in p for x in ("특별시", "광역시", "특별자치시", "특별자치도", "도")) or p.endswith("도"):
            ctpv = p
        elif any(x in p for x in ("시", "군", "구")):
            sgg = p
    if not ctpv and parts:
        # 고양시만 있는 경우 등
        if "서울" in r:
            ctpv = "서울특별시"
        elif "경기" in r:
            ctpv = "경기도"
        elif "인천" in r:
            ctpv = "인천광역시"
        elif "부산" in r:
            ctpv = "부산광역시"
        elif "대구" in r:
            ctpv = "대구광역시"
        elif "광주" in r:
            ctpv = "광주광역시"
        elif "대전" in r:
            ctpv = "대전광역시"
        elif "울산" in r:
            ctpv = "울산광역시"
        elif "세종" in r:
            ctpv = "세종특별자치시"
        elif "강원" in r:
            ctpv = "강원특별자치도"
        elif "충북" in r or "충청북" in r:
            ctpv = "충청북도"
        elif "충남" in r or "충청남" in r:
            ctpv = "충청남도"
        elif "전북" in r or "전라북" in r:
            ctpv = "전북특별자치도"
        elif "전남" in r or "전라남" in r:
            ctpv = "전라남도"
        elif "경북" in r or "경상북" in r:
            ctpv = "경상북도"
        elif "경남" in r or "경상남" in r:
            ctpv = "경상남도"
        elif "제주" in r:
            ctpv = "제주특별자치도"
    if not sgg and parts:
        for p in parts:
            if p != ctpv and (p.endswith("시") or p.endswith("군") or p.endswith("구")):
                sgg = p
                break
    return ctpv, sgg


def ai_welfare_eligibility_hint(provider, api_key, profile: dict, items: list, limit: int = 8):
    """복지 목록을 프로필에 맞춰 요약·우선순위 (확정 판정 금지)"""
    if not items or not api_key:
        return items[:limit], ""
    slim = []
    for i, it in enumerate(items[:25]):
        if not isinstance(it, dict):
            continue
        slim.append({
            "i": i,
            "name": it.get("servNm") or "",
            "summary": (it.get("servDgst") or "")[:80],
            "life": it.get("lifeNmArray") or "",
            "target": it.get("trgterIndvdlNmArray") or "",
            "theme": it.get("intrsThemaNmArray") or "",
            "sgg": it.get("sggNm") or "",
        })
    try:
        prompt = (
            f"프로필: {json.dumps(profile, ensure_ascii=False)}\n"
            f"복지목록: {json.dumps(slim, ensure_ascii=False)}\n"
            "청소년·보호자가 확인해볼 만한 순으로 고르고, 수급 확정은 하지 말 것.\n"
            f'JSON: {{"idx":[...], "note":"한줄 안내"}} 최대 {limit}개.'
        )
        raw = get_ai_response(
            provider, api_key,
            "복지 안내 보조. 확정 판정 금지. JSON만.",
            prompt,
        )
        m = re.search(r"\{[\s\S]*\}", raw or "")
        if m:
            obj = json.loads(m.group(0))
            idxs = obj.get("idx") or []
            out = []
            for i in idxs:
                try:
                    ii = int(i)
                    if 0 <= ii < len(items):
                        out.append(items[ii])
                except Exception:
                    pass
            return (out[:limit] if out else items[:limit]), str(obj.get("note") or "")
    except Exception:
        pass
    return items[:limit], ""



# ---------- 주소 → 위경도 (지오코딩) ----------
_GEOCODE_CACHE = {}



def detect_region_from_ip():
    """
    접속 IP 기반 시·도 추정.
    반환 예: {"region":"경기도 성남시", "ctpv":"경기도", "sgg":"성남시", "source":"ip"}
    """
    import requests
    out = {"region": "", "ctpv": "", "sgg": "", "source": ""}
    try:
        res = requests.get("https://ipapi.co/json/", timeout=8, verify=False)
        data = res.json() or {}
        if data.get("error"):
            raise RuntimeError(data.get("reason") or "ipapi error")
        city = (data.get("city") or "").strip()
        region = (data.get("region") or "").strip()  # often province
        country = (data.get("country_code") or "")
        if country and country != "KR":
            out["source"] = "ip_non_kr"
            return out
        # region이 영문일 수 있음 → 간단 매핑
        eng_map = {
            "Seoul": "서울특별시", "Busan": "부산광역시", "Daegu": "대구광역시",
            "Incheon": "인천광역시", "Gwangju": "광주광역시", "Daejeon": "대전광역시",
            "Ulsan": "울산광역시", "Sejong": "세종특별자치시",
            "Gyeonggi-do": "경기도", "Gyeonggi": "경기도",
            "Gangwon-do": "강원특별자치도", "Gangwon": "강원특별자치도",
            "Chungcheongbuk-do": "충청북도", "Chungcheongnam-do": "충청남도",
            "Jeollabuk-do": "전북특별자치도", "Jeollanam-do": "전라남도",
            "Gyeongsangbuk-do": "경상북도", "Gyeongsangnam-do": "경상남도",
            "Jeju-do": "제주특별자치도", "Jeju": "제주특별자치도",
        }
        ctpv = eng_map.get(region, region)
        # 한글 도시명 보정
        city_map = {
            "Seongnam-si": "성남시", "Seongnam": "성남시", "Suwon": "수원시",
            "Goyang": "고양시", "Yongin": "용인시", "Bucheon": "부천시",
            "Ansan": "안산시", "Anyang": "안양시", "Seoul": "서울특별시",
        }
        sgg = city_map.get(city, city)
        if sgg.endswith("-si"):
            sgg = sgg.replace("-si", "시")
        if ctpv or sgg:
            out["ctpv"] = ctpv
            out["sgg"] = sgg if sgg != ctpv else ""
            out["region"] = f"{ctpv} {out['sgg']}".strip()
            out["source"] = "ip"
        return out
    except Exception:
        try:
            res = requests.get("https://ipinfo.io/json", timeout=8, verify=False)
            data = res.json() or {}
            region = (data.get("region") or "").strip()
            city = (data.get("city") or "").strip()
            out["region"] = f"{region} {city}".strip()
            out["ctpv"] = region
            out["sgg"] = city
            out["source"] = "ipinfo"
            return out
        except Exception:
            return out


def region_sido_token(region: str) -> str:
    """시·도 단위 토큰 (지도 필터용)"""
    ctpv, sgg = split_region_ctpv_sgg(region or "")
    if ctpv:
        return ctpv
    r = str(region or "")
    for x in ("서울", "부산", "대구", "인천", "광주", "대전", "울산", "세종",
              "경기", "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주"):
        if x in r:
            return x
    return r[:2] if r else ""


def geocode_address(address: str):
    """
    주소 문자열 → (lat, lon) 또는 None
    우선순위: 카카오 > VWorld > Nominatim(OSM)
    secrets: KAKAO_REST_KEY / VWORLD_API_KEY
    """
    import requests
    from urllib.parse import quote

    addr = " ".join(str(address or "").split())
    if not addr or len(addr) < 4:
        return None
    if addr in _GEOCODE_CACHE:
        return _GEOCODE_CACHE[addr]

    # 1) Kakao
    try:
        kakao = (st.secrets.get("KAKAO_REST_KEY", "") or "").strip()
    except Exception:
        kakao = ""
    if kakao:
        try:
            res = requests.get(
                "https://dapi.kakao.com/v2/local/search/address.json",
                headers={"Authorization": f"KakaoAK {kakao}"},
                params={"query": addr},
                timeout=12,
                verify=False,
            )
            docs = (res.json() or {}).get("documents") or []
            if docs:
                lat = float(docs[0].get("y"))
                lon = float(docs[0].get("x"))
                _GEOCODE_CACHE[addr] = (lat, lon)
                return lat, lon
            # 키워드 검색 보조
            res2 = requests.get(
                "https://dapi.kakao.com/v2/local/search/keyword.json",
                headers={"Authorization": f"KakaoAK {kakao}"},
                params={"query": addr},
                timeout=12,
                verify=False,
            )
            docs2 = (res2.json() or {}).get("documents") or []
            if docs2:
                lat = float(docs2[0].get("y"))
                lon = float(docs2[0].get("x"))
                _GEOCODE_CACHE[addr] = (lat, lon)
                return lat, lon
        except Exception:
            pass

    # 2) VWorld
    try:
        vkey = (st.secrets.get("VWORLD_API_KEY", "") or st.secrets.get("VWORLD_KEY", "") or "").strip()
    except Exception:
        vkey = ""
    if vkey:
        try:
            res = requests.get(
                "https://api.vworld.kr/req/address",
                params={
                    "service": "address",
                    "request": "getcoord",
                    "version": "2.0",
                    "crs": "epsg:4326",
                    "address": addr,
                    "format": "json",
                    "type": "road",
                    "key": vkey,
                },
                timeout=12,
                verify=False,
            )
            data = res.json() or {}
            point = (
                ((data.get("response") or {}).get("result") or {}).get("point")
                or {}
            )
            if point.get("x") and point.get("y"):
                lon = float(point["x"])
                lat = float(point["y"])
                _GEOCODE_CACHE[addr] = (lat, lon)
                return lat, lon
            # 지번 재시도
            res = requests.get(
                "https://api.vworld.kr/req/address",
                params={
                    "service": "address",
                    "request": "getcoord",
                    "version": "2.0",
                    "crs": "epsg:4326",
                    "address": addr,
                    "format": "json",
                    "type": "parcel",
                    "key": vkey,
                },
                timeout=12,
                verify=False,
            )
            data = res.json() or {}
            point = (
                ((data.get("response") or {}).get("result") or {}).get("point")
                or {}
            )
            if point.get("x") and point.get("y"):
                lon = float(point["x"])
                lat = float(point["y"])
                _GEOCODE_CACHE[addr] = (lat, lon)
                return lat, lon
        except Exception:
            pass

    # 3) Nominatim (키 없음, 사용 시 User-Agent 필요)
    try:
        res = requests.get(
            "https://nominatim.openstreetmap.org/search",
            params={"q": addr, "format": "json", "limit": 1, "countrycodes": "kr"},
            headers={"User-Agent": "KYWA-DataFusion/1.0 (youth support)"},
            timeout=15,
            verify=False,
        )
        arr = res.json() or []
        if arr:
            lat = float(arr[0]["lat"])
            lon = float(arr[0]["lon"])
            _GEOCODE_CACHE[addr] = (lat, lon)
            return lat, lon
    except Exception:
        pass

    _GEOCODE_CACHE[addr] = None
    return None


def resolve_lat_lon(row: dict, name_keys=None, addr_keys=None):
    """
    API 행에서 위경도 확보.
    1) lat/lot/lng 필드
    2) 주소 필드 지오코딩
    3) 시설명+지역 조합 지오코딩
    """
    if not isinstance(row, dict):
        return None, None, ""
    lat = lon = None
    for la, lo in (("lat", "lot"), ("lat", "lng"), ("latitude", "longitude"), ("ypos", "xpos")):
        try:
            if row.get(la) not in (None, "", 0, "0") and row.get(lo) not in (None, "", 0, "0"):
                lat = float(row.get(la))
                lon = float(row.get(lo))
                if lat and lon:
                    return lat, lon, "api"
        except Exception:
            pass

    name_keys = name_keys or ("fcltyNm", "fcltNm", "prgrmNm", "servNm", "name")
    addr_keys = addr_keys or (
        "addr1", "addr2", "adres", "address", "mainActvPlcNm", "rdnmadr", "lnmadr"
    )
    addr_parts = []
    for k in addr_keys:
        v = row.get(k)
        if v and str(v).strip() and str(v) not in addr_parts:
            addr_parts.append(str(v).strip())
    # 시도·시군구 보조
    for k in ("ctpvNm", "sggNm", "sido", "sigungu"):
        v = row.get(k)
        if v and str(v).strip() and str(v) not in addr_parts:
            addr_parts.insert(0, str(v).strip())
    name = ""
    for k in name_keys:
        if row.get(k):
            name = str(row.get(k)).strip()
            break
    queries = []
    if addr_parts:
        queries.append(" ".join(addr_parts))
    if name and addr_parts:
        queries.append(f"{name} {' '.join(addr_parts[:2])}")
    if name:
        queries.append(name)
    for q in queries:
        got = geocode_address(q)
        if got:
            return got[0], got[1], "geocode"
    return None, None, ""


def naver_map_search_link(query: str) -> str:
    from urllib.parse import quote
    q = (query or "").strip()
    if not q:
        return ""
    return f"https://map.naver.com/v5/search/{quote(q)}"


def page_header(title: str, caption: str = ""):
    cap_html = f"<p>{caption}</p>" if caption else ""
    st.markdown(
        f'<div class="page-header"><h2>{title}</h2>{cap_html}</div>',
        unsafe_allow_html=True,
    )

def result_card(title: str, body_markdown: str):
    st.markdown(f'<div class="result-card"><h3>{title}</h3></div>', unsafe_allow_html=True)
    st.markdown(body_markdown)


def cm_h1(icon: str, roman: str = "", title: str = ""):
    """대목차: 아이콘 + 제목 (로마숫자 미사용)"""
    # 호환: cm_h1(icon, title) 또는 cm_h1(icon, roman, title)
    if title == "" and roman and not roman.startswith(("I", "V", "X")) and len(roman) > 2:
        title = roman
        roman = ""
    label = f"{icon} {title}".strip()
    st.markdown(f'<div class="cm-h1">{label}</div>', unsafe_allow_html=True)


def cm_h2(num: str, title: str):
    """중목차: 숫자"""
    st.markdown(
        f'<div class="cm-h2">{num} {title}</div>',
        unsafe_allow_html=True,
    )


def cm_card(title: str, body: str = "", meta: str = "", links: str = ""):
    """항목 카드 박스"""
    import html as _html
    parts = ['<div class="cm-card">']
    if title:
        parts.append(f'<div class="cm-card-title">{_html.escape(str(title))}</div>')
    if meta:
        parts.append(f'<div class="cm-meta">{_html.escape(str(meta))}</div>')
    if body:
        parts.append(f'<div class="cm-body">{_html.escape(str(body)).replace(chr(10), "<br/>")}</div>')
    if links:
        parts.append(f'<div class="cm-meta">{links}</div>')  # 링크는 신뢰 HTML
    parts.append('</div>')
    st.markdown("".join(parts), unsafe_allow_html=True)


def cm_empty(msg: str):
    st.markdown(
        f'<div class="cm-card"><div class="cm-meta">{msg}</div></div>',
        unsafe_allow_html=True,
    )




def show_examples(items):
    """메뉴 사용 예시 (D)"""
    st.markdown("**사용 예시**")
    for t in items:
        st.markdown(f'<div class="example-box">• {t}</div>', unsafe_allow_html=True)


def load_feedback():
    if not FEEDBACK_DB.exists():
        return {"ratings": [], "bad_questions": []}
    try:
        with open(FEEDBACK_DB, "r", encoding="utf-8") as f:
            data = json.load(f)
        data.setdefault("ratings", [])
        data.setdefault("bad_questions", [])
        return data
    except Exception:
        return {"ratings": [], "bad_questions": []}


def save_feedback(data):
    with open(FEEDBACK_DB, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def get_ai_response(provider, api_key, system_prompt, user_prompt):
    if not api_key:
        raise ValueError("API 키가 입력되지 않았습니다.")
    import httpx
    http_client = httpx.Client(verify=False, timeout=90.0)
    if provider == "OpenAI (ChatGPT)":
        from openai import OpenAI
        client = OpenAI(api_key=api_key, http_client=http_client)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            temperature=0.3,
        )
        return response.choices[0].message.content
    elif provider == "Claude (Anthropic)":
        from anthropic import Anthropic
        client = Anthropic(api_key=api_key, http_client=http_client)
        response = client.messages.create(
            model="claude-3-5-sonnet-20241022", max_tokens=1500, system=system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )
        return response.content[0].text
    elif provider == "Groq":
        from groq import Groq
        client = Groq(api_key=api_key, http_client=http_client)
        response = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            temperature=0.3,
        )
        return response.choices[0].message.content
    raise ValueError("지원하지 않는 AI 제공자입니다.")


def transcribe_audio_openai(api_key: str, uploaded_audio) -> str:
    """OpenAI Whisper로 음성 → 텍스트"""
    from openai import OpenAI
    import httpx
    client = OpenAI(api_key=api_key, http_client=httpx.Client(verify=False, timeout=180.0))
    # Whisper는 파일 객체에 name 속성이 필요
    bio = BytesIO(uploaded_audio.getvalue())
    bio.name = uploaded_audio.name
    result = client.audio.transcriptions.create(
        model="whisper-1",
        file=bio,
        language="ko",
    )
    return (result.text or "").strip()


def extract_text_from_file(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    if name.endswith(".pdf"):
        import pdfplumber
        pages_text, tables_text = [], []
        with pdfplumber.open(uploaded_file) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                t = page.extract_text() or ""
                if t.strip():
                    pages_text.append(f"[페이지 {i}]\n{t}")
                try:
                    tables = page.extract_tables() or []
                except Exception:
                    tables = []
                for t_idx, table in enumerate(tables, 1):
                    if not table:
                        continue
                    rows = []
                    for row in table:
                        cells = [(c or "").replace("\n", " ").strip() for c in row]
                        if any(cells):
                            rows.append(" | ".join(cells))
                    if rows:
                        tables_text.append(f"[페이지 {i} / 표 {t_idx}]\n" + "\n".join(rows))
        content = "\n\n".join(pages_text)
        if tables_text:
            content += "\n\n===== 추출된 표 =====\n\n" + "\n\n".join(tables_text)
        return content
    if name.endswith(".docx"):
        from docx import Document
        doc = Document(uploaded_file)
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for ti, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                cells = [(c.text or "").replace("\n", " ").strip() for c in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"[DOCX 표 {ti}]\n" + "\n".join(rows))
        return "\n".join(parts)
    raise ValueError("지원 형식: txt, pdf, docx")


def chunk_text(text: str, chunk_size: int = 450, overlap: int = 80):
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    numbered_parts = re.split(r"(?m)(?=^\s*\d+\s*[\.\)、]\s*)", text)
    numbered_parts = [p.strip() for p in numbered_parts if p.strip()]
    if len(numbered_parts) >= 5:
        chunks = []
        for part in numbered_parts:
            if len(part) <= chunk_size * 1.5:
                chunks.append(part)
            else:
                start = 0
                while start < len(part):
                    end = start + chunk_size
                    piece = part[start:end].strip()
                    if piece:
                        chunks.append(piece)
                    if end >= len(part):
                        break
                    start = max(0, end - overlap)
        return chunks
    split_pattern = r"(?=\n\s*(?:제\s*\d+\s*조|제\s*\d+\s*항|별표|부칙|①|②|③|④|⑤))"
    parts = re.split(split_pattern, text)
    raw_chunks, buf = [], ""
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if not buf:
            buf = part
        elif len(buf) + 1 + len(part) <= chunk_size:
            buf = f"{buf}\n{part}"
        else:
            raw_chunks.append(buf)
            buf = part
        while len(buf) > chunk_size:
            cut = buf[:chunk_size]
            m = re.search(r"[\n\.。;；][^\n\.。;；]*$", cut)
            cut_pos = (m.start() + 1) if (m and m.start() > chunk_size * 0.6) else chunk_size
            raw_chunks.append(buf[:cut_pos].strip())
            next_start = max(cut_pos - overlap, 0)
            buf = buf[next_start:].strip()
    if buf:
        raw_chunks.append(buf)
    chunks = []
    for ch in raw_chunks:
        ch = ch.strip()
        if not ch:
            continue
        if chunks and len(chunks[-1]) < 80:
            merged = (chunks[-1] + "\n" + ch).strip()
            if len(merged) <= chunk_size * 1.3:
                chunks[-1] = merged
            else:
                chunks.append(ch)
        else:
            chunks.append(ch)
    if not chunks:
        start = 0
        while start < len(text):
            end = start + chunk_size
            piece = text[start:end].strip()
            if piece:
                chunks.append(piece)
            if end >= len(text):
                break
            start = max(0, end - overlap)
    return chunks


def extract_reg_date(file_name: str):
    name = Path(file_name).stem
    for p in [
        r"(20\d{2})[.\-_/년\s]*(\d{1,2})[.\-_/월\s]*(\d{1,2})",
        r"(?<!\d)(\d{2})[.\-_/년\s]*(\d{1,2})[.\-_/월\s]*(\d{1,2})",
    ]:
        m = re.search(p, name)
        if not m:
            continue
        y, mo, d = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if y < 100:
            y += 2000
        try:
            return datetime(y, mo, d)
        except ValueError:
            continue
    return None


def normalize_reg_name(file_name: str) -> str:
    stem = Path(file_name).stem
    stem = re.sub(r"\(.*?\)", " ", stem)
    stem = re.sub(r"(20\d{2}|\d{2})[.\-_/년\s]*\d{1,2}[.\-_/월\s]*\d{1,2}", " ", stem)
    stem = re.sub(r"(개정|일부개정|전부개정|신설)", " ", stem)
    stem = re.sub(r"[\s_\-]+", "", stem)
    return stem.strip() or Path(file_name).stem


def canonical_source_name(file_name: str) -> str:
    return f"{normalize_reg_name(file_name)}{Path(file_name).suffix.lower() or '.pdf'}"


def choose_latest_reg_files(paths):
    groups = {}
    for p in paths:
        key = normalize_reg_name(p.name)
        dt = extract_reg_date(p.name) or datetime.min
        prev = groups.get(key)
        if prev is None or dt > prev[0]:
            groups[key] = (dt, p)
    return groups


def get_openai_client(api_key: str):
    from openai import OpenAI
    import httpx
    return OpenAI(api_key=api_key, http_client=httpx.Client(verify=False, timeout=90.0))


def build_or_load_collection(name="kywa_regs"):
    import chromadb
    client = chromadb.PersistentClient(path=str(VECTOR_DIR))
    return client.get_or_create_collection(name=name)


def index_regulation_file(file_name: str, text: str, api_key: str, collection_name="kywa_regs"):
    client = get_openai_client(api_key)
    collection = build_or_load_collection(collection_name)
    chunks = chunk_text(text)
    if not chunks:
        raise ValueError("추출된 텍스트가 없습니다.")
    existing = collection.get(where={"source": file_name})
    if existing and existing.get("ids"):
        collection.delete(ids=existing["ids"])
    ids, docs, metas, embeds = [], [], [], []
    for start in range(0, len(chunks), 20):
        batch = chunks[start:start + 20]
        emb_res = client.embeddings.create(model="text-embedding-3-small", input=batch)
        for j, chunk in enumerate(batch):
            i = start + j
            ids.append(f"{file_name}_{i}")
            docs.append(chunk)
            metas.append({"source": file_name, "chunk": i, "collection": collection_name})
            embeds.append(emb_res.data[j].embedding)
    collection.add(ids=ids, documents=docs, metadatas=metas, embeddings=embeds)
    return len(chunks)


def search_regs(question: str, api_key: str, top_k: int = 10, collection_name="kywa_regs"):
    client = get_openai_client(api_key)
    collection = build_or_load_collection(collection_name)
    q_emb = client.embeddings.create(model="text-embedding-3-small", input=question).data[0].embedding
    result = collection.query(query_embeddings=[q_emb], n_results=top_k)
    docs = result.get("documents", [[]])[0]
    metas = result.get("metadatas", [[]])[0]
    return list(zip(docs, metas))


def extract_keywords(text: str):
    stop = {"및", "또는", "관련", "대한", "위해", "하는", "에서", "으로", "하다", "싶어", "어떻게", "알려", "안내", "해줘"}
    tokens = re.findall(r"[가-힣A-Za-z0-9]{2,}", text)
    return [t for t in tokens if t not in stop][:12]


def hybrid_rerank(hits, question: str):
    kws = extract_keywords(question)
    scored = []
    for idx, (doc, meta) in enumerate(hits):
        score = 100 - idx
        d = doc or ""
        for kw in kws:
            if kw in d:
                score += 3
            if kw in str((meta or {}).get("source", "")):
                score += 2
        scored.append((score, doc, meta))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [(d, m) for _, d, m in scored]


def merge_hits(hit_lists, max_total=14, max_per_source=4):
    seen, merged, per_source = set(), [], {}
    for hits in hit_lists:
        for doc, meta in hits:
            source = (meta or {}).get("source", "")
            chunk = (meta or {}).get("chunk", "")
            key = (source, chunk, (doc or "")[:120])
            if key in seen:
                continue
            if per_source.get(source, 0) >= max_per_source:
                continue
            seen.add(key)
            per_source[source] = per_source.get(source, 0) + 1
            merged.append((doc, meta))
            if len(merged) >= max_total:
                return merged
    return merged


ERP_KEYWORDS = [
    "ERP", "erp", "더존", "회계전표", "전표", "구매요청", "발주", "입고", "재고",
    "품목", "거래처", "전자세금계산서", "결산", "계정과목", "예산통제", "품의연동",
    "급여전산", "인사ERP", "전산입력", "ERP매뉴얼", "ERP 매뉴얼",
]

def is_erp_question(text: str) -> bool:
    if not text:
        return False
    t = text.lower()
    return any(k.lower() in t for k in ERP_KEYWORDS)


def search_auto_route(question: str, rewritten: str, api_key: str):
    use_erp = is_erp_question(question) or is_erp_question(rewritten or "")
    col = "kywa_erp" if use_erp else "kywa_regs"
    hits1 = search_regs(question, api_key, top_k=10, collection_name=col)
    hits2 = search_regs(rewritten, api_key, top_k=10, collection_name=col) if rewritten and rewritten != question else []
    merged = merge_hits([hits1, hits2], max_total=14, max_per_source=4)
    merged = hybrid_rerank(merged, f"{question} {rewritten}")
    return merged, col

def prefer_hits_by_question(hits, question: str):
    """질문 핵심어가 출처/본문에 많은 조각을 앞으로"""
    keys = re.findall(r"[가-힣A-Za-z0-9]{2,}", question or "")
    stop = {"관련", "대한", "위해", "하는", "에서", "으로", "어떻게", "알려", "안내", "해줘", "싶어", "방법"}
    keys = [k for k in keys if k not in stop][:12]
    ranked = []
    for doc, meta in hits:
        text = f"{(meta or {}).get('source', '')} {doc or ''}"
        score = sum(1 for k in keys if k in text)
        ranked.append((score, doc, meta))
    ranked.sort(key=lambda x: x[0], reverse=True)
    return [(d, m) for _, d, m in ranked]


def _parse_data_go_items(data):
    """공공데이터포털 JSON 응답에서 item 리스트 추출"""
    items = None
    if isinstance(data, dict):
        body = data.get("response", {}).get("body") or data.get("body") or data
        if isinstance(body, dict):
            items = body.get("items", {})
            if isinstance(items, dict):
                items = items.get("item")
            if items is None:
                items = body.get("item")
        if items is None:
            for k in ("data", "result", "list"):
                if k in data:
                    items = data[k]
                    break
    if items is None:
        return []
    if isinstance(items, dict):
        items = [items]
    if not isinstance(items, list):
        return []
    return items


def fetch_training_facility_safety(
    service_key: str,
    ctpv_nm: str = "",
    fclty_nm: str = "",
    evl_yr: str = "",
    page_no: int = 1,
    num_of_rows: int = 50,
):
    """svc001 수련시설 정보 및 안전점검 종합평가 결과 (단건 페이지)"""
    import requests

    url = "https://apis.data.go.kr/B552713/svc001/getTrFcltyInfoSftyInspcCmphsEvlRsltInfo"
    params = {
        "serviceKey": service_key,
        "pageNo": page_no,
        "numOfRows": num_of_rows,
        "returnType": "JSON",
    }
    if (ctpv_nm or "").strip():
        params["ctpvNm"] = ctpv_nm.strip()
    if (fclty_nm or "").strip():
        params["fcltyNm"] = fclty_nm.strip()
    if (evl_yr or "").strip():
        params["evlYr"] = evl_yr.strip()

    res = requests.get(url, params=params, timeout=30, verify=False)
    res.raise_for_status()
    data = res.json()
    return _parse_data_go_items(data), data


def _data_go_total_count(data) -> int:
    try:
        body = (data.get("response") or data).get("body") or data.get("body") or {}
        tc = body.get("totalCount") or body.get("totalcount")
        return int(tc) if tc is not None else 0
    except Exception:
        return 0


def fetch_training_facility_safety_all(
    service_key: str,
    ctpv_nm: str = "",
    fclty_nm: str = "",
    evl_yr: str = "",
    page_size: int = 100,
):
    """평가연도(+시도) 해당 안전평가 시설 전체 조회"""
    all_items = []
    page = 1
    total = None
    last_raw = None
    while True:
        items, raw = fetch_training_facility_safety(
            service_key,
            ctpv_nm=ctpv_nm,
            fclty_nm=fclty_nm,
            evl_yr=evl_yr,
            page_no=page,
            num_of_rows=page_size,
        )
        last_raw = raw
        if total is None:
            total = _data_go_total_count(raw)
        if not items:
            break
        all_items.extend(items)
        if total and len(all_items) >= total:
            break
        if len(items) < page_size:
            break
        page += 1
        if page > 200:  # 안전장치
            break
    return all_items, {"pages": page, "totalCount": total or len(all_items), "last": last_raw}


def _parse_svc004_response(text: str, ctype: str = ""):
    """svc004 응답(JSON/XML) → (items, meta)"""
    import xml.etree.ElementTree as ET
    from urllib.parse import unquote

    text = (text or "").strip()
    # BOM 제거
    if text.startswith("\ufeff"):
        text = text.lstrip("\ufeff")

    meta = {"content_type": ctype, "snippet": text[:500]}

    # JSON
    if "json" in (ctype or "").lower() or text.startswith("{"):
        import json
        data = json.loads(text)
        items = _parse_data_go_items(data)
        # resultCode 추출
        try:
            hdr = (data.get("response") or data).get("header") or {}
            meta["resultCode"] = hdr.get("resultCode") or hdr.get("resultcode")
            meta["resultMsg"] = hdr.get("resultMsg") or hdr.get("resultmsg")
        except Exception:
            pass
        meta["count"] = len(items)
        return items, {**meta, "data": data}

    # XML
    root = ET.fromstring(text)

    def local(tag):
        return tag.split("}")[-1] if tag else tag

    result_code = result_msg = total_count = None
    for el in root.iter():
        t = local(el.tag).lower()
        if t == "resultcode" and el.text:
            result_code = el.text.strip()
        elif t == "resultmsg" and el.text:
            result_msg = el.text.strip()
        elif t == "totalcount" and el.text:
            total_count = el.text.strip()

    xml_items = []
    for item_el in root.iter():
        if local(item_el.tag) != "item":
            continue
        row = {}
        for child in list(item_el):
            row[local(child.tag)] = (child.text or "").strip()
        if row:
            xml_items.append(row)

    meta.update({
        "raw_xml": True,
        "resultCode": result_code,
        "resultMsg": result_msg,
        "totalCount": total_count,
        "count": len(xml_items),
    })
    return xml_items, meta


def load_cert_file_cache(ttl_hours: int = CERT_CACHE_TTL_HOURS):
    """
    디스크 파일 캐시 로드.
    유효하면 (items, meta) 반환, 없거나 만료면 (None, meta).
    """
    import gzip
    import json
    from datetime import datetime, timezone

    meta = {"path": str(CERT_CACHE_FILE), "ttl_hours": ttl_hours}
    if not CERT_CACHE_FILE.exists():
        meta["status"] = "missing"
        return None, meta
    try:
        with gzip.open(CERT_CACHE_FILE, "rt", encoding="utf-8") as f:
            payload = json.load(f)
        saved_at = payload.get("saved_at") or ""
        items = payload.get("items") or []
        meta["saved_at"] = saved_at
        meta["count"] = len(items)
        # 만료 검사
        try:
            # ISO 또는 timestamp
            if isinstance(saved_at, (int, float)):
                saved_ts = float(saved_at)
            else:
                saved_ts = datetime.fromisoformat(str(saved_at).replace("Z", "+00:00")).timestamp()
            age_h = (time.time() - saved_ts) / 3600.0
            meta["age_hours"] = round(age_h, 2)
            if age_h > ttl_hours:
                meta["status"] = "expired"
                return None, meta
        except Exception as e:
            meta["status"] = f"date_error:{e}"
            return None, meta
        if not items:
            meta["status"] = "empty"
            return None, meta
        meta["status"] = "ok"
        return items, meta
    except Exception as e:
        meta["status"] = f"load_error:{e}"
        return None, meta


def save_cert_file_cache(items: list):
    """인증프로그램 전체 목록을 gzip JSON으로 저장"""
    import gzip
    import json
    from datetime import datetime, timezone

    CACHE_DIR.mkdir(exist_ok=True)
    payload = {
        "saved_at": datetime.now(timezone.utc).isoformat(),
        "count": len(items or []),
        "items": items or [],
    }
    tmp = CERT_CACHE_FILE.with_suffix(".tmp.gz")
    with gzip.open(tmp, "wt", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False)
    tmp.replace(CERT_CACHE_FILE)
    return {"path": str(CERT_CACHE_FILE), "count": payload["count"], "saved_at": payload["saved_at"]}


def clear_cert_file_cache():
    if CERT_CACHE_FILE.exists():
        CERT_CACHE_FILE.unlink()
        return True
    return False


def fetch_cert_programs_page(
    service_key: str,
    page_no: int = 1,
    num_of_rows: int = 1000,
):
    """svc004 단건 페이지 (ctpvNm 등 필터 파라미터 사용 금지)"""
    import requests
    from urllib.parse import unquote

    url = "https://apis.data.go.kr/B552713/svc004/getYthTrActvtyCrtfnPrgmRegImplInfo"
    key = (service_key or "").strip()
    if "%" in key:
        try:
            key = unquote(key)
        except Exception:
            pass
    params = {
        "serviceKey": key,
        "pageNo": str(page_no),
        "numOfRows": str(num_of_rows),
    }
    res = requests.get(url, params=params, timeout=60, verify=False)
    text = res.text or ""
    ctype = res.headers.get("Content-Type") or ""
    if res.status_code >= 400 and not text.strip():
        res.raise_for_status()
    items, meta = _parse_svc004_response(text, ctype)
    meta["http_status"] = res.status_code
    # totalCount JSON에서 보강
    if text.strip().startswith("{"):
        try:
            import json
            data = json.loads(text)
            meta["totalCount"] = _data_go_total_count(data) or meta.get("totalCount")
        except Exception:
            pass
    return items, meta


def fetch_cert_programs_all(
    service_key: str,
    page_size: int = 1000,
    progress_callback=None,
    start_page: int = 1,
    existing_items: list = None,
):
    """
    svc004 전체 데이터 조회 (7만건+).
    - progress_callback(page, collected, total)
    - start_page / existing_items 로 이어받기 가능
    - 매 페이지마다 session에 부분 저장은 호출측에서 처리
    """
    all_items = list(existing_items) if existing_items else []
    page = max(1, int(start_page or 1))
    total = None
    last_meta = {}
    while True:
        items, meta = fetch_cert_programs_page(
            service_key, page_no=page, num_of_rows=page_size
        )
        last_meta = meta
        if total is None:
            try:
                total = int(meta.get("totalCount") or 0) or None
            except Exception:
                total = None
        if not items:
            break
        all_items.extend(items)
        if progress_callback:
            try:
                progress_callback(page, len(all_items), total or 0, all_items)
            except TypeError:
                try:
                    progress_callback(page, len(all_items), total or 0)
                except Exception:
                    pass
            except Exception:
                pass
        if total and len(all_items) >= total:
            break
        if len(items) < page_size:
            break
        page += 1
        if page > 500:
            break
    return all_items, {
        "pages": page,
        "totalCount": total or len(all_items),
        "count": len(all_items),
        "last": last_meta,
        "next_page": page + 1,
    }


def fetch_cert_programs(
    service_key: str,
    ctpv_nm: str = "",
    fclty_nm: str = "",
    page_no: int = 1,
    num_of_rows: int = 100,
):
    """하위 호환: 단건 페이지 + 클라이언트 시도 필터"""
    items, meta = fetch_cert_programs_page(service_key, page_no=page_no, num_of_rows=num_of_rows)
    if (ctpv_nm or "").strip() and items:
        n_ctpv = normalize_ctpv(ctpv_nm)
        items = [p for p in items if normalize_ctpv(p.get("ctpvNm") or "") == n_ctpv]
    if (fclty_nm or "").strip() and items:
        q = normalize_fclty_name(fclty_nm)
        items = [p for p in items if q and q in normalize_fclty_name(p.get("fcltyNm") or "")]
    meta["count"] = len(items)
    return items, meta


def normalize_fclty_name(name: str) -> str:
    """시설명 정규화 (공백·괄호·시도 접두만 정리. 센터/수련관/문화의집은 구분 유지)"""
    s = (name or "").strip()
    s = re.sub(r"\(.*?\)", "", s)
    # 앞에 붙는 시도·시군구 접두 제거 (시설 유형 명칭은 유지)
    for prefix in [
        "서울특별시", "부산광역시", "대구광역시", "인천광역시", "광주광역시",
        "대전광역시", "울산광역시", "세종특별자치시", "경기도", "강원특별자치도",
        "강원도", "충청북도", "충청남도", "전북특별자치도", "전라북도", "전라남도",
        "경상북도", "경상남도", "제주특별자치도", "제주도",
    ]:
        if s.startswith(prefix):
            s = s[len(prefix):].strip()
            break
    s = re.sub(r"[\s\-_/·.,]+", "", s)
    return s.lower()


def normalize_ctpv(name: str) -> str:
    """시도명 정규화 (특별자치 등 표기 차이 완화)"""
    s = (name or "").strip()
    s = s.replace("특별자치도", "도").replace("특별자치시", "시")
    s = s.replace("광역시", "").replace("특별시", "")
    s = re.sub(r"\s+", "", s)
    return s


def year_in_cert_period(yr: str, p: dict) -> bool:
    """조회 연도가 인증기간(시작~종료)에 포함되면 True. 날짜 없으면 True."""
    if not (yr or "").strip():
        return True
    y = yr.strip()[:4]
    if not y.isdigit():
        return True
    yi = int(y)
    bg = str(p.get("certPrdBgngYmd") or p.get("actvBgngYmd") or "").strip()
    ed = str(p.get("certPrdEndYmd") or p.get("actvEndYmd") or p.get("endYmd") or "").strip()
    if not bg and not ed:
        return True
    try:
        by = int(bg[:4]) if len(bg) >= 4 and bg[:4].isdigit() else None
        ey = int(ed[:4]) if len(ed) >= 4 and ed[:4].isdigit() else None
    except ValueError:
        return True
    if by is not None and ey is not None:
        return by <= yi <= ey
    if by is not None:
        return by <= yi
    if ey is not None:
        return yi <= ey
    return True


def build_map_links(name: str, lat, lon, addr: str) -> str:
    """네이버 지도만 · 주소 기반 새 탭 링크"""
    import urllib.parse

    if not (addr or "").strip():
        return ""
    q = urllib.parse.quote(addr.strip())
    return f"[네이버지도](https://map.naver.com/v5/search/{q})"


def region_to_ctpv(region: str) -> str:
    """진로맵 지역 키워드 → svc001 ctpvNm"""
    r = (region or "").strip()
    if not r:
        return ""
    mapping = {
        "서울": "서울특별시", "부산": "부산광역시", "대구": "대구광역시",
        "인천": "인천광역시", "광주": "광주광역시", "대전": "대전광역시",
        "울산": "울산광역시", "세종": "세종특별자치시", "경기": "경기도",
        "고양": "경기도", "성남": "경기도", "용인": "경기도", "수원": "경기도",
        "강원": "강원특별자치도", "충북": "충청북도", "충남": "충청남도",
        "전북": "전북특별자치도", "전남": "전라남도", "경북": "경상북도",
        "경남": "경상남도", "제주": "제주특별자치도",
    }
    for k, v in mapping.items():
        if k in r:
            return v
    if r.endswith(("특별시", "광역시", "특별자치시", "도", "특별자치도")):
        return r
    return r


def match_facility_record(org_name: str, facilities: list):
    """시설명 정규화 매칭 → 가장 적합한 시설 1건"""
    n = normalize_fclty_name(org_name)
    if not n or not facilities:
        return None
    exact, partial = None, None
    for f in facilities:
        if not isinstance(f, dict):
            continue
        fn = normalize_fclty_name(f.get("fcltyNm") or f.get("시설명") or "")
        if not fn:
            continue
        if fn == n:
            exact = f
            break
        if (n in fn or fn in n) and min(len(n), len(fn)) >= 4 and partial is None:
            partial = f
    return exact or partial


def _pick_field(row: dict, *keys) -> str:
    """대소문자 무시 필드 추출"""
    if not isinstance(row, dict):
        return ""
    lower = {str(k).lower(): v for k, v in row.items()}
    for k in keys:
        v = lower.get(str(k).lower())
        if v is None or v == "":
            continue
        s = str(v).strip()
        if s and s.lower() not in ("null", "none", "-", "nan"):
            return s
    return ""


def normalize_youth_program(row: dict) -> dict:
    """여가부 응답만으로 표시용 필드 정리 (svc001 매칭 없음)"""
    if not isinstance(row, dict):
        return {}
    name = _pick_field(
        row, "prgrmNm", "pgmNm", "programNm", "progrmNm", "pgm", "title", "프로그램명", "actvNm", "name", "atName"
    )
    facility = _pick_field(
        row, "organNm", "orgNm", "org", "기관명", "fcltyNm", "시설명", "operInstNm",
        "mainActvPlcNm", "insttNm", "placeNm"
    )
    target = _pick_field(row, "target", "참가대상", "trget", "trgt", "obj", "object")
    fee = _pick_field(row, "price", "fee", "cost", "참가비", "요금", "수강료", "ppPrtcpCst")
    time_info = _pick_field(
        row, "sdate", "actvBgngYmd", "actvTime", "time", "시간", "활동일자", "operDayCnt", "nldgActvHr"
    )
    end_info = _pick_field(row, "edate", "actvEndYmd", "endYmd", "종료")
    if end_info and time_info and end_info not in time_info:
        time_info = f"{time_info} ~ {end_info}"
    sido = _pick_field(row, "sido", "ctpvNm", "시도구분", "시도")
    sgg = _pick_field(row, "sigungu", "sggNm", "시군구구분", "시군구")
    addr = _pick_field(
        row, "addr", "address", "addr1", "주소", "roadAddr", "location", "mainActvPlcNm", "plc"
    )
    if not addr and (sido or sgg):
        addr = f"{sido} {sgg}".strip()
    tel = _pick_field(row, "telno", "tel", "전화", "전화번호", "fxno")
    return {
        "name": name,
        "facility": facility,
        "target": target,
        "fee": fee,
        "time": time_info,
        "addr": addr,
        "tel": tel,
        "sido": sido,
        "sgg": sgg,
        "raw": row,
    }


def region_to_zip_prefixes(region: str):
    """지역 키워드 → 법정동 앞자리(느슨한 매칭용)"""
    r = (region or "").strip()
    table = {
        "서울": ["11"], "부산": ["26"], "대구": ["27"], "인천": ["28"],
        "광주": ["29"], "대전": ["30"], "울산": ["31"], "세종": ["36"],
        "경기": ["41"], "고양": ["41"], "성남": ["41"], "용인": ["41"], "수원": ["41"],
        "강원": ["51", "42"], "충북": ["43"], "충남": ["44"],
        "전북": ["52", "45"], "전남": ["46"], "경북": ["47"], "경남": ["48"], "제주": ["50"],
    }
    out = []
    for k, prefixes in table.items():
        if k in r:
            out.extend(prefixes)
    return list(dict.fromkeys(out))


def school_to_age_range(school: str):
    """교급 → 대략 연령 (느슨)"""
    s = (school or "").strip()
    if s.startswith("초"):
        return 9, 13
    if s.startswith("중"):
        return 12, 16
    if s.startswith("고"):
        return 15, 19
    return 9, 34  # 청소년·청년 넓게


def fetch_ontong_youth_policies(
    api_key: str,
    keyword: str = "",
    region: str = "",
    school: str = "",
    page_num: int = 1,
    page_size: int = 20,
):
    """
    온통청년 정책 API
    https://www.youthcenter.go.kr/go/ythip/getPlcy
    필터: 연령·지역 느슨 적용
    """
    import requests
    key = (api_key or "").strip()
    if not key:
        return [], {"error": "YOUTH_CENTER_API_KEY(apiKeyNm) 없음"}

    params = {
        "apiKeyNm": key,
        "pageNum": page_num,
        "pageSize": page_size,
        "pageType": "1",
        "rtnType": "json",
    }
    if keyword:
        params["plcyKywdNm"] = keyword
        params["plcyNm"] = keyword

    try:
        res = requests.get(
            "https://www.youthcenter.go.kr/go/ythip/getPlcy",
            params=params,
            timeout=45,
            verify=False,
        )
        text = res.text or ""
        meta = {"http_status": res.status_code}
        data = {}
        try:
            data = res.json() if text.strip().startswith("{") else {}
        except Exception as e:
            meta["parse_error"] = str(e)
            meta["preview"] = text[:300]
            return [], meta

        # 응답 구조 유연 처리
        items = []
        for path in (
            data.get("youthPolicyList"),
            (data.get("result") or {}).get("youthPolicyList"),
            data.get("data"),
            data.get("list"),
        ):
            if isinstance(path, list):
                items = path
                break
            if isinstance(path, dict):
                # 단일 또는 래핑
                inner = path.get("youthPolicy") or path.get("item") or path
                if isinstance(inner, list):
                    items = inner
                    break
                if isinstance(inner, dict) and (inner.get("plcyNm") or inner.get("plcyNo")):
                    items = [inner]
                    break

        age_min_u, age_max_u = school_to_age_range(school)
        zip_pref = region_to_zip_prefixes(region)

        filtered = []
        for it in items:
            if not isinstance(it, dict):
                continue
            # 연령 느슨: 제한 없거나 구간 겹치면 OK
            try:
                tmin = int(str(it.get("sprtTrgtMinAge") or "0") or 0)
            except Exception:
                tmin = 0
            try:
                tmax = int(str(it.get("sprtTrgtMaxAge") or "99") or 99)
            except Exception:
                tmax = 99
            age_lmt = str(it.get("sprtTrgtAgeLmtYn") or "").upper()
            age_ok = True
            if age_lmt in ("Y", "1", "TRUE"):
                # 겹침: user_max >= tmin and user_min <= tmax
                age_ok = age_max_u >= tmin and age_min_u <= tmax
            # 지역 느슨: zip 비어 있으면 전국, 있으면 접두 일치
            zip_cd = str(it.get("zipCd") or "")
            region_ok = True
            if zip_pref and zip_cd.strip():
                codes = [c.strip() for c in zip_cd.replace(" ", ",").split(",") if c.strip()]
                region_ok = any(
                    any(c.startswith(p) or p.startswith(c[:2]) for p in zip_pref)
                    for c in codes
                ) or any(c in ("00000", "99", "99999") for c in codes)
            if age_ok and region_ok:
                filtered.append(it)

        # 너무 적으면 연령만 통과한 것 / 원본 상위
        if len(filtered) < 3:
            loose = []
            for it in items:
                if not isinstance(it, dict):
                    continue
                try:
                    tmin = int(str(it.get("sprtTrgtMinAge") or "0") or 0)
                    tmax = int(str(it.get("sprtTrgtMaxAge") or "99") or 99)
                except Exception:
                    tmin, tmax = 0, 99
                if age_max_u >= tmin and age_min_u <= tmax:
                    loose.append(it)
            filtered = loose or items[:8]

        meta["raw_count"] = len(items)
        meta["filtered"] = len(filtered)
        meta["age_range"] = [age_min_u, age_max_u]
        meta["zip_pref"] = zip_pref
        return filtered[:12], meta
    except Exception as e:
        return [], {"error": str(e)}


def match_programs_to_facility(fac_name: str, fac_ctpv: str, programs: list):
    """1) 시도(정규화)+시설명 완전일치 2) 포함관계"""
    n_fac = normalize_fclty_name(fac_name)
    n_ctpv = normalize_ctpv(fac_ctpv)
    exact, partial = [], []
    for p in programs:
        if not isinstance(p, dict):
            continue
        p_ctpv = normalize_ctpv(p.get("ctpvNm") or "")
        if n_ctpv and p_ctpv and n_ctpv != p_ctpv:
            continue
        n_p = normalize_fclty_name(p.get("fcltyNm") or "")
        if not n_p:
            continue
        if n_fac and n_p == n_fac:
            exact.append(p)
        elif n_fac and n_p and (n_fac in n_p or n_p in n_fac) and min(len(n_fac), len(n_p)) >= 4:
            partial.append(p)
    if exact:
        return exact, "높음"
    if partial:
        return partial, "중"
    return [], "없음"


def rewrite_question_for_search(provider, api_key, question: str) -> str:
    prompt = f"""다음 사용자 질문을 공공기관 규정/매뉴얼 검색용 질의로 바꿔주세요.
절차형이면 취소, 변경, 재신청, 사후결재, 복명, 결과보고, 기한, 서식, 메뉴경로를 보강.
짧게 한 줄만 출력.
질문: {question}"""
    try:
        return get_ai_response(provider, api_key, "규정 검색 질의 재작성기", prompt).strip()
    except Exception:
        return question


def cache_key(question: str) -> str:
    return hashlib.sha256(question.strip().encode("utf-8")).hexdigest()


def has_any_result() -> bool:
    return bool(
        st.session_state.get("doc_result")
        or st.session_state.get("draft_result")
        or st.session_state.get("meeting_result")
        or st.session_state.get("meeting_actions")
        or st.session_state.get("email_result")
        or st.session_state.get("reg_answer")
    )


# ---------- 세션 ----------
for k, v in {
    "logged_in": False, "username": "", "last_active": None, "prev_tool": None,
    "image_generating": False, "reg_generating": False,
    "reg_answer": "", "reg_hits": [], "reg_question": "", "reg_rewritten": "",
    "reg_error": "", "reg_progress": "", "reg_route": "", "qa_cache": {},
    "doc_result": "", "draft_result": "", "meeting_result": "", "email_result": "",
    "meeting_transcript": "", "meeting_actions": "",
    "carry_summary": "", "carry_draft": "", "qa_type": "일반",
    "fac_items": [], "fac_raw": None, "fac_ai_summary": "",
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ---------- 로그인 ----------
if not st.session_state.logged_in:
    st.markdown("<div style='height: 4.5rem;'></div>", unsafe_allow_html=True)
    outer_l, outer_c, outer_r = st.columns([0.15, 3.7, 0.15])
    with outer_c:
        t1, t2, t3 = st.columns([0.9, 2.0, 1.2])
        with t1:
            if (ASSETS / "login_ai.png").exists():
                st.image(str(ASSETS / "login_ai.png"), width=150)
        with t2:
            if (ASSETS / "logo.png").exists():
                st.image(str(ASSETS / "logo.png"), use_container_width=True)
            else:
                st.markdown("<h2 style='text-align:center;'>KYWA</h2>", unsafe_allow_html=True)
            st.markdown(
                """<div style="text-align:center; margin-top:0.5rem;">
                <span style="display:inline-block; background: rgba(148,163,184,0.25); color:#e5eefb;
                padding: 0.5rem 1.2rem; border-radius: 8px; font-size: 1.3rem; font-weight: 600;">
                한국청소년활동진흥원 · 데이터 융복합 서비스</span></div>""",
                unsafe_allow_html=True,
            )
        with t3:
            if (ASSETS / "login_bot.png").exists():
                st.image(str(ASSETS / "login_bot.png"), width=260)

    st.markdown("<div style='height: 2.0rem;'></div>", unsafe_allow_html=True)
    st.markdown('<div class="section-line"></div>', unsafe_allow_html=True)
    st.markdown("<div style='height: 1.0rem;'></div>", unsafe_allow_html=True)


    left, center, right = st.columns([1.0, 1.8, 1.0])
    with center:
        mode = st.radio(
            "계정",
            ["로그인", "회원가입"],
            horizontal=True,
            label_visibility="collapsed",
            key="auth_mode",
        )
        if mode == "로그인":
            st.markdown(
                '<div style="background:#1e3a5f;border:1px solid #3b82f6;border-radius:10px;'
                'padding:0.55rem 0.85rem;margin-bottom:0.6rem;color:#bfdbfe;font-weight:700;">'
                '🔑 로그인</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                """<style>
                div[data-testid="stForm"] button {
                  background: #2563eb !important; color: #fff !important;
                  border: 2px solid #93c5fd !important; font-weight: 800 !important;
                }
                </style>""",
                unsafe_allow_html=True,
            )
            with st.form("login_form_v2"):
                login_id = st.text_input("아이디", key="f_login_id")
                login_pw = st.text_input("비밀번호", type="password", key="f_login_pw")
                login_submit = st.form_submit_button("로그인", type="primary", use_container_width=True)
                if login_submit:
                    lid = (st.session_state.get("f_login_id") or login_id or "").strip()
                    lpw = (st.session_state.get("f_login_pw") or login_pw or "").strip()
                    if not lid or not lpw:
                        st.warning("아이디와 비밀번호를 입력해주세요.")
                    else:
                        users = load_users()
                        if lid in users and users[lid].get("password") == hash_pw(lpw):
                            u = users[lid]
                            is_admin = (lid == ADMIN_USER) or (u.get("role") == "admin")
                            approved = is_admin or bool(u.get("approved", False))
                            if not approved:
                                st.warning("관리자 승인 대기 중입니다. 승인 후 로그인할 수 있습니다.")
                            else:
                                st.session_state.logged_in = True
                                st.session_state.username = lid
                                st.session_state.last_active = time.time()
                                try:
                                    inc_login_count()
                                except Exception:
                                    pass
                                st.rerun()
                        else:
                            st.error("아이디 또는 비밀번호가 올바르지 않습니다.")
        else:
            st.markdown(
                '<div style="background:#14532d;border:1px solid #22c55e;border-radius:10px;'
                'padding:0.55rem 0.85rem;margin-bottom:0.6rem;color:#bbf7d0;font-weight:700;">'
                '📝 회원가입 (관리자 승인 후 이용)</div>',
                unsafe_allow_html=True,
            )
            st.markdown(
                """<style>
                div[data-testid="stForm"] button {
                  background: #16a34a !important; color: #fff !important;
                  border: 2px solid #86efac !important; font-weight: 800 !important;
                }
                </style>""",
                unsafe_allow_html=True,
            )
            st.caption("내부 사용자 계정 생성 · 관리자 승인 후 로그인 가능")
            with st.form("signup_form_v2"):
                signup_id = st.text_input("새 아이디", key="f_signup_id")
                signup_pw = st.text_input("새 비밀번호", type="password", key="f_signup_pw")
                signup_pw2 = st.text_input("비밀번호 확인", type="password", key="f_signup_pw2")
                signup_submit = st.form_submit_button("회원가입 신청", type="primary", use_container_width=True)
                if signup_submit:
                    sid = (st.session_state.get("f_signup_id") or signup_id or "").strip()
                    spw = (st.session_state.get("f_signup_pw") or signup_pw or "").strip()
                    spw2 = (st.session_state.get("f_signup_pw2") or signup_pw2 or "").strip()
                    if not sid or not spw:
                        st.warning("아이디와 비밀번호를 입력해주세요.")
                    elif len(spw) < 4:
                        st.warning("비밀번호는 4자 이상 입력해주세요.")
                    elif spw != spw2:
                        st.error("비밀번호 확인이 일치하지 않습니다.")
                    else:
                        users = load_users()
                        if sid in users:
                            st.error("이미 존재하는 아이디입니다.")
                        else:
                            try:
                                users[sid] = {
                                    "password": hash_pw(spw),
                                    "name": sid,
                                    "approved": True if sid == ADMIN_USER else False,
                                    "role": "admin" if sid == ADMIN_USER else "user",
                                }
                                save_users(users)
                                if sid == ADMIN_USER:
                                    st.success("관리자 계정이 생성되었습니다. 로그인으로 전환해 접속하세요.")
                                else:
                                    st.success("가입 신청이 접수되었습니다. 관리자 승인 후 로그인할 수 있습니다.")
                            except Exception as e:
                                st.error(f"저장 중 오류: {e}")
    st.markdown("<div style='height: 3.5rem;'></div>", unsafe_allow_html=True)
    st.stop()


# ---------- 사이드바 ----------
if (ASSETS / "logo.png").exists():
    st.sidebar.image(str(ASSETS / "logo.png"), use_container_width=True)
else:
    st.sidebar.markdown("### 🔗 KYWA 데이터")

st.sidebar.caption(f"로그인: {st.session_state.username}")
if st.session_state.get("last_active"):
    remain_sec = max(0, int(SESSION_TIMEOUT_MIN * 60 - (time.time() - st.session_state.last_active)))
    st.sidebar.caption(f"세션 남은 시간: {remain_sec // 60:02d}분 {remain_sec % 60:02d}초")

if st.sidebar.button("로그아웃", use_container_width=True, key="btn_logout"):
    for k in list(st.session_state.keys()):
        del st.session_state[k]
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.markdown("#### 사용할 도구")

# 데이터 융복합 전용 메뉴
MENU_TOP = [
    "🏠 처음 화면",
]
MENU_FUSION = "🔗 KYWA 데이터 융복합 서비스"
MENU_FUSION_SUB = [
    "🏕️ 수련시설 안전 및 인증프로그램 현황",
    "🧭 청소년 통합 지원 서비스",
]
MENU_BOTTOM = [
    "🔐 관리자 · 가입승인 / 비밀번호",
]

if "nav_tool" not in st.session_state:
    st.session_state.nav_tool = "🏠 처음 화면"

_fusion_open = (
    st.session_state.nav_tool == MENU_FUSION
    or st.session_state.nav_tool in MENU_FUSION_SUB
)

def _nav_item(label: str, indent: bool = False):
    active = st.session_state.nav_tool == label
    display = f"　└ {label}" if indent else label
    btn_type = "primary" if active else "secondary"
    if st.sidebar.button(
        display,
        key=f"nav_{'sub_' if indent else ''}{label}",
        use_container_width=True,
        type=btn_type,
    ):
        if label == MENU_FUSION and MENU_FUSION_SUB:
            st.session_state.nav_tool = MENU_FUSION_SUB[0]
        else:
            st.session_state.nav_tool = label
        st.rerun()

for m in MENU_TOP:
    _nav_item(m)
_nav_item(MENU_FUSION)
if _fusion_open:
    for s in MENU_FUSION_SUB:
        _nav_item(s, indent=True)
for m in MENU_BOTTOM:
    _nav_item(m)

tool = st.session_state.nav_tool
main_tool = (
    MENU_FUSION
    if (tool == MENU_FUSION or tool in MENU_FUSION_SUB)
    else tool
)
if tool == MENU_FUSION and MENU_FUSION_SUB:
    tool = MENU_FUSION_SUB[0]

# 세션 만료 + 자동새로고침
# ※ 긴 API 조회(수련시설 등) 중에는 자동새로고침·세션만료를 멈춤
if st.session_state.logged_in:
    now = time.time()
    # 장기 작업 중이면 활동시간 계속 갱신 (로그아웃 방지)
    if st.session_state.get("fac_fetching") or st.session_state.get("long_job"):
        st.session_state.last_active = now
    last = st.session_state.get("last_active")
    if last is None:
        st.session_state.last_active = now
    elif now - last > SESSION_TIMEOUT_MIN * 60:
        st.session_state.logged_in = False
        st.session_state.username = ""
        st.session_state.last_active = None
        st.warning(f"{SESSION_TIMEOUT_MIN}분 동안 사용이 없어 자동 로그아웃되었습니다.")
        st.rerun()
    # 자동새로고침: 장기 작업·결과 표시·특정 메뉴에서는 중지
    if (
        not st.session_state.get("reg_generating", False)
        and not st.session_state.get("fac_fetching", False)
        and not st.session_state.get("long_job", False)
        and not has_any_result()
        and tool not in (
            "🏕️ 수련시설 안전 및 인증프로그램 현황",
            "🧭 청소년 통합 지원 서비스",
            "🔗 KYWA 데이터 융복합 서비스",
        )
        and main_tool != "🔗 KYWA 데이터 융복합 서비스"
    ):
        st_autorefresh(interval=60 * 1000, key="session_refresh")

if "prev_scroll_tool" not in st.session_state:
    st.session_state.prev_scroll_tool = tool
if st.session_state.prev_scroll_tool != tool:
    st.session_state.prev_scroll_tool = tool
    components.html(
        """<script>
        const main = window.parent.document.querySelector('section.main');
        if (main) main.scrollTo({top: 0, behavior: 'instant'});
        window.parent.scrollTo(0, 0);
        </script>""",
        height=0,
    )

if st.session_state.prev_tool is None:
    st.session_state.prev_tool = tool
elif tool != st.session_state.prev_tool:
    touch_session()
    st.session_state.prev_tool = tool

st.sidebar.markdown("---")
st.sidebar.markdown("### AI 설정")
provider = st.sidebar.selectbox("AI 제공자 선택", ["OpenAI (ChatGPT)", "Claude (Anthropic)", "Groq"])
api_key = ""
try:
    if provider == "OpenAI (ChatGPT)":
        api_key = st.secrets["OPENAI_API_KEY"]
    elif provider == "Claude (Anthropic)":
        api_key = st.secrets["ANTHROPIC_API_KEY"]
    else:
        api_key = st.secrets["GROQ_API_KEY"]
except Exception:
    api_key = ""
if api_key:
    st.sidebar.success("API 키 로드 완료 (서버 저장)")
else:
    st.sidebar.error("API 키가 설정되지 않았습니다. secrets.toml 확인")

if st.session_state.username == ADMIN_USER:
    st.sidebar.markdown("---")
    st.sidebar.markdown("### 사용실적")
    usage = load_usage()
    st.sidebar.metric("누적 로그인", f"{usage.get('total_logins', 0):,}회")
    st.sidebar.caption("메뉴별 AI 실행")
    menu_counts = usage.get("menu_counts", {})
    if menu_counts:
        for name, cnt in sorted(menu_counts.items(), key=lambda x: x[1], reverse=True):
            st.sidebar.write(f"· {name}: **{cnt}**회")
    else:
        st.sidebar.caption("아직 실행 기록이 없습니다.")

st.sidebar.markdown("---")
st.sidebar.caption("디지털정보부 · 데이터 융복합 서비스")

# ---------- 메인 ----------
if tool == "🏠 처음 화면":
    st.markdown("<div style='height: 3.5rem;'></div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1.1, 2.6, 1.1])
    with c1:
        if (ASSETS / "kiumee.png").exists():
            st.image(str(ASSETS / "kiumee.png"), use_container_width=True)
    with c2:
        st.markdown("""
        <div style="text-align:center; padding-top: 0.4rem;">
            <div class="main-title">KYWA<br>데이터 융복합 서비스</div>
            <div class="sub-title">한국청소년활동진흥원 · 디지털정보부</div>
        </div>
        """, unsafe_allow_html=True)
        if st.session_state.get("username"):
            st.markdown(
                f'<div style="text-align:center;"><span class="badge-ok">로그인: {st.session_state.username}</span></div>',
                unsafe_allow_html=True,
            )
    with c3:
        if (ASSETS / "irumee.png").exists():
            st.image(str(ASSETS / "irumee.png"), use_container_width=True)
    st.markdown('<div class="section-line"></div>', unsafe_allow_html=True)
    a, b = st.columns(2)
    with a:
        st.markdown(
            '<div class="feature-card"><h4>🏕️ 수련시설 안전·인증</h4>'
            '<p>안전평가 + 인증프로그램 매칭 조회</p></div>',
            unsafe_allow_html=True,
        )
    with b:
        st.markdown(
            '<div class="feature-card"><h4>🧭 청소년 통합 지원</h4>'
            '<p>대화형 진로·체험·활동 안내 도우미</p></div>',
            unsafe_allow_html=True,
        )
    st.markdown("<div style='height: 1.25rem;'></div>", unsafe_allow_html=True)
    st.info("왼쪽 메뉴에서 서비스를 선택하세요. 회원가입 후 관리자 승인이 필요합니다.")







elif tool == "🔐 관리자 · 가입승인 / 비밀번호":
    page_header("🔐 관리자 · 가입승인 / 비밀번호", "관리자 전용")
    if st.session_state.username != ADMIN_USER:
        st.error("관리자만 접근할 수 있습니다.")
    else:
        st.markdown("### 회원가입 승인")
        st.caption(
            "⚠️ Streamlit Cloud는 앱 재배포 시 로컬 회원파일이 초기화될 수 있습니다. "
            "승인 후 아래 **회원 JSON 내보내기**를 받아 secrets의 USERS_JSON에 붙여 두면 복구됩니다."
        )

        users = load_users()
        # 관리자 계정 자동 승인 플래그
        if ADMIN_USER in users:
            users[ADMIN_USER]["approved"] = True
            users[ADMIN_USER]["role"] = "admin"
            save_users(users)
        pending = [
            uid for uid, u in users.items()
            if uid != ADMIN_USER and not u.get("approved", False)
        ]
        approved_users = [
            uid for uid, u in users.items()
            if uid != ADMIN_USER and u.get("approved", False)
        ]
        st.write(f"승인 대기 **{len(pending)}**명 · 승인 완료 **{len(approved_users)}**명")
        if pending:
            for uid in pending:
                c1, c2, c3 = st.columns([3, 1, 1])
                with c1:
                    st.markdown(f"**{uid}**")
                with c2:
                    if st.button("승인", key=f"appr_{uid}"):
                        users[uid]["approved"] = True
                        save_users(users)
                        st.success(f"{uid} 승인 완료")
                        st.rerun()
                with c3:
                    if st.button("거절·삭제", key=f"rej_{uid}"):
                        del users[uid]
                        save_users(users)
                        st.warning(f"{uid} 삭제")
                        st.rerun()
        else:
            st.caption("대기 중인 가입 신청이 없습니다.")
        if approved_users:
            with st.expander("승인된 사용자"):
                for uid in approved_users:
                    c1, c2 = st.columns([3, 1])
                    with c1:
                        st.write(uid)
                    with c2:
                        if st.button("승인 취소", key=f"unappr_{uid}"):
                            users[uid]["approved"] = False
                            save_users(users)
                            st.rerun()
        st.markdown("---")
        st.markdown("### 사용실적")
        usage = load_usage()
        c1, c2 = st.columns(2)
        with c1:
            st.metric("누적 로그인", f"{usage.get('total_logins', 0):,}회")
        with c2:
            st.metric("AI 실행 합계", f"{sum(usage.get('menu_counts', {}).values()):,}회")
        st.markdown("---")
        st.markdown("### 회원 데이터 백업·복구")
        users_now = load_users()
        st.download_button(
            "📥 회원 JSON 내보내기",
            data=json.dumps(users_now, ensure_ascii=False, indent=2),
            file_name="users_backup.json",
            mime="application/json",
            key="dl_users_json",
        )
        up = st.file_uploader("회원 JSON 가져오기 (복구)", type=["json"], key="up_users_json")
        if up is not None and st.button("가져오기 적용", key="btn_import_users"):
            try:
                imported = json.loads(up.read().decode("utf-8"))
                if not isinstance(imported, dict):
                    st.error("JSON 최상위는 객체(dict)여야 합니다.")
                else:
                    cur = load_users()
                    cur.update(imported)
                    save_users(cur)
                    st.success(f"복구 완료: {len(imported)}명 병합")
                    st.rerun()
            except Exception as e:
                st.error(f"가져오기 실패: {e}")
        st.markdown("---")
        st.markdown("### 비밀번호 재설정")

        users = load_users()
        if users:
            target_id = st.selectbox("사용자", list(users.keys()))
            new_pw = st.text_input("새 비밀번호", type="password", key="admin_new_pw")
            new_pw2 = st.text_input("확인", type="password", key="admin_new_pw2")
            if st.button("재설정", type="primary", key="btn_admin_reset"):
                touch_session()
                if new_pw and new_pw == new_pw2:
                    users[target_id]["password"] = hash_pw(new_pw)
                    save_users(users)
                    st.success(f"{target_id} 재설정 완료")
                else:
                    st.error("비밀번호를 확인하세요.")

elif tool == "🏕️ 수련시설 안전 및 인증프로그램 현황":
    page_header("🏕️ 수련시설 안전 및 인증프로그램 현황")
    show_examples([
        "시도·연도 필수 입력 후 조회",
        "시설명 정규화 매칭으로 인증프로그램 연결",
        "주소·좌표 기반 지도 새 탭 링크",
    ])

    try:
        data_key = st.secrets.get("DATA_GO_KR_KEY", "")
    except Exception:
        data_key = ""
    try:
        data_key_004 = st.secrets.get("DATA_GO_KR_KEY_SVC004", "") or data_key
    except Exception:
        data_key_004 = data_key

    if not data_key:
        st.warning("`.streamlit/secrets.toml` 에 DATA_GO_KR_KEY 를 설정하세요.")
    st.caption("svc004 키가 다르면 DATA_GO_KR_KEY_SVC004 를 추가로 설정할 수 있습니다.")

    c1, c2, c3 = st.columns(3)
    with c1:
        ctpv = st.selectbox(
            "시도 *",
            [
                "",
                "서울특별시", "부산광역시", "대구광역시", "인천광역시",
                "광주광역시", "대전광역시", "울산광역시", "세종특별자치시",
                "경기도", "강원특별자치도", "충청북도", "충청남도",
                "전북특별자치도", "전라남도", "경상북도", "경상남도", "제주특별자치도",
            ],
            key="fac_ctpv",
        )
    with c2:
        year = st.text_input("평가연도 *", placeholder="예: 2023", key="fac_year")
    with c3:
        fclty = st.text_input("시설명 (선택)", key="fac_name")

    page_size = st.slider("한페이지에 조회건수", 1, 100, 20, 1, key="fac_page_size")

    # 파일 캐시 상태 + 강제 갱신
    _fi, _fm = load_cert_file_cache()
    if _fm.get("status") == "ok":
        st.caption(
            f"📁 인증프로그램 파일 캐시: {_fm.get('count', 0):,}건 · "
            f"경과 {_fm.get('age_hours', '?')}시간 / 유효 {CERT_CACHE_TTL_HOURS}시간"
        )
    else:
        st.caption(f"📁 인증프로그램 파일 캐시: 없음 ({_fm.get('status')})")
    if st.button("인증프로그램 캐시 강제 갱신", key="btn_fac_cache_clear"):
        clear_cert_file_cache()
        st.session_state.pop("fac_cert_all_cache", None)
        st.session_state.pop("fac_cert_all_done", None)
        st.session_state["fac_cert_next_page"] = 1
        st.success("캐시를 삭제했습니다. 다음 조회 시 API에서 다시 받습니다.")
        st.rerun()

    if st.button("조회하기", type="primary", key="btn_fac_search"):
        touch_session()
        st.session_state["fac_ai_summary"] = ""
        st.session_state["fac_merged_rows"] = []
        st.session_state["fac_prog_detail"] = {}
        st.session_state["fac_page"] = 1
        if not data_key:
            st.error("API 키가 없습니다.")
        elif not (ctpv or "").strip() or not (year or "").strip():
            st.warning("시도·연도는 필수입니다.")
        else:
            safety_items, safety_raw = [], None
            prog_items, prog_raw = [], None
            err_msgs = []

            # 1) 안전평가: 평가연도 해당 시설 전체
            try:
                st.session_state["fac_fetching"] = True
                st.session_state["long_job"] = True
                touch_session()
                with st.spinner("안전평가 시설 전체 조회 중..."):
                    safety_items, safety_raw = fetch_training_facility_safety_all(
                        data_key,
                        ctpv_nm=ctpv.strip(),
                        fclty_nm=(fclty or "").strip(),
                        evl_yr=year.strip(),
                        page_size=100,
                    )
                touch_session()
            except Exception as e:
                err_msgs.append(f"안전평가 API: {type(e).__name__}: {e}")

            # 2) 인증프로그램: 파일캐시(24h) → 세션 → API 전체 조회
            try:
                st.session_state["fac_fetching"] = True
                st.session_state["long_job"] = True
                touch_session()

                prog_items = None
                prog_raw = {}

                # (1) 세션 메모리 (같은 로그인 세션)
                mem = st.session_state.get("fac_cert_all_cache")
                if st.session_state.get("fac_cert_all_done") and isinstance(mem, list) and len(mem) > 1000:
                    prog_items = mem
                    prog_raw = {"source": "session", "count": len(mem)}

                # (2) 디스크 파일 캐시 (24시간)
                if prog_items is None:
                    file_items, file_meta = load_cert_file_cache()
                    if file_items:
                        prog_items = file_items
                        prog_raw = {"source": "file", **file_meta}
                        st.session_state["fac_cert_all_cache"] = file_items
                        st.session_state["fac_cert_all_done"] = True
                        st.info(
                            f"파일 캐시 사용 · {len(file_items):,}건 · "
                            f"저장 후 {file_meta.get('age_hours', '?')}시간 경과 "
                            f"(유효 {CERT_CACHE_TTL_HOURS}시간)"
                        )
                    else:
                        st.caption(
                            f"파일 캐시 없음/만료: {file_meta.get('status')} "
                            f"→ API 전체 조회를 시작합니다."
                        )

                # (3) API 전체 조회 (+ 이어받기)
                if prog_items is None:
                    prog_bar = st.progress(0, text="인증프로그램 전체 조회 중... (자동새로고침 중지)")
                    status_txt = st.empty()
                    existing = mem if isinstance(mem, list) else []
                    start_pg = int(st.session_state.get("fac_cert_next_page") or 1)
                    if existing and start_pg > 1:
                        status_txt.info(f"이전 수집 {len(existing):,}건부터 이어받습니다. (페이지 {start_pg}~)")

                    def _prog(page, collected, total, items_so_far=None):
                        touch_session()
                        st.session_state["fac_fetching"] = True
                        if items_so_far is not None:
                            st.session_state["fac_cert_all_cache"] = list(items_so_far)
                            st.session_state["fac_cert_next_page"] = page + 1
                            st.session_state["fac_cert_all_done"] = False
                        if total:
                            prog_bar.progress(
                                min(1.0, collected / max(total, 1)),
                                text=f"인증프로그램 조회 {collected:,} / {total:,} (페이지 {page})",
                            )
                        else:
                            prog_bar.progress(
                                min(0.99, page / 100),
                                text=f"인증프로그램 조회 {collected:,}건 (페이지 {page})",
                            )
                        status_txt.caption(f"페이지 {page} · 수집 {collected:,}건 · 세션 유지 중")

                    prog_items, prog_raw = fetch_cert_programs_all(
                        data_key_004,
                        page_size=1000,
                        progress_callback=_prog,
                        start_page=start_pg if existing else 1,
                        existing_items=existing if existing else None,
                    )
                    st.session_state["fac_cert_all_cache"] = prog_items
                    st.session_state["fac_cert_all_done"] = True
                    st.session_state["fac_cert_next_page"] = 1
                    # 디스크 저장
                    try:
                        save_info = save_cert_file_cache(prog_items)
                        prog_raw = {**(prog_raw or {}), "source": "api", "file_cache": save_info}
                        st.success(
                            f"인증프로그램 {len(prog_items):,}건 조회·파일 캐시 저장 완료 "
                            f"({CERT_CACHE_TTL_HOURS}시간 유효)"
                        )
                    except Exception as se:
                        st.warning(f"파일 캐시 저장 실패(조회 결과는 사용 가능): {se}")
                    prog_bar.progress(1.0, text=f"인증프로그램 조회 완료 {len(prog_items):,}건")
            except Exception as e:
                err_msgs.append(f"인증프로그램 API: {type(e).__name__}: {e}")
                partial = st.session_state.get("fac_cert_all_cache") or []
                if partial:
                    err_msgs.append(
                        f"부분 수집 {len(partial):,}건 보관됨. 다시「조회하기」를 누르면 이어서 받습니다."
                    )
                prog_items, prog_raw = (partial if partial else []), {"error": str(e)}
            finally:
                st.session_state["fac_fetching"] = False
                st.session_state["long_job"] = False
                touch_session()

            st.session_state["fac_raw"] = {"safety": safety_raw, "program": prog_raw}
            for m in err_msgs:
                st.warning(m)

            if not prog_items:
                st.warning("인증프로그램 API 결과가 0건입니다.")
                with st.expander("프로그램 API 디버그", expanded=True):
                    st.write(prog_raw if isinstance(prog_raw, dict) else str(prog_raw))

            # 매칭용: 전체 인증프로그램 사용 (시설명 기준 매칭)
            filtered_prog = [p for p in prog_items if isinstance(p, dict)]
            # 시도가 있으면 클라이언트 필터로 후보 축소
            if (ctpv or "").strip():
                n_ctpv = normalize_ctpv(ctpv)
                sido_prog = [
                    p for p in filtered_prog
                    if normalize_ctpv(p.get("ctpvNm") or "") == n_ctpv
                ]
                if sido_prog:
                    filtered_prog = sido_prog

            def _fmt_ymd(s):
                s = re.sub(r"[^0-9]", "", str(s or ""))
                if len(s) >= 8:
                    return f"{s[:4]}-{s[4:6]}-{s[6:8]}"
                return s or ""

            def _expire_status(ed: str, query_yr: str) -> str:
                """인증만료일 연도 < 조회연도 → 만료, 그 외 인증중"""
                digits = re.sub(r"[^0-9]", "", str(ed or ""))
                qy = re.sub(r"[^0-9]", "", str(query_yr or ""))[:4]
                if len(digits) < 4 or not qy:
                    return "-"
                try:
                    ey = int(digits[:4])
                    qyi = int(qy)
                except ValueError:
                    return "-"
                return "만료" if ey < qyi else "인증중"

            yr = year.strip()
            merged = []
            detail_map = {}
            for it in safety_items:
                if not isinstance(it, dict):
                    continue
                name = it.get("fcltyNm") or ""
                sido = it.get("ctpvNm") or ctpv
                addr = f"{it.get('addr1') or ''} {it.get('addr2') or ''}".strip()
                matched, conf = match_programs_to_facility(name, sido, filtered_prog)

                # 프로그램별 줄바꿈 표기
                prog_lines, type_lines, expire_lines, status_lines = [], [], [], []
                seen_pn = set()
                for p in matched:
                    pn = (p.get("prgrmNm") or "").strip()
                    if not pn or pn in seen_pn:
                        continue
                    seen_pn.add(pn)
                    prog_lines.append(pn)
                    type_lines.append((p.get("actvTypeNm") or "").strip() or "-")
                    ed = str(
                        p.get("certPrdEndYmd") or p.get("endYmd") or p.get("actvEndYmd") or ""
                    ).strip()
                    expire_lines.append(_fmt_ymd(ed) if ed else "-")
                    status_lines.append(_expire_status(ed, yr))

                prog_names = "\n".join(prog_lines) if prog_lines else "-"
                type_summary = "\n".join(type_lines) if type_lines else "-"
                expire_summary = "\n".join(expire_lines) if expire_lines else "-"
                status_summary = "\n".join(status_lines) if status_lines else "-"

                map_md = build_map_links(name, it.get("lat"), it.get("lot"), addr)
                row = {
                    "시도": sido,
                    "시설명": name,
                    "평가연도": it.get("evlYr") or year,
                    "시설종류": it.get("fcltKnNm") or "",
                    "종합등급": it.get("snthsEvlGrd") or "",
                    "건축": it.get("archEvlGrd") or "",
                    "전기": it.get("prvtmEvlGrd") or "",
                    "토목": it.get("civilEvlGrd") or "",
                    "소방": it.get("ffEvlGrd") or "",
                    "가스": it.get("gasEvlGrd") or "",
                    "위생": it.get("sanEvlGrd") or "",
                    "기계": it.get("mchnEvlGrd") or "",
                    "수용인원": it.get("actcNope") if it.get("actcNope") is not None else "",
                    "숙박정원": it.get("ldgPscpCnt") if it.get("ldgPscpCnt") is not None else "",
                    "전화": it.get("telno") or "",
                    "주소": addr,
                    "운영기관": it.get("operInstNm") or "",
                    "인증프로그램명": prog_names,
                    "활동유형": type_summary,
                    "인증만료일": expire_summary,
                    "인증만료여부": status_summary,
                    "매칭신뢰도": conf,
                    "지도": map_md,
                    "인증프로그램수": len(prog_lines),
                }
                merged.append(row)
                detail_map[name] = matched

            st.session_state["fac_merged_rows"] = merged
            st.session_state["fac_prog_detail"] = detail_map
            st.session_state["fac_items"] = safety_items
            st.session_state["fac_page"] = 1
            if merged:
                inc_menu_count("수련시설 안전평가 및 인증현황 조회")
                matched_n = sum(1 for r in merged if (r.get("인증프로그램수") or 0) > 0)
                st.success(
                    f"시설 {len(merged)}건 · 인증프로그램 {len(prog_items):,}건 · "
                    f"매칭된 시설 {matched_n}건"
                )
            else:
                st.warning("표시할 시설이 없습니다.")
                with st.expander("원본 응답(디버그)"):
                    st.json({"safety": safety_raw, "program": prog_raw})

    merged = st.session_state.get("fac_merged_rows") or []
    if merged:
        import pandas as pd
        import math

        page_size = int(st.session_state.get("fac_page_size") or page_size or 20)
        page_size = max(1, min(100, page_size))
        total = len(merged)
        total_pages = max(1, math.ceil(total / page_size))
        if "fac_page" not in st.session_state:
            st.session_state["fac_page"] = 1
        # 페이지 범위 보정
        cur = int(st.session_state.get("fac_page") or 1)
        if cur < 1:
            cur = 1
        if cur > total_pages:
            cur = total_pages
        st.session_state["fac_page"] = cur

        start = (cur - 1) * page_size
        end = min(start + page_size, total)
        page_rows = merged[start:end]

        st.markdown("### 조회 결과")
        st.caption(f"전체 {total}건 · {cur}/{total_pages}페이지 · 이 페이지 {len(page_rows)}건")

        header = [
            "시도", "시설명", "평가연도", "종합등급", "건축", "전기", "소방", "가스",
            "인증프로그램명", "활동유형", "인증만료일", "인증만료여부", "전화", "주소", "지도",
        ]
        multiline_cols = {"인증프로그램명", "활동유형", "인증만료일", "인증만료여부"}
        html = [
            '<div style="overflow-x:auto; width:100%; margin:8px 0 16px 0;">',
            '<table style="border-collapse:collapse; width:max-content; min-width:100%; '
            'font-size:12px; line-height:1.35; color:#e8e8e8;">',
            '<thead><tr>',
        ]
        for h in header:
            html.append(
                f'<th style="border:1px solid #555; background:#1e1e2e; padding:6px 8px; '
                f'white-space:nowrap; text-align:center; font-weight:600;">{h}</th>'
            )
        html.append("</tr></thead><tbody>")
        for r in page_rows:
            html.append("<tr>")
            for h in header:
                v = r.get(h, "")
                if h == "지도":
                    cell = str(v) if v else "-"
                    if cell.startswith("[") and "](" in cell:
                        try:
                            label = cell.split("](")[0][1:]
                            href = cell.split("](")[1].rstrip(")")
                            cell = (
                                f'<a href="{href}" target="_blank" rel="noopener" '
                                f'style="color:#7eb8ff;">{label}</a>'
                            )
                        except Exception:
                            pass
                else:
                    cell = str(v) if v is not None else ""
                    if not cell:
                        cell = "-"
                    # 줄바꿈 → <br>, HTML 이스케이프
                    cell = (
                        cell.replace("&", "&amp;")
                        .replace("<", "&lt;")
                        .replace(">", "&gt;")
                        .replace("\n", "<br>")
                    )
                nowrap = h not in multiline_cols and h not in ("시설명", "주소")
                ws = "nowrap" if nowrap else "normal"
                html.append(
                    f'<td style="border:1px solid #555; padding:5px 8px; white-space:{ws}; '
                    f'vertical-align:top;">{cell}</td>'
                )
            html.append("</tr>")
        html.append("</tbody></table></div>")
        st.markdown("".join(html), unsafe_allow_html=True)

        # 페이지 탭
        st.markdown("---")
        st.caption("페이지 이동")
        # 한 줄에 최대 10개 버튼
        btn_cols = st.columns(min(10, total_pages) + 2)
        with btn_cols[0]:
            if st.button("◀", key="fac_prev", disabled=(cur <= 1)):
                st.session_state["fac_page"] = cur - 1
                st.rerun()
        # 현재 구간 페이지 번호
        window = 10
        start_p = max(1, cur - window // 2)
        end_p = min(total_pages, start_p + window - 1)
        start_p = max(1, end_p - window + 1)
        for i, p in enumerate(range(start_p, end_p + 1)):
            with btn_cols[min(i + 1, len(btn_cols) - 2)]:
                label = f"[{p}]" if p == cur else str(p)
                if st.button(label, key=f"fac_page_btn_{p}"):
                    st.session_state["fac_page"] = p
                    st.rerun()
        with btn_cols[-1]:
            if st.button("▶", key="fac_next", disabled=(cur >= total_pages)):
                st.session_state["fac_page"] = cur + 1
                st.rerun()

        with st.expander("전체 컬럼 표 (다운로드용)"):
            df_show = pd.DataFrame(merged).drop(columns=["지도"], errors="ignore")
            st.dataframe(df_show, use_container_width=True, hide_index=True)

        df_dl = pd.DataFrame(merged).drop(columns=["지도"], errors="ignore")
        st.download_button(
            "CSV 다운로드 (전체)",
            data=df_dl.to_csv(index=False).encode("utf-8-sig"),
            file_name="수련시설_안전평가_인증현황.csv",
            mime="text/csv",
            key="dl_fac_csv",
        )

        # 시설별 프로그램 상세
        names = [r["시설명"] for r in merged if (r.get("인증프로그램수") or 0) > 0]
        if names:
            pick = st.selectbox("인증프로그램 상세 보기 (시설)", names, key="fac_detail_pick")
            detail = (st.session_state.get("fac_prog_detail") or {}).get(pick) or []
            if detail:
                drows = []
                for p in detail:
                    drows.append({
                        "프로그램명": p.get("prgrmNm") or "",
                        "인증번호": p.get("certNo") or "",
                        "활동유형": p.get("actvTypeNm") or "",
                        "인증시작": p.get("certPrdBgngYmd") or "",
                        "인증종료": p.get("certPrdEndYmd") or "",
                        "숙박": p.get("ldgYn") or "",
                        "이동": p.get("mvmnYn") or "",
                        "고위험": p.get("highRiskYn") or "",
                        "주활동장소": p.get("mainActvPlcNm") or "",
                    })
                st.dataframe(pd.DataFrame(drows), use_container_width=True, hide_index=True)

        if api_key and st.button("조회 결과 AI 요약", key="btn_fac_ai"):
            touch_session()
            sample = pd.DataFrame(merged).drop(columns=["지도"], errors="ignore").head(30).to_csv(index=False)
            try:
                summary = get_ai_response(
                    provider,
                    api_key,
                    "공공기관 데이터 요약. 사실만 간결히.",
                    "다음 수련시설 안전평가·인증현황 결과를 3~6줄로 요약하세요. "
                    "등급·매칭·프로그램 수 특징만. 추측 금지.\n\n" + sample,
                )
                st.session_state["fac_ai_summary"] = summary
            except Exception as e:
                st.error(f"요약 오류: {e}")

        if st.session_state.get("fac_ai_summary"):
            result_card("AI 요약", st.session_state["fac_ai_summary"])

elif tool == "🧭 청소년 통합 지원 서비스":
    page_header("🧭 청소년 통합 지원 서비스", "진로·체험 · 복지혜택 · 지역 활동 지도")
    # 접속지(IP) 기반 시·도 추정 — 대화에서 지역을 먼저 묻지 않음
    if "career_ip_region" not in st.session_state:
        st.session_state.career_ip_region = detect_region_from_ip()
    if st.session_state.career_ip_region.get("region"):
        if not (st.session_state.get("career_profile") or {}).get("region"):
            p = dict(st.session_state.get("career_profile") or {})
            p["region"] = st.session_state.career_ip_region["region"]
            st.session_state.career_profile = p


    if "career_msgs" not in st.session_state:
        st.session_state.career_msgs = []
    if "career_profile" not in st.session_state:
        st.session_state.career_profile = {}
    if "career_guide" not in st.session_state:
        st.session_state.career_guide = ""
    if "career_ggomgil" not in st.session_state:
        st.session_state.career_ggomgil = []
    if "career_youth" not in st.session_state:
        st.session_state.career_youth = []
        st.session_state.career_welfare = None
        st.session_state.career_welfare_note = ""
        st.session_state.career_welfare_meta = {}

    # 관리자: 청소년활동 일괄 캐시 + 진로솔루션 스크래핑
    if st.session_state.get("username") == ADMIN_USER:
        with st.expander("관리자 · 청소년활동 API 일괄 수집(캐시)"):
            y_all, y_meta = load_youth_file_cache()
            st.caption(
                f"캐시 상태: {(y_meta or {}).get('status', '-')} · "
                f"건수: {(y_meta or {}).get('count', 0)} · "
                f"저장: {(y_meta or {}).get('saved_at', '-')}"
            )
            ykey = get_youth_api_key()
            if not ykey:
                st.warning("YOUTH_ACTIV_KEY 또는 DATA_GO_KR_KEY 가 Secrets에 필요합니다.")
            st.info(
                "① **이어받기**: 중단해도 중간 저장분에서 재개 가능. "
                "② **운영·예정만**: API에 sdate~edate를 넣어 **호출 단계**에서 건수를 줄입니다. "
                "페이지당 100~1000 권장."
            )
            mode = st.radio(
                "기간 필터 (API 요청)",
                ["운영·예정만 (권장)", "전체 기간"],
                horizontal=True,
                key="youth_collect_mode",
            )
            from datetime import timedelta as _td
            _today = datetime.now().date()
            if mode.startswith("운영"):
                _ds = (_today - _td(days=7)).strftime("%Y%m%d")
                _de = (_today + _td(days=365)).strftime("%Y%m%d")
            else:
                _ds, _de = "", ""
            d1, d2 = st.columns(2)
            with d1:
                sdate_in = st.text_input("sdate 활동시작일(YYYYMMDD)", value=_ds, key="youth_sdate")
            with d2:
                edate_in = st.text_input("edate 활동종료일(YYYYMMDD)", value=_de, key="youth_edate")
            st.caption("비우면 기간 조건 없이 호출합니다. API가 기간 파라미터를 지원해야 효과가 있습니다.")
            c_a, c_b = st.columns(2)
            with c_a:
                max_pages = st.number_input(
                    "최대 페이지",
                    min_value=1,
                    max_value=5000,
                    value=1500,
                    step=50,
                    key="youth_max_pages",
                    help="12만건 ÷ 페이지당건수 보다 여유 있게",
                )
            with c_b:
                rows_pp = st.number_input(
                    "페이지당 건수",
                    min_value=10,
                    max_value=1000,
                    value=100,
                    step=10,
                    key="youth_rows",
                    help="공공 API는 보통 100~1000 제한",
                )
            st.caption(
                f"예상 최대 수집량: 약 {int(max_pages) * int(rows_pp):,} 건 "
                f"(실제는 API totalCount에서 멈춤)"
            )
            last_page = int((y_meta or {}).get("last_page") or 0)
            is_complete = bool((y_meta or {}).get("complete"))
            cached_n = int((y_meta or {}).get("count") or 0)
            if cached_n:
                st.warning(
                    f"캐시 **{cached_n:,}건** · 마지막 페이지 **{last_page}** · "
                    f"상태: {'완료 표시' if is_complete else '미완료(이어받기 가능)'}"
                )
                if is_complete and cached_n < 20000:
                    st.error(
                        "완료로 표시되었지만 건수가 적습니다. "
                        "이전 수집이 **봉사(Vol) API** 등으로 조기 종료됐을 수 있습니다. "
                        "아래에서 **강제 이어받기** 또는 **처음부터(검색 API)** 를 사용하세요."
                    )
            resume = st.checkbox(
                "기존 캐시에서 이어받기",
                value=bool(cached_n and (not is_complete or cached_n < 20000)),
                key="youth_resume",
            )
            force_resume = st.checkbox(
                "강제 이어받기 (complete=true 여도 다음 페이지부터 계속)",
                value=False,
                key="youth_force_resume",
            )
            by_sido = st.checkbox(
                "시도별 분할 수집 (권장 · 전국 단일 한도 ~4천건 우회)",
                value=True,
                key="youth_by_sido",
            )
            endpoint_mode = st.selectbox(
                "수집 API",
                [
                    "search_only (프로그램 검색 API · 권장)",
                    "auto (검색→신고→봉사 순)",
                    "singo_only (신고 프로그램)",
                    "vol_only (봉사 · 비권장)",
                ],
                index=0,
                key="youth_endpoint_mode",
            )
            ck_every = st.number_input(
                "중간 저장 주기(페이지)",
                min_value=1,
                max_value=50,
                value=5,
                key="youth_ckpt",
                help="N페이지마다 파일로 저장 → 중단 후에도 이어받기 가능",
            )
            if st.button("청소년활동 전체 수집·캐시 저장", type="primary", key="btn_youth_scrape"):
                if not ykey:
                    st.error("API 키가 없습니다.")
                else:
                    prog = st.progress(0)
                    status = st.empty()
                    def _cb(page, n, total, extra=None):
                        extra = extra or {}
                        status.write(
                            f"[{extra.get('sido','')}] p{extra.get('page', page)} · "
                            f"누적 {n:,}건 · 신규 {extra.get('new', '-')} "
                            + (f"/ 약 {total:,} " if total else "")
                            + (f"· {(extra.get('url') or '')[-40:]}")
                        )
                        # 시도 수 기준 대략 진행률
                        prog.progress(min(0.99, n / 120000.0) if n else 0.01)
                    existing = []
                    start_page = 1
                    do_resume = resume or force_resume
                    if do_resume and y_all:
                        existing = list(y_all)
                        start_page = max(1, last_page + 1)
                        status.write(f"이어받기: {len(existing):,}건 보유 → {start_page}페이지부터")
                    emode = "search_only"
                    if "auto" in (endpoint_mode or ""):
                        emode = "auto"
                    elif "singo" in (endpoint_mode or ""):
                        emode = "singo_only"
                    elif "vol" in (endpoint_mode or ""):
                        emode = "vol_only"
                    with st.spinner("청소년활동 API 수집 중… (중단해도 중간 저장분 유지)"):
                        items, meta = fetch_youth_programs_all(
                            ykey,
                            num_of_rows=int(rows_pp),
                            max_pages=int(max_pages),
                            progress_cb=_cb,
                            start_page=start_page,
                            existing_items=existing,
                            checkpoint_every=int(ck_every),
                            sdate=(sdate_in or "").strip(),
                            edate=(edate_in or "").strip(),
                            endpoint_mode=emode,
                            by_sido=bool(st.session_state.get("youth_by_sido", True)),
                        )
                    if items:
                        info = save_youth_file_cache(
                            items,
                            extra_meta={
                                "last_page": meta.get("last_page") or meta.get("pages"),
                                "num_of_rows": int(rows_pp),
                                "totalCount": meta.get("totalCount"),
                                "complete": bool(meta.get("complete")),
                            },
                        )
                        done = "완료" if meta.get("complete") else "미완료(이어받기 가능)"
                        st.success(
                            f"저장: {info.get('count'):,}건 · 마지막페이지 {info.get('last_page')} · {done}"
                        )
                        st.json({k: meta.get(k) for k in ("collected", "pages", "totalCount", "complete", "last_page", "url", "resultCode") if k in (meta or {}) or True})
                    else:
                        st.error("수집 결과가 비었습니다. API 키·엔드포인트를 확인하세요.")
                        st.write(meta)
            if y_all:
                st.markdown("#### 캐시 로컬 저장")
                st.caption("아래 버튼으로 PC에 다운로드할 수 있습니다.")
                # gzip json 원본
                try:
                    raw_gz = YOUTH_CACHE_FILE.read_bytes() if YOUTH_CACHE_FILE.exists() else None
                except Exception:
                    raw_gz = None
                cdl1, cdl2, cdl3 = st.columns(3)
                with cdl1:
                    if raw_gz:
                        st.download_button(
                            "📦 JSON.GZ 원본",
                            data=raw_gz,
                            file_name="youth_programs.json.gz",
                            mime="application/gzip",
                            key="dl_youth_gz",
                        )
                with cdl2:
                    import pandas as pd
                    try:
                        df_y = pd.json_normalize(y_all)
                        csv_y = df_y.to_csv(index=False).encode("utf-8-sig")
                        st.download_button(
                            "📥 CSV 다운로드",
                            data=csv_y,
                            file_name=f"youth_programs_{len(y_all)}.csv",
                            mime="text/csv",
                            key="dl_youth_csv",
                        )
                    except Exception as e:
                        st.caption(f"CSV 변환 실패: {e}")
                with cdl3:
                    try:
                        js = json.dumps(
                            {"count": len(y_all), "items": y_all},
                            ensure_ascii=False,
                        ).encode("utf-8")
                        st.download_button(
                            "📄 JSON 다운로드",
                            data=js,
                            file_name=f"youth_programs_{len(y_all)}.json",
                            mime="application/json",
                            key="dl_youth_json",
                        )
                    except Exception as e:
                        st.caption(f"JSON 변환 실패: {e}")
            if y_all and st.button("청소년활동 캐시 삭제", key="btn_youth_cache_del"):
                try:
                    YOUTH_CACHE_FILE.unlink(missing_ok=True)
                    st.success("캐시 삭제됨")
                    st.rerun()
                except Exception as e:
                    st.error(str(e))
        with st.expander("관리자 · 진로솔루션 자료 스크래핑"):
            st.caption("진로상담 > 진로솔루션 게시글 목록·상세·PDF 요약을 로컬 캐시에 저장합니다.")
            jmeta = {}
            if JINSOL_META.exists():
                try:
                    jmeta = json.loads(JINSOL_META.read_text(encoding="utf-8"))
                except Exception:
                    jmeta = {}
            if jmeta:
                st.write(f"캐시: {jmeta.get('count', 0)}건 · 텍스트 {jmeta.get('with_text', 0)}건 · {jmeta.get('built_at', '-')}")
            if st.button("자료 스크래핑", key="btn_jinsol_scrape"):
                prog = st.empty()
                try:
                    def _cb(msg):
                        prog.info(msg)
                    summary = scrape_jinsol_solutions(progress_cb=_cb)
                    prog.success(
                        f"완료: {summary.get('count', 0)}건 "
                        f"(본문추출 {summary.get('with_text', 0)}건)"
                    )
                except Exception as e:
                    prog.error(f"스크래핑 오류: {e}")

    c_reset, _ = st.columns([1, 3])
    with c_reset:
        if st.button("대화 초기화", key="btn_career_reset"):
            for k in list(st.session_state.keys()):
                if str(k).startswith("career_"):
                    del st.session_state[k]
            st.session_state.career_clarify_count = 0
            st.rerun()

    # 첫 인사 — 가볍게, 세부 취조 금지
    if not st.session_state.career_msgs:
        st.session_state.career_msgs.append({
            "role": "assistant",
            "content": (
                "안녕! 나는 진로맵 길잡이야.\n"
                "초·중·고 중 어디에 가까운지, 요즘 관심 있는 것(예: 요리, 로봇, 그림)만 편하게 말해줘.\n"
                "지역을 알면 근처 체험도 같이 찾아줄게. 없어도 괜찮아!"
            ),
        })

    for msg in st.session_state.career_msgs:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    user_turns = sum(1 for m in st.session_state.career_msgs if m["role"] == "user")
    if "career_clarify_count" not in st.session_state:
        st.session_state.career_clarify_count = 0

    prompt = st.chat_input("관심이나 학교·지역을 적어 주세요")
    if prompt:
        touch_session()
        st.session_state.career_msgs.append({"role": "user", "content": prompt})
        user_turns += 1

        user_text = " ".join(
            m["content"] for m in st.session_state.career_msgs if m["role"] == "user"
        )
        has_school = any(
            k in user_text for k in ("초등", "중학", "고등", "초 ", "중 ", "고 ", "학년", "초등학교", "중학교", "고등학교")
        )

        # ---------- 모호 답변·학교명 추가 질의 (최대 2회) ----------
        # 관심: 동음이의/짧은 명사/넓은 주제 · 학교: 동일 교명 다지역 가능 시 지역 확인
        need_clarify = False
        clarify_q = ""
        max_clarify = 2
        can_clarify = (
            st.session_state.career_clarify_count < max_clarify
            and user_turns <= 5
            and not st.session_state.get("career_guide")
        )

        import re as _re_c
        last_u = prompt.strip()
        # 규칙 1) 짧은 관심 명사·동음이의 후보
        AMBIG_WORDS = {
            "새", "차", "공", "별", "빛", "눈", "장", "배", "말", "밤", "달", "강", "산",
            "과학", "예술", "게임", "운동", "음악", "그림", "동물", "컴퓨터", "로봇",
        }
        region_tokens = (
            "서울", "부산", "대구", "인천", "광주", "대전", "울산", "세종",
            "경기", "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주",
            "시", "군", "구", "특별", "광역",
        )
        has_region = any(r in user_text for r in region_tokens)
        school_name_hit = bool(
            _re_c.search(
                r"([가-힣]{2,12}(초등학교|중학교|고등학교|초등|중학|고등))",
                user_text,
            )
        )
        short_interest = (
            len(last_u) <= 6
            or any(w in last_u for w in AMBIG_WORDS)
            or (
                len(last_u.split()) <= 2
                and any(w in last_u for w in AMBIG_WORDS)
            )
        )

        if can_clarify:
            # 규칙 기반 우선
            if school_name_hit and not has_region and st.session_state.career_clarify_count == 0:
                need_clarify = True
                m_sch = _re_c.search(
                    r"([가-힣]{2,12}(초등학교|중학교|고등학교|초등|중학|고등))",
                    user_text,
                )
                sch_nm = m_sch.group(1) if m_sch else "학교"
                clarify_q = (
                    f"‘{sch_nm}’이(가) 여러 지역에 있을 수 있어.\n"
                    "어느 시·도(또는 시·군·구) 쪽인지 알려주면 근처 체험을 더 정확히 찾을게!"
                )
            elif short_interest and st.session_state.career_clarify_count < max_clarify:
                # 짧은 관심어는 AI 또는 고정 질문으로 한 번 더
                need_clarify = True
                if api_key:
                    try:
                        hist = "\n".join(
                            f"{m['role']}: {m['content']}"
                            for m in st.session_state.career_msgs[-8:]
                        )
                        raw_c = get_ai_response(
                            provider,
                            api_key,
                            (
                                "진로 대화 판별기. JSON만 출력.\n"
                                "need_clarify=true 인 경우:\n"
                                "1) 관심어가 짧거나 동음이의(예: 새=조류/새로움, 차=자동차/茶)\n"
                                "2) 범위가 너무 넓음(과학, 예술, 게임만)\n"
                                "3) 학교명만 있고 지역이 없음\n"
                                "이미 충분히 구체적이면 need_clarify=false.\n"
                                "question은 반말 1~2문장, 2~3개 선택지를 가볍게 제시.\n"
                                "개인정보·부모·일정은 묻지 말 것."
                            ),
                            (
                                f"대화:\n{hist}\n\n"
                                'JSON: {"need_clarify":true/false,"reason":"",'
                                '"question":"추가질문"}'
                            ),
                        )
                        mjson = _re_c.search(r"\{[\s\S]*\}", raw_c or "")
                        if mjson:
                            cobj = json.loads(mjson.group(0))
                            # 규칙으로 이미 true면 AI가 false해도 유지하되 질문만 AI 것으로
                            if cobj.get("question"):
                                clarify_q = str(cobj.get("question")).strip()
                            if cobj.get("need_clarify") is False and not short_interest:
                                need_clarify = False
                    except Exception:
                        pass
                if not clarify_q:
                    clarify_q = (
                        f"‘{last_u}’에 관심 있구나! 어떤 쪽에 더 가까워?\n"
                        "예) 직접 관찰·키우기 / 그리거나 만들기 / 관련 직업·과학 알아보기 "
                        "중 고르거나, 다른 뜻이면 짧게 알려줘."
                    )
            elif can_clarify and api_key:
                # 그 외에도 AI가 모호하다고 보면 1회 질문
                try:
                    hist = "\n".join(
                        f"{m['role']}: {m['content']}"
                        for m in st.session_state.career_msgs[-8:]
                    )
                    raw_c = get_ai_response(
                        provider,
                        api_key,
                        (
                            "진로 대화 판별기. JSON만.\n"
                            "관심·학교·지역 정보가 추천하기에 충분한지 판단.\n"
                            "모호/동음이의/학교만 있고 지역 없음 → need_clarify=true.\n"
                            "충분하면 false. 질문은 반말 짧은 선택형."
                        ),
                        (
                            f"대화:\n{hist}\n\n"
                            'JSON: {"need_clarify":true/false,"question":"..."}'
                        ),
                    )
                    mjson = _re_c.search(r"\{[\s\S]*\}", raw_c or "")
                    if mjson:
                        cobj = json.loads(mjson.group(0))
                        need_clarify = bool(cobj.get("need_clarify"))
                        clarify_q = str(cobj.get("question") or "").strip()
                except Exception:
                    need_clarify = False

        if need_clarify and clarify_q:
            st.session_state.career_clarify_count = int(st.session_state.career_clarify_count) + 1
            st.session_state.career_msgs.append({"role": "assistant", "content": clarify_q})
            st.rerun()

        # 추가 질의가 필요 없거나 한도 도달 → 본 분석 파이프라인
        else:
            # ============================================================
            # 파이프라인
            # ① 대화 메타 수집 → ② 조회 → ③ AI 분석
            # → ④ 분석 키워드로 꿈길·청소년활동 추천
            # ============================================================
            import re as _re
            hist = "\n".join(f"{m['role']}: {m['content']}" for m in st.session_state.career_msgs)
            profile = {
                "school": "", "region": "", "keywords": [], "goal": "",
                "activity_keywords": [],
            }
            guide = ""
            analysis = {}

            # ① 메타데이터 추출
            if api_key:
                try:
                    raw = get_ai_response(
                        provider,
                        api_key,
                        "진로 메타 추출기. JSON만 출력. 설명 금지.",
                        (
                            "대화에서 아래 JSON만 출력:\n"
                            '{"school":"초|중|고|모름","region":"시도시군구키워드",'
                            '"keywords":["관심키워드최대5"],'
                            '"goal":"한줄목표",'
                            '"activity_keywords":["체험검색용키워드최대8"]}\n'
                            "세부분류(서양/동양 등)가 없어도 관심 분야를 넓게 잡을 것.\n"
                            "예: 요리 → 조리,요리,제과,제빵,셰프,쿠킹,음식,디저트\n\n"
                            + hist
                        ),
                    )
                    mjson = _re.search(r"\{[\s\S]*\}", raw or "")
                    if mjson:
                        profile = {**profile, **json.loads(mjson.group(0))}
                except Exception as e:
                    st.warning(f"메타 추출 참고: {e}")
            else:
                text_all = " ".join(
                    m["content"] for m in st.session_state.career_msgs if m["role"] == "user"
                )
                for tok in ["초등", "초 ", "중학", "중 ", "고등", "고 "]:
                    if tok.strip() in text_all:
                        profile["school"] = "초" if "초" in tok else ("중" if "중" in tok else "고")
                        break
                for r in ["서울", "경기", "인천", "부산", "대구", "광주", "대전", "울산", "세종",
                          "강원", "충북", "충남", "전북", "전남", "경북", "경남", "제주"]:
                    if r in text_all:
                        profile["region"] = r
                        break
                profile["keywords"] = [w for w in text_all.replace(",", " ").split() if len(w) >= 2][:5]
                profile["activity_keywords"] = list(profile["keywords"])

            school = str(profile.get("school") or "")
            region = str(profile.get("region") or "")
            kws = profile.get("keywords") or []
            if isinstance(kws, str):
                kws = [kws]
            act_kws = profile.get("activity_keywords") or kws
            if isinstance(act_kws, str):
                act_kws = [act_kws]
            # 검색용 키워드 합치기 + 관심어 확장 (예: 새 → 조류, 탐조, 동물…)
            search_kws = []
            for k in list(kws) + list(act_kws):
                k = str(k).strip()
                if k and k not in search_kws:
                    search_kws.append(k)
            search_kws = expand_interest_keywords(search_kws)
            kw0 = search_kws[0] if search_kws else ""

            # ② API 수집 (프로그램 목록보다 먼저)
            cn_key = get_career_net_key()
            jobs, jobs_meta, counsel_list, counsel_detail, majors = [], {}, [], {}, []
            if cn_key:
                try:
                    # 관심 키워드마다 직업 검색 (빈 키워드 조회 금지 → 인기직업 오염 방지)
                    seen_job = set()
                    for sk in (search_kws[:8] or []):
                        if not sk:
                            continue
                        if school.startswith("초") or school.startswith("중"):
                            part, jobs_meta = fetch_careernet_junior_jobs(cn_key, keyword=sk)
                            if not part:
                                part, jobs_meta = fetch_careernet_jobs(cn_key, keyword=sk)
                        else:
                            part, jobs_meta = fetch_careernet_jobs(cn_key, keyword=sk)
                            if not part:
                                part, jobs_meta = fetch_careernet_junior_jobs(cn_key, keyword=sk)
                        for j in (part or []):
                            nm = j.get("job_nm") or j.get("JOB_NM") or j.get("jobNm") or str(j)
                            if nm not in seen_job:
                                seen_job.add(nm)
                                jobs.append(j)
                        if len(jobs) >= 20:
                            break
                    counsel_list, _ = fetch_careernet_counsel(cn_key)
                    if counsel_list and search_kws:
                        scored = []
                        for c in counsel_list:
                            memo = str(c.get("memo") or c.get("question") or "")
                            score = sum(1 for k in search_kws if k and k in memo)
                            scored.append((score, c))
                        scored.sort(key=lambda x: -x[0])
                        counsel_list = [c for s, c in scored[:8] if s > 0] or counsel_list[:5]
                    else:
                        counsel_list = (counsel_list or [])[:5]
                    # 상세 Q&A 최대 3건
                    counsel_details = []
                    for c in counsel_list[:5]:
                        code = c.get("code") or c.get("con_cd") or ""
                        if not code:
                            continue
                        det, _ = fetch_careernet_counsel_detail(cn_key, code)
                        if det and (det.get("question") or det.get("answer") or det.get("memo")):
                            counsel_details.append(det)
                        if len(counsel_details) >= 3:
                            break
                    if counsel_details:
                        counsel_detail = counsel_details[0]
                    st.session_state.career_counsel_details = counsel_details
                    # 학과·학교·진로교육: 전체 캐시 구축 후 키워드/AI 매칭
                    # 기존 빈 캐시면 강제 재수집
                    force_rebuild = False
                    try:
                        if not CN_MAJOR_CACHE.exists() or len(load_cn_cache("majors")) < 10:
                            force_rebuild = True
                    except Exception:
                        force_rebuild = True
                    cache_sum = build_careernet_cache(cn_key, force=force_rebuild)
                    st.session_state.career_cn_cache_meta = cache_sum

                    cand_majors = pick_from_cn_cache("majors", search_kws, limit=25)
                    if len(cand_majors) < 5:
                        cand_majors = load_cn_cache("majors")[:40]
                    majors = ai_pick_from_cache(
                        provider, api_key, "학과", search_kws, cand_majors, limit=10
                    ) if api_key else cand_majors[:10]

                    # 캐시 매칭 실패 시 실시간 searchTitle 조회 (자동차 등)
                    if len(majors) < 3 and search_kws:
                        live = []
                        for gubun in ("univ_list", "대학교", "전문대학"):
                            for sk in search_kws[:5]:
                                try:
                                    items, _ = fetch_careernet_majors(
                                        cn_key, keyword=sk, gubun=gubun, page=1, per_page=30
                                    )
                                    live.extend(items or [])
                                except Exception:
                                    pass
                        if live:
                            # 캐시에 병합 저장
                            try:
                                existing = load_cn_cache("majors")
                                bucket = {}
                                _merge_cn_items(bucket, existing, ["major", "facilName", "mClass", "name"])
                                _merge_cn_items(bucket, live, ["major", "facilName", "mClass", "name"])
                                CN_MAJOR_CACHE.write_text(
                                    json.dumps(list(bucket.values()), ensure_ascii=False),
                                    encoding="utf-8",
                                )
                            except Exception:
                                pass
                            majors = ai_pick_from_cache(
                                provider, api_key, "학과", search_kws, live, limit=10
                            ) if api_key else live[:10]

                    major_details = []
                    for m in (majors or [])[:5]:
                        seq = str(
                            m.get("majorSeq") or m.get("seq") or m.get("mjrSeq")
                            or m.get("major_seq") or ""
                        )
                        mn = (
                            m.get("major") or m.get("majorNm") or m.get("mjrNm")
                            or m.get("name") or ""
                        )
                        det, _ = fetch_careernet_major_view(
                            cn_key, major_seq=seq, major_nm=str(mn)
                        )
                        if det:
                            major_details.append(det)
                    st.session_state.career_major_details = major_details
                    st.session_state.career_majors = majors[:10]

                    cand_schools = pick_from_cn_cache(
                        "schools", search_kws + ([region] if region else []), limit=25
                    )
                    if len(cand_schools) < 5:
                        cand_schools = load_cn_cache("schools")[:40]
                    schools = ai_pick_from_cache(
                        provider, api_key, "학교", search_kws, cand_schools, limit=8
                    ) if api_key else cand_schools[:8]
                    st.session_state.career_schools = schools[:8]

                    # 진로교육자료: 사전 캐시 전체에서 키워드 후보 → AI가 흥미·적성 향상 자료 선별
                    cand_cose = pick_from_cn_cache("cose", search_kws, limit=40)
                    if len(cand_cose) < 10:
                        full_cose = load_cn_cache("cose")
                        # 키워드 약한 매칭으로 후보 확장
                        seen = {_cn_item_key(x, ["dataTitle", "title", "seq"]) for x in cand_cose}
                        for it in full_cose:
                            if not isinstance(it, dict):
                                continue
                            blob = " ".join(str(it.get(k) or "") for k in (
                                "dataTitle", "title", "subject", "author", "activityType", "achieveType"
                            ))
                            if any(str(k) in blob for k in search_kws if k):
                                key = _cn_item_key(it, ["dataTitle", "title", "seq"])
                                if key not in seen:
                                    cand_cose.append(it)
                                    seen.add(key)
                            if len(cand_cose) >= 50:
                                break
                        if len(cand_cose) < 10:
                            cand_cose = (cand_cose + full_cose)[:60]
                    cose_list = ai_pick_from_cache(
                        provider, api_key, "진로교육자료", search_kws, cand_cose, limit=8
                    ) if api_key else cand_cose[:8]
                    st.session_state.career_cose = cose_list[:8]
                except Exception as e:
                    jobs_meta = {"error": str(e)}

            # 관련도 낮은 직업 제거 (쇼핑호스트·가사도우미 등)
            jobs = filter_jobs_by_keywords(jobs, search_kws, min_score=2, limit=8)
            st.session_state.career_jobs = jobs
            st.session_state.career_jobs_meta = jobs_meta
            st.session_state.career_counsel = counsel_list
            st.session_state.career_counsel_detail = counsel_detail
            majors_raw = filter_majors_by_keywords(
                st.session_state.get("career_majors") or majors or [],
                search_kws,
                min_score=2,
                limit=40,
            )
            major_rows_all = normalize_major_rows(majors_raw)
            st.session_state.career_majors_all = major_rows_all
            top_majors = ai_pick_top_majors(
                provider, api_key, search_kws, major_rows_all, limit=5
            ) if api_key else major_rows_all[:5]
            st.session_state.career_majors = top_majors

            def _major_brief(m):
                if not isinstance(m, dict):
                    return {"name": str(m)}
                name = (
                    m.get("major") or m.get("majorNm") or m.get("mjrNm")
                    or m.get("name") or m.get("facilName") or ""
                )
                desc = (
                    m.get("summary") or m.get("majorSummary") or m.get("intro")
                    or m.get("content") or m.get("description") or m.get("l<fim-middle>clsfNm")
                    or ""
                )
                return {
                    "name": name,
                    "desc": str(desc)[:200],
                    "lcls": m.get("lclsfNm") or m.get("large") or "",
                    "mcls": m.get("mclsfNm") or m.get("middle") or "",
                }

            major_details = st.session_state.get("career_major_details") or []
            schools = st.session_state.get("career_schools") or []
            cose_list = st.session_state.get("career_cose") or []

            def _school_brief(s):
                return {
                    "name": s.get("schoolNm") or s.get("schulNm") or s.get("name") or s.get("schoolName") or "",
                    "type": s.get("schoolType") or s.get("schulKndScNm") or s.get("type") or "",
                    "region": s.get("region") or s.get("adres") or s.get("areaNm") or s.get("ctpvNm") or "",
                    "desc": str(s.get("schoolInfo") or s.get("intro") or s.get("content") or "")[:150],
                }

            def _major_detail_brief(d):
                return {
                    "name": d.get("major") or d.get("majorNm") or d.get("mjrNm") or d.get("name") or "",
                    "summary": str(
                        d.get("summary") or d.get("majorSummary") or d.get("intro")
                        or d.get("content") or d.get("description") or ""
                    )[:300],
                    "career": str(d.get("career") or d.get("job") or d.get("relateJob") or "")[:150],
                }

            def _cose_brief(c):
                return {
                    "title": c.get("title") or c.get("subject") or c.get("name") or "",
                    "summary": str(c.get("summary") or c.get("content") or c.get("intro") or "")[:150],
                }

            cn_brief = {
                "jobs": [
                    {
                        "name": j.get("job_nm") or j.get("JOB_NM") or j.get("jobNm") or "",
                        "work": (j.get("work") or j.get("JOB_SUMMARY") or "")[:150],
                        "top": j.get("top_nm") or j.get("aptit_name") or "",
                    }
                    for j in (jobs or [])[:6]
                ],
                "counsel": [
                    {"q": c.get("memo") or c.get("question") or ""}
                    for c in (counsel_list or [])[:4]
                ],
                "counsel_detail": {
                    "q": (counsel_detail or {}).get("question") or "",
                    "a": ((counsel_detail or {}).get("answer") or "")[:500],
                },
                "majors": [_major_brief(m) for m in (majors or [])[:8]],
                "major_details": [_major_detail_brief(d) for d in major_details[:5]],
                "schools": [_school_brief(s) for s in schools[:6]],
                "career_edu_materials": [_cose_brief(c) for c in cose_list[:5]],
            }

            # 진로솔루션 캐시 참고
            jinsol_hits = pick_jinsol_for_keywords(search_kws, limit=4)
            st.session_state.career_jinsol = jinsol_hits
            cn_brief["jinsol"] = [
                {
                    "job": j.get("job_nm") or "",
                    "subject": j.get("subject") or "",
                    "year": j.get("year") or "",
                    "excerpt": (j.get("text") or "")[:900],
                    "url": j.get("url") or "",
                }
                for j in jinsol_hits
            ]

            # ③ AI 분석 + 진학 조언 (학과·진로솔루션 반영)
            admission_advice = ""
            if api_key:
                try:
                    raw_an = get_ai_response(
                        provider,
                        api_key,
                        (
                            "진로·진학 분석가. 초등·중학생도 이해하게. 개조식. "
                            "관심 분야를 넓게 안내. 제공된 직업·학과·진로솔루션 사실만 사용하고 없는 학과명은 지어내지 말 것. "
                            "상담사례는 참고만 하고 인용 나열 금지. "
                            "jinsol(진로솔루션) 내용이 있으면 직업 소개·준비 팁을 더 구체적으로 풀어 쓸 것."
                        ),
                        (
                            f"학생메타: {json.dumps(profile, ensure_ascii=False)}\n"
                            f"수집데이터: {json.dumps(cn_brief, ensure_ascii=False)}\n\n"
                            "출력 형식:\n"
                            "## 관심 분야 안내\n"
                            "(5~8줄)\n\n"
                            "## 관련 직업\n"
                            "(2~4개, 제공 데이터 기준. 진로솔루션 내용이 있으면 반영)\n\n"
                            "## 진로솔루션 상세 안내\n"
                            "(jinsol 데이터가 있으면 직업별 3~8줄 구체 안내. 없으면 이 제목 생략)\n\n"
                            "## 관련 학과와 진학 조언\n"
                            "(학과 목록·학과상세·학교정보를 바탕으로 초·중·고에 맞는 진학 조언 10~18줄. "
                            "초등: 흥미·기초체험, 중등: 과목·동아리, 고등: 학과·학교유형. "
                            "제공된 학교·학과명만 사용. 없는 대학/학과 창작 금지.)\n\n"
                            "## 학교·진로교육 참고\n"
                            "(학교정보·진로교육자료가 있으면 2~5줄 요약. 없으면 생략)\n\n"
                            "## 교육·체험·활동 팁\n"
                            "(4~8줄)\n\n"
                            "## 추천 방향\n"
                            "(2~3가지)\n\n"
                            "```json\n"
                            '{"search_keywords":["체험검색키워드6~8개"],'
                            '"job_names":["추천직업"],'
                            '"major_names":["추천학과"],"reason":"한줄"}\n'
                            "```"
                        ),
                    )
                    guide = raw_an or ""
                    mjson = _re.search(r"```json\s*(\{[\s\S]*?\})\s*```", guide)
                    if not mjson:
                        mjson = _re.search(r"(\{[^{}]*search_keywords[^{}]*\})", guide)
                    if mjson:
                        try:
                            analysis = json.loads(mjson.group(1) if mjson.lastindex else mjson.group(0))
                        except Exception:
                            analysis = {}
                    guide = _re.sub(r"```json[\s\S]*?```", "", guide).strip()
                    guide = _re.sub(
                        r"\{[^{}]*\"search_keywords\"[^{}]*\}", "", guide
                    ).strip()
                    # 진학 조언 섹션 분리 저장
                    m_adm = _re.search(
                        r"##\s*관련 학과와 진학 조언\s*([\s\S]*?)(?=\n##\s|\Z)", guide
                    )
                    if m_adm:
                        admission_advice = m_adm.group(1).strip()
                except Exception as e:
                    guide = f"AI 분석 오류: {e}"
            else:
                guide = "사이드바에 AI API 키를 입력하면 분석을 생성합니다."
            st.session_state.career_admission = admission_advice

            # 분석에서 나온 검색어 우선
            if isinstance(analysis.get("search_keywords"), list):
                for k in analysis["search_keywords"]:
                    k = str(k).strip()
                    if k and k not in search_kws:
                        search_kws.append(k)
            if isinstance(analysis.get("job_names"), list):
                for k in analysis["job_names"]:
                    k = str(k).strip()
                    if k and k not in search_kws:
                        search_kws.append(k)

            st.session_state.career_profile = profile
            st.session_state.career_guide = guide
            st.session_state.career_analysis = analysis
            st.session_state.career_search_kws = search_kws

            # 지역 키워드 완화: 고양시 → 고양, 경기
            region_candidates = []
            if region:
                region_candidates.append(region)
                if "고양" in region:
                    region_candidates.extend(["고양", "경기"])
                if region.endswith("시") and len(region) > 2:
                    region_candidates.append(region[:-1])

            # ④ 분석 키워드 → 꿈길 · 청소년활동 (관련 키워드 우선, 무관 목록 채우지 않음)
            hits = []
            try:
                # 1) 지역 + 키워드
                for reg in region_candidates or [""]:
                    hits = search_ggomgil(
                        school_level=school, region_kw=reg, keywords=search_kws, top_n=15
                    )
                    if hits:
                        break
                # 2) 키워드만 (전국)
                if not hits and search_kws:
                    hits = search_ggomgil(
                        school_level=school, region_kw="", keywords=search_kws, top_n=15
                    )
                # 3) 키워드 하나씩
                if not hits and search_kws:
                    for sk in search_kws[:6]:
                        hits = search_ggomgil(
                            school_level=school, region_kw="", keywords=[sk], top_n=15
                        )
                        if hits:
                            break
                # 지역만으로 아무 체험이나 채우지 않음 (무관 결과 방지)
            except Exception as e:
                st.session_state.career_ggomgil_err = str(e)
            if hits and api_key:
                try:
                    hits = ai_filter_ggomgil(provider, api_key, search_kws, hits, limit=10)
                except Exception:
                    hits = hits[:10]
            st.session_state.career_ggomgil = hits or []

            youth_items = []
            try:
                y_all, y_meta = load_youth_file_cache()
                if not y_all:
                    st.session_state.career_youth_meta = {
                        "error": "청소년활동 캐시가 없습니다. 관리자 메뉴에서 API 일괄 수집을 실행하세요.",
                        **(y_meta or {}),
                    }
                else:
                    cand = filter_youth_cache_by_keywords(
                        y_all,
                        search_kws,
                        region=region or "",
                        school=school or "",
                        limit=60,
                    )
                    # 후보가 너무 적으면 지역만이라도 재검색
                    if len(cand) < 5 and (region or school):
                        more = filter_youth_cache_by_keywords(
                            y_all,
                            keywords=[],
                            region=region or "",
                            school=school or "",
                            limit=40,
                        )
                        # 합치기
                        seen_k = set()
                        merged = []
                        for row in cand + more:
                            k = _youth_item_key(row) if isinstance(row, dict) else str(id(row))
                            if k in seen_k:
                                continue
                            seen_k.add(k)
                            merged.append(row)
                        cand = merged[:60]
                    if not cand:
                        # 최후: 지역 문자열 포함 항목
                        if region:
                            cand = [
                                it for it in (y_all or [])
                                if isinstance(it, dict)
                                and region[:2] in json.dumps(it, ensure_ascii=False)
                            ][:40]
                    if api_key and cand:
                        youth_items = ai_pick_youth_from_cache(
                            provider,
                            api_key,
                            search_kws,
                            region or "",
                            cand,
                            limit=12,
                            school=school or "",
                        )
                    else:
                        youth_items = cand[:12]
                    # 지역 강제 탈락 없음 — 점수 순 추천 유지
                    st.session_state.career_youth_meta = {
                        "cache_count": len(y_all),
                        "cache_status": (y_meta or {}).get("status"),
                        "candidates": len(cand) if y_all else 0,
                        "picked": len(youth_items),
                        "saved_at": (y_meta or {}).get("saved_at"),
                    }
            except Exception as e:
                st.session_state.career_youth_meta = {"error": str(e)}
            st.session_state.career_youth = youth_items or []
            # 온통청년 정책 API 연동은 현재 제외

            summary_line = (
                f"대화를 바탕으로 정리했어.\n"
                f"- 교급: **{school or '미상'}** · 지역: **{region or '미상'}**\n"
                f"- 관심: **{', '.join(kws) if kws else '미상'}**\n"
                f"- 검색키워드: **{', '.join(search_kws[:6]) if search_kws else '-'}**\n\n"
                f"직업정보·체험·청소년활동을 안내할게."
            )
            st.session_state.career_msgs.append({"role": "assistant", "content": summary_line})
            st.rerun()


    # 결과 영역 (계층 목차 + 카드 박스)
    if (
        st.session_state.get("career_guide")
        or st.session_state.get("career_ggomgil")
        or st.session_state.get("career_jobs")
        or st.session_state.get("career_counsel")
        or st.session_state.get("career_youth")
        or st.session_state.get("career_jinsol")
        or st.session_state.get("career_majors")
    ):
        st.markdown("---")
        roman = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII"]
        sec = 0

        # I. 진로 분석
        if st.session_state.get("career_guide"):
            cm_h1("🧭", "", "진로 분석")
            sec += 1
            st.markdown('<div class="cm-analysis-box cm-indent">', unsafe_allow_html=True)
            # 하위 ## 제목이 메인보다 커지지 않도록 HTML로 완만히 변환
            import re as _re_g
            import html as _html_g
            g = st.session_state["career_guide"]
            parts = []
            for line in (g or "").splitlines():
                if line.startswith("### "):
                    parts.append(f'<div class="cm-h3">· {_html_g.escape(line[4:].strip())}</div>')
                elif line.startswith("## "):
                    parts.append(f'<div class="cm-h2">{_html_g.escape(line[3:].strip())}</div>')
                elif line.startswith("# "):
                    parts.append(f'<div class="cm-h2">{_html_g.escape(line[2:].strip())}</div>')
                elif line.strip().startswith("- ") or line.strip().startswith("* "):
                    parts.append(
                        f'<div class="cm-body cm-indent">'
                        f'<span class="cm-dot">•</span>{_html_g.escape(line.strip()[2:])}</div>'
                    )
                elif line.strip():
                    parts.append(f'<div class="cm-body">{_html_g.escape(line)}</div>')
            st.markdown("".join(parts) if parts else f'<div class="cm-body">{_html_g.escape(g or "")}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        # 메타 한 줄
        prof = st.session_state.get("career_profile") or {}
        sk = st.session_state.get("career_search_kws") or []
        st.markdown(
            f'<div class="cm-meta cm-indent">교급 {prof.get("school") or "-"} · '
            f'지역 {prof.get("region") or "-"} · '
            f'관심 {", ".join(prof.get("keywords") or []) or "-"}</div>',
            unsafe_allow_html=True,
        )

        # II. 관련 직업
        cm_h1("💼", "", "관련 직업")
        sec += 1
        c_jobs = st.session_state.get("career_jobs") or []
        if not c_jobs:
            cm_empty("관련 직업 검색 결과가 없습니다.")
        else:
            for i, j in enumerate(c_jobs[:8], 1):
                nm = j.get("job_nm") or j.get("JOB_NM") or j.get("jobNm") or "(직업)"
                work = j.get("work") or j.get("JOB_SUMMARY") or j.get("job_summary") or ""
                top = j.get("top_nm") or j.get("aptit_name") or ""
                body = (work or "")[:220]
                meta = top if top else ""
                cm_card(f"{i}. {nm}", body=body, meta=meta)

        # III. 관련 학과
        cm_h1("🎓", "", "관련 학과")
        sec += 1
        majors = st.session_state.get("career_majors") or []
        majors_all = st.session_state.get("career_majors_all") or []
        if not majors and not majors_all:
            cm_empty("관련 학과를 찾지 못했습니다.")
        else:
            st.markdown(
                '<div class="cm-meta cm-indent">AI가 관심·유망도를 기준으로 선별한 추천 학과 (최대 5개)</div>',
                unsafe_allow_html=True,
            )
            for i, r in enumerate(majors[:5], 1):
                if not isinstance(r, dict):
                    continue
                title = r.get("소분류(학과)") or r.get("학과명") or "(학과)"
                meta = " · ".join(
                    x for x in [r.get("대분류") or r.get("계열"), r.get("중분류")] if x
                )
                cm_card(f"{i}. {title}", meta=meta)
            if majors_all:
                try:
                    import pandas as pd
                    df_all = pd.DataFrame(majors_all)
                    # 대·중·소 정렬
                    for col in ("대분류", "중분류", "소분류(학과)"):
                        if col not in df_all.columns:
                            df_all[col] = ""
                    df_all = df_all.sort_values(
                        by=["대분류", "중분류", "소분류(학과)"], kind="stable"
                    )
                    csv_bytes = df_all.to_csv(index=False).encode("utf-8-sig")
                    st.download_button(
                        label=f"📥 관련 학과 전체 CSV ({len(df_all)}건)",
                        data=csv_bytes,
                        file_name="related_majors.csv",
                        mime="text/csv",
                        key="dl_majors_csv",
                    )
                except Exception as e:
                    st.caption(f"CSV 준비 오류: {e}")

        # IV. 진학 조언
        if st.session_state.get("career_admission"):
            cm_h1("📌", "", "진학 조언")
            sec += 1
            st.markdown('<div class="cm-card cm-indent">', unsafe_allow_html=True)
            st.markdown(st.session_state["career_admission"])
            st.markdown('</div>', unsafe_allow_html=True)

        # V. 진로솔루션 참고 (정리된 카드)
        jinsol_hits = st.session_state.get("career_jinsol") or []
        if jinsol_hits:
            cm_h1("📘", "", "진로솔루션 참고")
            sec += 1
            for i, j in enumerate(jinsol_hits[:4], 1):
                title = j.get("job_nm") or "자료"
                sub = j.get("subject") or ""
                yr = " · ".join(
                    x for x in [str(j.get("year") or ""), str(j.get("month") or "")] if x
                )
                # 본문에서 잡음 줄이기: 앞부분만, 줄바꿈 정리
                raw = (j.get("text") or "").replace("\r", "\n")
                lines = [ln.strip() for ln in raw.split("\n") if ln.strip()]
                # 헤더성 짧은 줄 제외하고 문장 위주
                clean = []
                for ln in lines:
                    if len(ln) < 8 and any(x in ln for x in ("QnA", "솔루션", "vol", "Vol", "전문가")):
                        continue
                    clean.append(ln)
                    if sum(len(x) for x in clean) > 320:
                        break
                excerpt = " ".join(clean)[:320]
                if excerpt and len(" ".join(clean)) > 320:
                    excerpt += "…"
                link = j.get("url") or ""
                pdf = j.get("downloadfile") or ""
                if pdf and not str(pdf).startswith("http"):
                    pdf = "https://www.career.go.kr" + str(pdf)
                links = []
                if link:
                    links.append(f'<a href="{link}" target="_blank">웹에서 보기</a>')
                if pdf:
                    links.append(f'<a href="{pdf}" target="_blank">PDF</a>')
                meta = " · ".join(x for x in [sub, yr] if x and x != title)
                cm_card(
                    f"{i}. {title}",
                    body=excerpt,
                    meta=meta,
                    links=" · ".join(links),
                )

        # VI. 관련 학교
        schools = st.session_state.get("career_schools") or []
        cm_h1("🏫", "", "관련 학교 정보")
        sec += 1
        if not schools:
            cm_empty("관심·지역에 맞는 학교 정보가 없거나 매칭되지 않았습니다.")
        else:
            for i, s in enumerate(schools[:8], 1):
                if not isinstance(s, dict):
                    continue
                nm = s.get("schoolNm") or s.get("schoolName") or s.get("schulNm") or s.get("name") or ""
                typ = s.get("schoolType") or s.get("schoolGubun") or s.get("type") or ""
                region = s.get("region") or s.get("adres") or ""
                link = s.get("link") or s.get("collegeinfourl") or ""
                links = f'<a href="{link}" target="_blank">학교 정보</a>' if link and str(link).startswith("http") else ""
                cm_card(f"{i}. {nm}", meta=" · ".join(x for x in [typ, region] if x), links=links)

        # VII. 진로교육자료
        cose_list = st.session_state.get("career_cose") or []
        cm_h1("📚", "", "진로교육자료")
        sec += 1
        if not cose_list:
            cm_empty("관련 진로교육자료가 없습니다.")
        else:
            cn_key = get_career_net_key()
            for i, c in enumerate(cose_list[:8], 1):
                if not isinstance(c, dict):
                    continue
                title = (
                    c.get("title") or c.get("dataTitle") or c.get("subject")
                    or c.get("name") or ""
                )
                if not title:
                    continue
                seq = str(c.get("seq") or "").strip()
                author = c.get("author") or ""
                year = c.get("year") or ""
                files = list(c.get("files") or [])
                if not files:
                    files = parse_cose_att_files(c.get("attFile") or c.get("url") or "", seq)
                if cn_key and seq:
                    try:
                        det, _ = fetch_careernet_cose_view(cn_key, seq)
                        if det:
                            title = det.get("title") or title
                            more = det.get("files") or parse_cose_att_files(det.get("attFile") or "", seq)
                            if more:
                                seen = set(files)
                                for u in more:
                                    if u not in seen:
                                        files.append(u)
                                        seen.add(u)
                                files = parse_cose_att_files(",".join(files), seq)
                    except Exception:
                        pass
                body_files = [u for u in files if "no=0" not in u]
                cover = [u for u in files if "no=0" in u]
                main = (body_files or cover or [None])[0]
                links = []
                if main:
                    links.append(f'<a href="{main}" target="_blank">자료 파일</a>')
                if body_files and cover:
                    links.append(f'<a href="{cover[0]}" target="_blank">표지</a>')
                web = cose_search_page_url(title)
                if web:
                    links.append(f'<a href="{web}" target="_blank">자료실</a>')
                cm_card(
                    f"{i}. {title}",
                    meta=" · ".join(x for x in [author, year] if x),
                    links=" · ".join(links),
                )

        # IX. 복지혜택
        cm_h1("🏛️", "", "복지혜택 안내")
        sec += 1
        prof = st.session_state.get("career_profile") or {}
        user_region = (prof.get("region") or "").strip()
        if not user_region:
            ipr = st.session_state.get("career_ip_region") or detect_region_from_ip()
            st.session_state.career_ip_region = ipr
            if ipr.get("region"):
                user_region = ipr["region"]
                prof = dict(prof)
                prof["region"] = user_region
                st.session_state.career_profile = prof
                st.caption(f"접속 위치 추정: {user_region}")
        ctpv, sgg = split_region_ctpv_sgg(user_region)
        life, age_n = map_school_to_life_age(str(prof.get("school") or ""), "")
        interest_str = ", ".join(prof.get("interests") or []) if isinstance(prof.get("interests"), list) else str(prof.get("interests") or prof.get("interest") or "")
        thema_codes = map_interest_to_thema_codes(interest_str)
        wkey = get_welfare_api_key()
        wlist = st.session_state.get("career_welfare")
        if wlist is None and wkey:
            collected = []
            wmeta = {}
            for thema in (thema_codes + [""]):
                part, wmeta = fetch_welfare_list(
                    wkey, ctpv_nm=ctpv, sgg_nm=sgg, life_array=life,
                    age=age_n, search_wrd=(interest_str.split(",")[0].strip() if interest_str else "청소년"),
                    intrs_thema=thema, num_of_rows=20,
                )
                for p in part or []:
                    if isinstance(p, dict):
                        collected.append(p)
                if len(collected) >= 15:
                    break
            seen_s, uniq = set(), []
            for p in collected:
                sid = p.get("servId") or p.get("servNm")
                if sid in seen_s:
                    continue
                seen_s.add(sid)
                uniq.append(p)
            if api_key and uniq:
                picked, note = ai_welfare_eligibility_hint(provider, api_key, prof, uniq, limit=8)
                st.session_state.career_welfare = picked
                st.session_state.career_welfare_note = note
            else:
                st.session_state.career_welfare = uniq[:8]
                st.session_state.career_welfare_note = ""
            st.session_state.career_welfare_meta = wmeta
            wlist = st.session_state.career_welfare
        if not wkey:
            cm_empty("복지 API 키가 없습니다. secrets에 WELFARE_API_KEY 또는 DATA_GO_KR_KEY를 설정하세요.")
        elif not wlist:
            cm_empty(str((st.session_state.get("career_welfare_meta") or {}).get("error") or "복지 조회 결과가 없습니다."))
        else:
            if st.session_state.get("career_welfare_note"):
                st.caption(st.session_state.career_welfare_note)
            import pandas as pd
            rows = []
            for w in wlist:
                rows.append({
                    "서비스명": w.get("servNm") or "",
                    "지역": f"{w.get('ctpvNm') or ''} {w.get('sggNm') or ''}".strip(),
                    "생애주기": w.get("lifeNmArray") or "",
                    "대상": w.get("trgterIndvdlNmArray") or "",
                    "요약": (w.get("servDgst") or "")[:100],
                    "신청": w.get("aplyMtdNm") or "",
                    "상세링크": w.get("servDtlLink") or "",
                })
            dfw = pd.DataFrame(rows)
            cfg = {}
            if "상세링크" in dfw.columns:
                cfg["상세링크"] = st.column_config.LinkColumn("상세링크", display_text="바로가기")
            st.dataframe(dfw, use_container_width=True, hide_index=True, column_config=cfg or None)
            st.caption("※ 수급 가능 여부는 복지로·지자체에서 최종 확인하세요.")


        # X. 지도 — 진로체험·진로교육·청소년활동
        cm_h1("🗺️", "", "지역 교육 · 활동 · 체험 지도")
        sec += 1
        user_region = (st.session_state.get("career_profile") or {}).get("region") or user_region or ""
        if not user_region:
            ipr = st.session_state.get("career_ip_region") or {}
            user_region = ipr.get("region") or ""
        sido_tok = region_sido_token(user_region)
        map_points = []

        # 1) 꿈길: 진로체험 / 진로교육 구분
        for row in (st.session_state.get("career_ggomgil") or []):
            if not isinstance(row, dict):
                continue
            area = str(row.get("체험지역명") or row.get("지역") or "")
            place = str(row.get("체험처명") or row.get("운영기관명") or "")
            title = str(row.get("체험프로그램명") or row.get("프로그램명") or row.get("제목") or place)
            blob = f"{area} {place} {title}"
            if sido_tok and sido_tok not in blob and not region_matches_user(blob, user_region):
                continue
            try:
                bucket = classify_ggomgil_bucket(row)
            except Exception:
                bucket = "체험"
            kind = "진로교육" if bucket == "교육" else "진로체험"
            lat, lng = row.get("위도"), row.get("경도")
            try:
                lat = float(lat) if lat not in (None, "") else None
                lng = float(lng) if lng not in (None, "") else None
            except Exception:
                lat = lng = None
            if not lat or not lng:
                lat, lng, _ = resolve_lat_lon({
                    "addr1": area, "fcltyNm": place or title, "ctpvNm": sido_tok,
                })
            if lat and lng:
                map_points.append({
                    "lat": lat, "lon": lng,
                    "name": title or place,
                    "type": kind,
                    "addr": f"{place} {area}".strip(),
                    "meta": str(row.get("체험유형") or row.get("체험프로그램 직업유형") or ""),
                })

        # 2) 청소년활동
        for row in (st.session_state.get("career_youth") or []):
            if not isinstance(row, dict):
                continue
            info = normalize_youth_program(row)
            blob = " ".join(
                x for x in [info.get("sido"), info.get("sgg"), info.get("addr"), info.get("facility"), info.get("name")] if x
            )
            if blob and sido_tok and not region_matches_user(blob, user_region) and sido_tok not in blob:
                continue
            lat, lon, _ = resolve_lat_lon({
                "fcltyNm": info.get("facility"),
                "addr1": info.get("addr"),
                "ctpvNm": info.get("sido") or sido_tok,
                "sggNm": info.get("sgg"),
                "prgrmNm": info.get("name"),
            })
            if lat and lon:
                map_points.append({
                    "lat": lat, "lon": lon,
                    "name": info.get("name") or info.get("facility") or "활동",
                    "type": "청소년활동",
                    "addr": info.get("addr") or f"{info.get('sido','')} {info.get('sgg','')}".strip(),
                    "meta": info.get("target") or "",
                })

        # 좌표 0건이면 필터 완화
        if not map_points:
            for row in (st.session_state.get("career_ggomgil") or [])[:25]:
                if not isinstance(row, dict):
                    continue
                area = str(row.get("체험지역명") or "")
                place = str(row.get("체험처명") or "")
                title = str(row.get("체험프로그램명") or place)
                try:
                    kind = "진로교육" if classify_ggomgil_bucket(row) == "교육" else "진로체험"
                except Exception:
                    kind = "진로체험"
                lat, lng = row.get("위도"), row.get("경도")
                try:
                    lat = float(lat) if lat not in (None, "") else None
                    lng = float(lng) if lng not in (None, "") else None
                except Exception:
                    lat = lng = None
                if not lat or not lng:
                    lat, lng, _ = resolve_lat_lon({"addr1": area, "fcltyNm": place or title, "ctpvNm": sido_tok})
                if lat and lng:
                    map_points.append({
                        "lat": lat, "lon": lng, "name": title or place, "type": kind,
                        "addr": f"{place} {area}".strip(), "meta": "",
                    })
            for row in (st.session_state.get("career_youth") or [])[:25]:
                if not isinstance(row, dict):
                    continue
                info = normalize_youth_program(row)
                lat, lon, _ = resolve_lat_lon({
                    "fcltyNm": info.get("facility"), "addr1": info.get("addr"),
                    "ctpvNm": info.get("sido") or sido_tok, "sggNm": info.get("sgg"),
                    "prgrmNm": info.get("name"),
                })
                if lat and lon:
                    map_points.append({
                        "lat": lat, "lon": lon,
                        "name": info.get("name") or info.get("facility") or "활동",
                        "type": "청소년활동",
                        "addr": info.get("addr") or "",
                        "meta": info.get("target") or "",
                    })

        if not map_points:
            cm_empty(
                f"지도에 표시할 좌표가 없습니다. (기준 시·도: {sido_tok or user_region or '미확인'}) "
                "꿈길·청소년활동 데이터 또는 지오코딩 키를 확인하세요."
            )
        else:
            import pandas as pd
            # 유형별 스타일
            STYLE = {
                "진로체험": {"color": [37, 99, 235, 230], "radius": 55, "label": "● 진로체험"},
                "진로교육": {"color": [16, 185, 129, 230], "radius": 70, "label": "◆ 진로교육"},
                "청소년활동": {"color": [249, 115, 22, 230], "radius": 48, "label": "▲ 청소년활동"},
            }
            for p in map_points:
                stl = STYLE.get(p.get("type"), STYLE["청소년활동"])
                p["color"] = stl["color"]
                p["radius"] = stl["radius"]

            st.caption(
                f"시·도 기준 **{sido_tok or user_region or '전국'}** · {len(map_points)}곳 · "
                "마커에 마우스를 올리면 프로그램명이 표시됩니다."
            )
            st.markdown(
                '<div style="display:flex;gap:1rem;flex-wrap:wrap;margin:0.4rem 0 0.8rem 0;">'
                '<span style="color:#60a5fa;">● 진로체험</span>'
                '<span style="color:#34d399;">◆ 진로교육</span>'
                '<span style="color:#fb923c;">▲ 청소년활동</span>'
                "</div>",
                unsafe_allow_html=True,
            )

            try:
                import pydeck as pdk
                dfm = pd.DataFrame(map_points)
                mid_lat = float(dfm["lat"].mean())
                mid_lon = float(dfm["lon"].mean())
                layers = []
                for kind, stl in STYLE.items():
                    sub = dfm[dfm["type"] == kind]
                    if sub.empty:
                        continue
                    # 외곽 halo
                    layers.append(pdk.Layer(
                        "ScatterplotLayer",
                        data=sub,
                        get_position="[lon, lat]",
                        get_fill_color=stl["color"][:3] + [90],
                        get_radius=stl["radius"] * 2.2,
                        radius_min_pixels=8,
                        radius_max_pixels=36,
                        pickable=True,
                    ))
                    # 중심 마커 (크기 차이로 모양 구분감)
                    layers.append(pdk.Layer(
                        "ScatterplotLayer",
                        data=sub,
                        get_position="[lon, lat]",
                        get_fill_color=stl["color"],
                        get_radius=stl["radius"],
                        radius_min_pixels=5,
                        radius_max_pixels=18,
                        pickable=True,
                    ))
                tooltip = {
                    "html": (
                        "<div style='font-weight:700;margin-bottom:4px;'>{name}</div>"
                        "<div style='color:#93c5fd;margin-bottom:2px;'>{type}</div>"
                        "<div style='font-size:12px;opacity:0.9'>{addr}</div>"
                        "<div style='font-size:11px;opacity:0.75'>{meta}</div>"
                    ),
                    "style": {
                        "backgroundColor": "rgba(15,23,42,0.96)",
                        "color": "#e2e8f0",
                        "border": "1px solid #38bdf8",
                        "borderRadius": "12px",
                        "padding": "10px 12px",
                        "fontSize": "13px",
                        "boxShadow": "0 8px 24px rgba(0,0,0,0.35)",
                    },
                }
                deck = pdk.Deck(
                    map_style="https://basemaps.cartocdn.com/gl/positron-gl-style/style.json",
                    initial_view_state=pdk.ViewState(
                        latitude=mid_lat, longitude=mid_lon, zoom=10, pitch=40, bearing=-10
                    ),
                    layers=layers,
                    tooltip=tooltip,
                )
                st.pydeck_chart(deck, use_container_width=True)
            except Exception as e:
                st.warning(f"상세 지도 로드 실패: {e}")
                st.map(pd.DataFrame(map_points)[["lat", "lon"]].dropna())

            # 디자인 표 (카드형 HTML)
            st.markdown(
                """
                <style>
                .map-table-wrap { margin-top: 0.75rem; }
                .map-row {
                  display: grid;
                  grid-template-columns: 7.5rem 1fr auto;
                  gap: 0.75rem;
                  align-items: start;
                  background: linear-gradient(180deg, rgba(30,41,59,0.95), rgba(15,23,42,0.92));
                  border: 1px solid rgba(148,163,184,0.22);
                  border-radius: 14px;
                  padding: 0.85rem 1rem;
                  margin-bottom: 0.55rem;
                  box-shadow: 0 4px 14px rgba(0,0,0,0.18);
                }
                .map-badge {
                  display: inline-block;
                  font-size: 0.78rem; font-weight: 700;
                  padding: 0.28rem 0.55rem; border-radius: 999px;
                  white-space: nowrap;
                }
                .badge-exp { background: rgba(37,99,235,0.25); color: #93c5fd; border: 1px solid rgba(96,165,250,0.45); }
                .badge-edu { background: rgba(16,185,129,0.2); color: #6ee7b7; border: 1px solid rgba(52,211,153,0.4); }
                .badge-act { background: rgba(249,115,22,0.2); color: #fdba74; border: 1px solid rgba(251,146,60,0.45); }
                .map-title { font-weight: 700; color: #f8fafc; font-size: 0.98rem; margin-bottom: 0.2rem; }
                .map-sub { color: #94a3b8; font-size: 0.84rem; line-height: 1.45; }
                .map-link a {
                  display: inline-block; padding: 0.35rem 0.7rem; border-radius: 10px;
                  background: rgba(37,99,235,0.25); color: #93c5fd !important;
                  text-decoration: none; font-size: 0.82rem; font-weight: 600;
                  border: 1px solid rgba(96,165,250,0.4);
                }
                </style>
                """,
                unsafe_allow_html=True,
            )
            badge_cls = {"진로체험": "badge-exp", "진로교육": "badge-edu", "청소년활동": "badge-act"}
            html_rows = ['<div class="map-table-wrap">']
            for p in map_points[:40]:
                bcls = badge_cls.get(p.get("type"), "badge-act")
                link = naver_map_link(p["lat"], p["lon"], p["name"]) or naver_map_search_link(p.get("addr") or p["name"])
                meta = p.get("meta") or ""
                addr = p.get("addr") or ""
                sub = " · ".join(x for x in [addr, meta] if x)
                link_html = f'<div class="map-link"><a href="{link}" target="_blank">지도</a></div>' if link else ""
                html_rows.append(
                    f'<div class="map-row">'
                    f'<div><span class="map-badge {bcls}">{p.get("type")}</span></div>'
                    f'<div><div class="map-title">{p.get("name") or "-"}</div>'
                    f'<div class="map-sub">{sub}</div></div>'
                    f'{link_html}'
                    f'</div>'
                )
            html_rows.append("</div>")
            st.markdown("\n".join(html_rows), unsafe_allow_html=True)



else:
    page_header("🔍 준비 중", "추가 도구 개발 중")
    st.info("곧 새로운 기능이 추가될 예정입니다.")
